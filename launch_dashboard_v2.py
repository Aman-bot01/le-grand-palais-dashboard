"""
Game Launch Intelligence — v2 (Nucleus-styled rebuild) -- SAFE CLONE
Run:  streamlit run launch_dashboard_v2.py --server.port 8504

SYNTHETIC-DATA CLONE: this copy runs against a locally generated DuckDB
database of fully made-up games/locations/revenue (see synthetic_data.py) --
it never connects to any real company database. Safe to run, share, or
demo without exposing real business data.
"""
from __future__ import annotations
import datetime as dt
import io
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st

import engine as E
import launch as L

# ─── Page ──────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Le Grand Palais — Casino Intelligence (Synthetic Data)",
    page_icon="🎰",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Design tokens — modern casino: warm ivory, deep gold + forest accents.
# Light (not dark) on purpose: an earlier dark-charcoal pass read as too dark to
# read comfortably, so this keeps the casino gold/serif identity but on a light,
# high-contrast ivory ground (dark text on light bg, same as the original theme).
BG        = "#F8F4EA"
SURFACE   = "#FFFFFF"
S2        = "#F0E8D2"
BORDER    = "#E0D4B0"
TEXT      = "#211C10"
T2        = "#5C5133"
T3        = "#8F7F52"

GOLD      = "#9A7418"
GOLD_LT   = "rgba(154,116,24,0.12)"

RED       = "#A6362B"
RED_LT    = "rgba(166,54,43,0.10)"
AMBER     = "#8A5A0F"
AMBER_LT  = "rgba(138,90,15,0.10)"
GREEN     = "#2E6B4C"
GREEN_LT  = "rgba(46,107,76,0.10)"
# Diverging pair reserved strictly for week-over-week / peer movement.
# Never reused as a general chart series color.
MOVE_UP   = "#2E7D5F"
MOVE_DOWN = "#A6362B"

SERIF = "'Playfair Display', Georgia, serif"
SANS  = "'Inter',-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif"

PLT_FONT = dict(family="Inter,-apple-system,BlinkMacSystemFont,sans-serif", color=T2, size=18)
PLT_COLORWAY = [GOLD, MOVE_UP, MOVE_DOWN, "#6B4FA0", T2]

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@500;600;700&family=Inter:wght@400;500;600;700&display=swap');
html, body, [data-testid="stAppViewContainer"] {{
  background:{BG} !important;
  font-family:{SANS};
  color:{TEXT};
  font-size:19px;
}}
h1,h2,h3 {{ font-family:{SERIF}; font-weight:600; }}
[data-testid="stSidebar"] {{ background:{SURFACE} !important; border-right:1px solid {BORDER}; }}
[data-testid="stSidebar"] * {{ color:{TEXT} !important; }}
.block-container {{ padding-top:0.6rem !important; padding-bottom:1.5rem; max-width:97%; }}
#MainMenu, footer {{ visibility:hidden; }}
div[data-testid="stToolbar"] {{ visibility:hidden; }}
header[data-testid="stHeader"] {{ height:2.2rem; min-height:2.2rem; background:{BG}; }}
header[data-testid="stHeader"] > * {{ height:2.2rem; min-height:2.2rem; }}
div[data-testid="stAppViewContainer"] > section {{ padding-top:0 !important; }}

.eyebrow {{
  font-family:{SANS}; font-size:15px; font-weight:700; text-transform:uppercase;
  letter-spacing:.09em; color:{T3};
}}

/* top header */
.topbar {{
  border-bottom:1px solid {BORDER}; padding:4px 0 16px; margin-bottom:18px;
  display:flex; align-items:flex-end; justify-content:space-between;
}}
.topbar-title {{ font-family:{SERIF}; font-size:34px; font-weight:600; letter-spacing:-.01em; margin-top:4px; }}
.topbar-sub   {{ font-size:17px; color:{T3}; margin-top:2px; }}
.tpill {{
  font-size:15px; font-weight:700; padding:4px 13px; text-transform:uppercase; letter-spacing:.06em;
  border-radius:20px; background:{S2}; color:{T2}; border:1px solid {BORDER};
}}

/* section header */
.ribbon {{ border-left:3px solid var(--rb,{T2}); padding:2px 0 2px 14px; margin:16px 0 10px; }}
.ribbon h3 {{ font-family:{SERIF}; font-size:24px; font-weight:600; color:{TEXT}; margin:3px 0 0; }}
.ribbon p  {{ font-size:18px; color:{T2}; margin:3px 0 0; }}

/* KPI tiles */
.krow {{ display:flex; gap:14px; margin:0 0 16px; flex-wrap:wrap; }}
.kcard {{ flex:1; min-width:130px; background:{SURFACE}; border:1px solid {BORDER}; border-radius:10px; padding:14px 16px; }}
.kl {{ font-size:15px; font-weight:700; text-transform:uppercase; letter-spacing:.07em; color:{T3}; margin-bottom:8px; }}
.kv {{ font-family:{SERIF}; font-size:36px; font-weight:600; color:{TEXT}; line-height:1.1; }}
.kd {{ display:inline-block; font-size:17px; font-weight:600; margin-left:10px; }}
.kd.up   {{ color:{MOVE_UP}; }}
.kd.down {{ color:{MOVE_DOWN}; }}
.ks {{ font-size:17px; color:{T3}; margin-top:8px; }}

/* status badges — green/amber/red reserved for OK/WATCH/FLAG only */
.badge {{ display:inline-block; font-size:15px; font-weight:700; text-transform:uppercase; letter-spacing:.04em;
          padding:3px 10px; border-radius:20px; vertical-align:middle; }}
.b-g {{ background:{GREEN_LT}; color:{GREEN}; }}
.b-a {{ background:{AMBER_LT}; color:{AMBER}; }}
.b-r {{ background:{RED_LT};   color:{RED};   }}
.b-n {{ background:{S2};       color:{T3};    }}

/* leaderboard rows */
.lrow {{ display:flex; align-items:center; gap:12px; padding:9px 6px; border-bottom:1px solid {BORDER}; }}
.lrow:last-child {{ border-bottom:none; }}
.lrow.hl {{ background:{S2}; border-radius:6px; }}
.lrank {{ font-size:18px; color:{T3}; min-width:24px; }}
.lname {{ flex:1; font-weight:600; font-size:18px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }}
.lval  {{ font-size:18px; min-width:80px; text-align:right; }}
.lbar-w {{ width:100px; background:{S2}; border-radius:4px; height:7px; overflow:hidden; }}
.lbar  {{ height:100%; border-radius:4px; background:{T2}; }}

/* data table (st.dataframe wrapper) */
[data-testid="stDataFrame"] {{ border:1px solid {BORDER}; border-radius:8px; font-size:19px; }}
[data-testid="stDataFrame"] * {{ font-size:19px !important; }}
[data-testid="stPlotlyChart"] {{ background:{SURFACE}; border:1px solid {BORDER}; border-radius:10px; padding:8px 10px; }}

/* tabs */
.stTabs [data-baseweb="tab-list"] {{ background:{SURFACE}; border-bottom:1px solid {BORDER}; gap:0; padding:0 4px; }}
.stTabs [data-baseweb="tab"] {{ font-size:18px; font-weight:500; color:{T3}; border-bottom:2px solid transparent; padding:11px 18px; background:transparent !important; margin-bottom:-1px; }}
.stTabs [aria-selected="true"] {{ color:{TEXT} !important; border-bottom-color:{TEXT} !important; font-weight:600 !important; }}

/* widen main content area's use of horizontal space */
[data-testid="stHorizontalBlock"] {{ gap:1.2rem; }}
</style>
""", unsafe_allow_html=True)

pio_template = go.layout.Template(layout=dict(
    font=PLT_FONT, paper_bgcolor=SURFACE, plot_bgcolor=SURFACE,
    colorway=PLT_COLORWAY,
    xaxis=dict(gridcolor=BORDER, zerolinecolor=BORDER, tickfont=dict(size=17)),
    yaxis=dict(gridcolor=BORDER, zerolinecolor=BORDER, tickfont=dict(size=17)),
    hoverlabel=dict(bgcolor=SURFACE, bordercolor=BORDER, font=dict(color=TEXT, size=18)),
))
import plotly.io as pio
pio.templates["nucleus"] = pio_template
pio.templates.default = "nucleus"


# ─── Formatting helpers ─────────────────────────────────────────────────────
def _usd(v, d=0):
    if v is None or (isinstance(v, float) and (np.isnan(v) or np.isinf(v))):
        return "–"
    a = abs(v); s = "-" if v < 0 else ""
    if a >= 1_000_000: return f"{s}${a/1_000_000:.1f}M"
    if a >= 1_000:     return f"{s}${a/1_000:.0f}K"
    return f"{s}${a:,.{d}f}"

def _pct(v, d=1):
    """Format a share/ratio-style percentage. Every caller of this helper is a
    "current value as % of some reference" metric (Hold %, Profit Margin, a
    location's share of total Game Net, ...), never a +/- delta badge (those
    format independently in krow()) -- so clipping the display at just under
    100% here is a safe, universal guard against exactly the kind of
    >100%-reads-as-broken number this was written to catch."""
    if v is None or (isinstance(v, float) and (np.isnan(v) or np.isinf(v))):
        return "–"
    return f"{min(v, 99.9):.{d}f}%"

def badge(text, kind="n"):
    return f'<span class="badge b-{kind}">{text}</span>'

def krow(items):
    """items: list of dicts with label, value, and optional delta_pct, sub."""
    cards = []
    for it in items:
        delta_html = ""
        dp = it.get("delta_pct")
        if dp is not None and not (isinstance(dp, float) and np.isnan(dp)):
            cls = "up" if dp >= 0 else "down"
            arrow = "▲" if dp >= 0 else "▼"
            delta_html = f'<span class="kd {cls}">{arrow} {abs(dp):.1f}%</span>'
        cards.append(
            f'<div class="kcard"><div class="kl">{it["label"]}</div>'
            f'<div class="kv">{it["value"]}{delta_html}</div>'
            + (f'<div class="ks">{it.get("sub","")}</div>' if it.get("sub") else "")
            + '</div>')
    st.markdown(f'<div class="krow">{"".join(cards)}</div>', unsafe_allow_html=True)

def ribbon(title, sub="", color=T2, tag=""):
    html = (f'<div class="ribbon" style="--rb:{color}">'
            + (f'<span class="eyebrow">{tag}</span>' if tag else '')
            + f'<h3>{title}</h3>'
            + (f'<p>{sub}</p>' if sub else '')
            + '</div>')
    st.markdown(html, unsafe_allow_html=True)


def fill_band(fig, x, lo, hi, color, name, showlegend=True):
    fig.add_trace(go.Scatter(
        x=list(x) + list(x)[::-1], y=list(hi) + list(lo)[::-1],
        fill="toself", fillcolor=color, line=dict(width=0),
        mode="lines", name=name, showlegend=showlegend, hoverinfo="skip",
    ))

# ─── What's New — SQL loaders, copied verbatim from launch_dashboard.py ────
# (can't import that file directly — it executes a full Streamlit app on import)
@st.cache_data(ttl=86400, show_spinner="Loading latest releases…")
def load_whats_new(days_back: int = 90) -> pd.DataFrame:
    conn = E.get_connection()
    sql = f"""
    SELECT
        loc.ConfigStudio                              AS Studio,
        loc.ConfigPlatform                            AS ConfigPlatform,
        CASE
            WHEN loc.ConfigPlatform = 'PFH' AND loc.ConfigProduct = 'PFH + Sweeps' AND loc.Kiosk = 1 THEN 'Kiosk Only'
            WHEN loc.ConfigPlatform = 'PFH' AND loc.ConfigProduct = 'Kiosk Only'   AND loc.Kiosk = 1 THEN 'Kiosk Only'
            WHEN loc.ConfigPlatform = 'PFH' AND loc.ConfigProduct = 'PFH + Sweeps' AND loc.Kiosk = 0 THEN 'PFH + Sweeps'
            WHEN loc.ConfigPlatform = 'PFH' AND loc.ConfigProduct = 'PFH Only'     AND loc.Kiosk = 0 THEN 'PFH Only'
            WHEN loc.ConfigPlatform = 'V2'  AND loc.ConfigProduct = 'P2P'          AND loc.Kiosk = 0 THEN 'P2P'
            WHEN loc.ConfigPlatform = 'V2'  AND loc.ConfigProduct = 'PullTabs'     AND loc.Kiosk = 0 THEN 'PullTabs'
            WHEN loc.ConfigPlatform = 'V2'  AND loc.ConfigProduct = 'Class 2'      AND loc.Kiosk = 0 THEN 'Class 2'
            WHEN loc.ConfigPlatform = 'V2'  AND loc.ConfigProduct = 'HHR'          AND loc.Kiosk = 0 THEN 'HHR'
            WHEN loc.ConfigPlatform = 'V2'  AND loc.ConfigProduct = 'Sweeps'       AND loc.Kiosk = 0 THEN 'Sweeps'
            WHEN loc.ConfigPlatform = 'V1'  AND loc.ConfigProduct = 'Sweeps'       AND loc.Kiosk = 0 THEN 'Sweeps'
            WHEN loc.ConfigPlatform = 'V1'  AND loc.ConfigProduct = 'PFH Only'     AND loc.Kiosk = 0 THEN 'Sweeps'
            WHEN loc.ConfigPlatform = 'V1'  AND loc.ConfigProduct = 'P2P'          AND loc.Kiosk = 0 THEN 'P2P'
            WHEN loc.ConfigPlatform = 'V1'  AND loc.ConfigProduct = 'Got Skill'    AND loc.Kiosk = 0 THEN 'P2P'
            WHEN loc.ConfigPlatform = 'UNKNOWN' AND loc.ConfigProduct = 'Sweeps'   AND loc.Kiosk = 0 THEN 'Sweeps'
            ELSE NULL
        END                                           AS RequiredProduct,
        loc.Distributor,
        loc.AccountManager,
        u.Category,
        u.Action,
        u.Note,
        MAX(u.Date)                                   AS LastDate,
        COUNT(DISTINCT u.LocationId)                  AS Locations
    FROM CrmUpdateLogView u
    LEFT JOIN CrmLocationView loc ON u.LocationId = loc.LocationId
    WHERE u.Date >= DATEADD('day', -{days_back}, GETDATE())
    GROUP BY
        loc.ConfigStudio, loc.ConfigPlatform, loc.ConfigProduct, loc.Kiosk,
        loc.Distributor, loc.AccountManager, u.Category, u.Action, u.Note
    ORDER BY
        loc.ConfigStudio, loc.ConfigPlatform, u.Category, u.Action, MAX(u.Date) DESC
    """
    df_wn = E.query_df(conn, sql)
    conn.close()
    df_wn["LastDate"] = pd.to_datetime(df_wn["LastDate"])
    df_wn["PlatformProduct"] = df_wn["ConfigPlatform"].fillna("–")
    return df_wn

@st.cache_data(ttl=3600, show_spinner="Loading land-based game releases…")
def load_game_releases() -> pd.DataFrame:
    # SIMPLIFIED FOR THE DUCKDB CLONE: the original T-SQL query here used
    # CROSS APPLY STRING_SPLIT(...) + OUTER APPLY ... TOP 1 ... ORDER BY <priority
    # CASE> to (a) split multi-game '|'-delimited CrmUpdateLogView.Note values into
    # one release row per game, and (b) pick the single best GameCatalogView1 match
    # by name when several catalog rows share a name (e.g. a base game + its HR
    # test-rig variant). DuckDB has no APPLY operator and no STRING_SPLIT-as-a-
    # table-source. The synthetic CrmUpdateLogView this clone ships never writes
    # pipe-delimited multi-game Notes, so (a) is dropped entirely (every Note is
    # already one game); (b) is reproduced with a standard
    # ROW_NUMBER() OVER (PARTITION BY ... ORDER BY <same priority CASE>) + WHERE rn=1,
    # which is equivalent to "TOP 1 ... ORDER BY" per group. Output columns are
    # unchanged from the original.
    conn = E.get_connection()
    sql = """
    WITH fe AS (
        SELECT loc.ConfigStudio AS Studio, loc.ConfigPlatform AS Platform,
               u.Note, u.Platform AS CrmPlatform, MIN(u.Date) AS FirstEnableDate
        FROM CrmUpdateLogView u
        LEFT JOIN CrmLocationView loc ON u.LocationId = loc.LocationId
        WHERE u.Category = 'Game' AND u.Action = 'Enable'
        GROUP BY loc.ConfigStudio, loc.ConfigPlatform, u.Note, u.Platform
    ),
    ll AS (
        SELECT fe.Studio, fe.Platform, fe.Note, fe.CrmPlatform, fe.FirstEnableDate,
               COUNT(DISTINCT u2.LocationId) AS LocationsAtLaunch
        FROM fe
        JOIN CrmUpdateLogView u2
            ON  u2.Note = fe.Note AND u2.Action = 'Enable' AND u2.Category = 'Game'
            AND CAST(u2.Date AS DATE) = CAST(fe.FirstEnableDate AS DATE)
        GROUP BY fe.Studio, fe.Platform, fe.Note, fe.CrmPlatform, fe.FirstEnableDate
    ),
    cat AS (
        SELECT Name, Id AS GameId, Type AS GameType,
            CASE Product
                WHEN 'p2p' THEN 'PTP' WHEN 'hhr' THEN 'PTP' WHEN 'pulltabs' THEN 'PULL'
                WHEN 'sweeps' THEN 'SWPS' WHEN 'class2' THEN 'CLS2' WHEN 'gotskill' THEN 'GSKL'
                ELSE 'OTHER'
            END AS CrmGroup,
            CASE Product WHEN 'p2p' THEN 1 ELSE 2 END AS Prio
        FROM GameCatalogView1
    ),
    matched AS (
        SELECT ll.Studio, ll.Platform, ll.Note, ll.CrmPlatform, ll.FirstEnableDate,
               ll.LocationsAtLaunch, c.GameId, c.GameType, c.CrmGroup,
            ROW_NUMBER() OVER (
                PARTITION BY ll.Studio, ll.Platform, ll.Note, ll.CrmPlatform, ll.FirstEnableDate
                ORDER BY
                    CASE
                        WHEN ll.CrmPlatform = 'V2 Pay to Play' AND c.CrmGroup = 'PTP'  THEN 1
                        WHEN ll.CrmPlatform = 'V2 Pull-Tabs'   AND c.CrmGroup = 'PULL' THEN 1
                        WHEN ll.CrmPlatform IN ('V1 Sweeps', 'V2 Sweeps', 'PFH Sweeps') AND c.CrmGroup = 'SWPS' THEN 1
                        WHEN ll.CrmPlatform = 'V2 Class 2'     AND c.CrmGroup = 'CLS2' THEN 1
                        WHEN ll.CrmPlatform = 'V1 Got Skill'   AND c.CrmGroup = 'GSKL' THEN 1
                        ELSE 2
                    END, c.Prio, c.GameId
            ) AS rn
        FROM ll
        LEFT JOIN cat c ON LOWER(TRIM(c.Name)) = LOWER(TRIM(ll.Note))
    )
    SELECT
        m.Studio, m.Platform AS PlatformName, m.CrmPlatform, m.GameId,
        m.GameType AS Config, m.CrmGroup, m.Note AS GameName,
        m.FirstEnableDate, m.LocationsAtLaunch, th.Products
    FROM matched m
    LEFT JOIN (
        SELECT dp.GameId, STRING_AGG(dp.ProductName, ', ') AS Products
        FROM (
            SELECT DISTINCT GameId, ProductName FROM TaskHandlerBetSpinSummary
            WHERE CasinoName IN ('vendor1', 'vendor2') AND ProductName IS NOT NULL
        ) dp
        GROUP BY dp.GameId
    ) th ON th.GameId = m.GameId
    WHERE m.rn = 1
    ORDER BY m.FirstEnableDate DESC
    """
    df = E.query_df(conn, sql)
    conn.close()
    df["FirstEnableDate"] = pd.to_datetime(df["FirstEnableDate"])
    df["LastDate"] = df["FirstEnableDate"]
    df["Locations"] = df["LocationsAtLaunch"]
    df["Note"] = df["GameName"]
    df["PlatformProduct"] = df["PlatformName"].fillna("–")
    df["Category"] = "Game"
    return df

@st.cache_data(ttl=3600, show_spinner="Loading EdgeLabs releases…")
def load_edgelabs_releases(days_back: int = 90) -> pd.DataFrame:
    try:
        conn = E.get_connection()
        sql = f"""
        SELECT g.Id, g.Name, g.Type, fs.PlatformName, fs.FirstSpinDate
        FROM (
          SELECT TRY_CAST(GameId AS INT) AS GameIdInt, PlatformName, MIN(Date) AS FirstSpinDate
          FROM BetSpinSummaryCashView3
          WHERE FreeGameCampaignId IS NULL
          GROUP BY TRY_CAST(GameId AS INT), PlatformName
        ) fs
        JOIN GameCatalogView1 g ON fs.GameIdInt = g.Id
        WHERE fs.FirstSpinDate >= DATEADD('day', -{days_back}, GETDATE())
          AND fs.FirstSpinDate <= GETDATE()
        ORDER BY fs.FirstSpinDate
        """
        df_el = E.query_df(conn, sql)
        conn.close()
        if df_el.empty:
            return pd.DataFrame()
        df_el["FirstSpinDate"] = pd.to_datetime(df_el["FirstSpinDate"])
        df_el = df_el.rename(columns={"Name": "GameName", "FirstSpinDate": "LastDate", "Id": "GameId", "Type": "Config"})
        df_el["Studio"] = df_el["PlatformName"].fillna("Online")
        df_el["PlatformProduct"] = df_el["PlatformName"]
        df_el["Note"] = df_el["GameName"]
        df_el["Locations"] = 0
        df_el["Products"] = None
        df_el["Category"] = "Game"
        return df_el[["Studio", "PlatformProduct", "PlatformName", "GameId", "Config",
                       "GameName", "Note", "LastDate", "Locations", "Products", "Category"]].reset_index(drop=True)
    except Exception:
        return pd.DataFrame()

# ─── Peer matching + Quick Score helpers, copied verbatim from launch_dashboard.py ──
def load_sql_catalog() -> pd.DataFrame:
    """Load GameCatalogView1 via SQL; returns empty DataFrame on failure."""
    try:
        conn = E.get_connection()
        cat_sql = E.load_game_catalog_with_fallback(conn)
        conn.close()
        return cat_sql
    except Exception:
        return pd.DataFrame()

@st.cache_data(ttl=86400, show_spinner=False)
def find_peers_scaled(df, target_id, kpi="bet_decay", n_weeks=None, top_k=5,
                       exclude_hr=True, scale_tolerance=2.0, min_candidate_weeks=4):
    """Mechanic-aware peer matching via find_peers_v2; falls back to legacy find_peers
    if the SQL catalog is unavailable."""
    catalog = load_sql_catalog()
    if not catalog.empty and "Id" in catalog.columns:
        result = L.find_peers_v2(
            df, target_id, catalog, kpi=kpi, n_weeks=n_weeks, top_k=top_k,
            exclude_hr=exclude_hr, scale_tolerance=scale_tolerance,
            min_candidate_weeks=min_candidate_weeks,
        )
        peers = result["peers"]
        for p in peers:
            p["scale_constrained"] = result["mechanic_matched"]
            # Which tier actually matched (skin family / platform+product+orientation / ...
            # / loose) — previously computed by find_peers_v2 and silently dropped here,
            # so the UI never showed how thin or solid a peer set really was.
            p["peer_family"] = result["family"]
            p["peer_fallback_reason"] = result["fallback_reason"]
        return peers
    w0 = df[df["launch_week"] == 0].set_index("game_id")["bet_handle"]
    target_w0 = float(w0.get(target_id, np.nan))
    if np.isnan(target_w0) or target_w0 <= 0:
        res = L.find_peers(df, target_id, kpi=kpi, n_weeks=n_weeks, top_k=top_k,
                            exclude_hr=exclude_hr, min_candidate_weeks=min_candidate_weeks)
        for r in res:
            r["scale_constrained"] = False
            r["peer_family"] = "whole fleet (no SQL catalog)"
            r["peer_fallback_reason"] = "SQL catalog unavailable — no attribute matching possible"
        return res
    lo, hi = target_w0 / scale_tolerance, target_w0 * scale_tolerance
    scaled_ids = set(w0[(w0 >= lo) & (w0 <= hi)].index.tolist())
    df_sc = df[df["game_id"].isin(scaled_ids)]
    res = L.find_peers(df_sc, target_id, kpi=kpi, n_weeks=n_weeks, top_k=top_k,
                        exclude_hr=exclude_hr, min_candidate_weeks=min_candidate_weeks)
    if len(res) < max(2, top_k // 2):
        res = L.find_peers(df, target_id, kpi=kpi, n_weeks=n_weeks, top_k=top_k,
                            exclude_hr=exclude_hr, min_candidate_weeks=min_candidate_weeks)
        for r in res:
            r["scale_constrained"] = False
            r["peer_family"] = "whole fleet (no SQL catalog, scale too thin)"
            r["peer_fallback_reason"] = "SQL catalog unavailable and scale-matched pool was too small"
    else:
        for r in res:
            r["scale_constrained"] = True
            r["peer_family"] = "scale-matched (no SQL catalog)"
            r["peer_fallback_reason"] = "SQL catalog unavailable — matched by launch scale only"
    return res

@st.cache_data(ttl=3600, show_spinner="Loading player counts…")
def load_edgelabs_player_weeks(game_id: int, launch_date: str, platform: str = "EdgeLabs") -> pd.DataFrame:
    try:
        conn = E.get_connection()
        sql = f"""
        SELECT DATEDIFF('week', '{launch_date}', CAST("Date" AS DATE)) AS launch_week,
            COUNT(DISTINCT AccountNumber) AS unique_players
        FROM BetSpinSummaryCashView3
        WHERE PlatformName = '{platform}' AND TRY_CAST(GameId AS INT) = {int(game_id)}
          AND FreeGameCampaignId IS NULL AND AccountNumber IS NOT NULL
        GROUP BY DATEDIFF('week', '{launch_date}', CAST("Date" AS DATE))
        ORDER BY launch_week
        """
        d = E.query_df(conn, sql); conn.close()
        return d
    except Exception:
        return pd.DataFrame()

def _tenure_band(weeks_live: int) -> str:
    """🌿 New Release (<13 wks) / 🌳 Floor Regular (13–52 wks) / 🌲 House Classic (52+ wks) — see
    feedback_report_analysis_format memory. Gates which of the other metrics on this
    page are trustworthy, per our standing rule: don't score a metric that needs
    history the game doesn't have yet, flag it as insufficient instead."""
    if weeks_live < 13:
        return "🌿 New Release"
    if weeks_live < 52:
        return "🌳 Floor Regular"
    return "🌲 House Classic"

_band_cache: dict = {}
def bkpi(peer_df, kpi: str) -> pd.DataFrame:
    """Peer percentile bands for `kpi` over peer_df, computed once and cached per session.

    Cache key is content-based (the actual set of peer game_ids + row count), NOT
    id(peer_df) -- Python reuses memory addresses after garbage collection, so an
    id()-based key can silently collide across different DataFrames in a tight loop
    and return a stale/wrong game's percentile bands. See reference_bkpi_id_cache_bug.
    """
    if peer_df.empty or kpi not in peer_df.columns:
        return pd.DataFrame()
    cache_key = (tuple(sorted(peer_df["game_id"].unique())), len(peer_df), kpi)
    if cache_key not in _band_cache:
        _band_cache[cache_key] = L.bands_kpi(peer_df, kpi)
    return _band_cache[cache_key]

def _flag_cls(actual, p25, p75, p10=None, p90=None, higher_is_better=True):
    if actual is None or (isinstance(actual, float) and np.isnan(actual)):
        return "n"
    if p10 is not None and p90 is not None:
        if actual < p10: return "r" if higher_is_better else "g"
        if actual > p90: return "g" if higher_is_better else "r"
    if actual < p25: return "a" if higher_is_better else "g"
    if actual > p75: return "g" if higher_is_better else "a"
    return "g"

def _classify_game_platform(pp: str):
    p = str(pp).strip(); pl = p.lower()
    if "edgelabs" in pl: return ("EdgeLabs", "EdgeLabs")
    if p == "Pong" or pl == "pong": return ("PFH", "PFH · Pong")
    if "pfh" in pl: return ("PFH", "PFH · Sweeps")
    if pl.startswith("v1"):
        if "got skill" in pl or "gotskill" in pl: return ("V1", "Got Skill")
        return ("V1", "Sweeps")
    if pl.startswith("v2"):
        if "pull" in pl: return ("V2", "Pull-Tabs")
        if "pay to play" in pl or "paytoplay" in pl: return ("V2", "Pay to Play")
        if "class 2" in pl or "class2" in pl: return ("V2", "Class 2")
        if "hhr" in pl or "horse" in pl: return ("V2", "HHR")
        if "sweeps" in pl: return ("V2", "Sweeps")
        return ("V2", p)
    return ("Uncategorized", p)

def plotly_base(fig, h=425, ml=48, mr=14, mt=10, mb=40):
    fig.update_layout(
        height=h, margin=dict(l=ml, r=mr, t=mt, b=mb), template="nucleus",
        # Explicit, not just via the template -- this plotly.js build doesn't reliably
        # merge template-level paper_bgcolor/plot_bgcolor, so charts that only set the
        # template (not an explicit value) were rendering with plotly's own white
        # default. SURFACE matches the [data-testid="stPlotlyChart"] container CSS so
        # the chart blends into its card instead of sitting in a white box.
        paper_bgcolor=SURFACE, plot_bgcolor=SURFACE,
        font=PLT_FONT,
        xaxis=dict(automargin=True, gridcolor=BORDER, zerolinecolor=BORDER),
        yaxis=dict(automargin=True, gridcolor=BORDER, zerolinecolor=BORDER),
        legend=dict(font=dict(size=10), bgcolor="rgba(0,0,0,0)", orientation="h",
                    yanchor="bottom", y=1.02, xanchor="right", x=1),
    )
    return fig


def _hex_lerp(c1, c2, t):
    r1, g1, b1 = int(c1[1:3], 16), int(c1[3:5], 16), int(c1[5:7], 16)
    r2, g2, b2 = int(c2[1:3], 16), int(c2[3:5], 16), int(c2[5:7], 16)
    r, g, b = (int(a + (bb - a) * t) for a, bb in ((r1, r2), (g1, g2), (b1, b2)))
    return f"#{r:02x}{g:02x}{b:02x}"


def _hex_alpha(c, a):
    r, g, b = int(c[1:3], 16), int(c[3:5], 16), int(c[5:7], 16)
    return f"rgba({r},{g},{b},{a})"


def _treemap_fig(labels, values, hover_prefix="$", height=260):
    """Composition/concentration chart — replaces pies with many/uneven slices,
    where relative size is easier to read as tile area than as a wedge angle."""
    n = len(labels)
    cols = [PLT_COLORWAY[i % len(PLT_COLORWAY)] for i in range(n)]
    fig = go.Figure(go.Treemap(
        labels=labels, parents=[""] * n, values=list(values),
        marker=dict(colors=cols, line=dict(width=1, color=SURFACE)),
        textinfo="label+percent root", textfont=dict(size=13, color=SURFACE),
        hovertemplate=f"<b>%{{label}}</b><br>{hover_prefix}%{{value:,.0f}} (%{{percentRoot}})<extra></extra>",
    ))
    fig.update_layout(margin=dict(l=4, r=4, t=8, b=8), height=height,
                       paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)")
    return fig


def _ridgeline_fig(groups, unit="", height=330, color_lo=T3, color_hi=TEXT):
    """groups: [(label, values-array), ...] ordered oldest→newest, rendered bottom→top.
    Shows how a metric's whole-fleet distribution shape shifts across tenure —
    a box plot only shows one snapshot, this shows the shift itself."""
    from scipy.stats import gaussian_kde
    all_vals = np.concatenate([np.asarray(v, dtype=float) for _, v in groups])
    lo, hi = np.nanpercentile(all_vals, 1), np.nanpercentile(all_vals, 99)
    if not np.isfinite(lo) or not np.isfinite(hi) or lo == hi:
        lo, hi = float(np.nanmin(all_vals)), float(np.nanmax(all_vals) + 1)
    x_grid = np.linspace(lo, hi, 200)
    n = len(groups)
    gap = 1.0
    fig = go.Figure()
    for i, (label, vals) in enumerate(groups):
        y0 = i * gap
        try:
            dens = gaussian_kde(np.asarray(vals, dtype=float))(x_grid)
        except Exception:
            continue
        dmax = dens.max()
        dens = (dens / dmax * gap * 0.92) if dmax > 0 else dens
        color = _hex_lerp(color_lo, color_hi, i / max(n - 1, 1))
        fig.add_trace(go.Scatter(x=x_grid, y=[y0] * len(x_grid), mode="lines",
                                  line=dict(width=0), showlegend=False, hoverinfo="skip"))
        fig.add_trace(go.Scatter(x=x_grid, y=y0 + dens, mode="lines", line=dict(color=color, width=1.6),
                                  fill="tonexty", fillcolor=_hex_alpha(color, 0.30), name=label,
                                  showlegend=False, hovertemplate=f"{label}<br>%{{x:.1f}}{unit}<extra></extra>"))
    fig.update_yaxes(tickvals=[i * gap for i in range(n)], ticktext=[g[0] for g in groups],
                      showgrid=False, zeroline=False, automargin=True)
    fig.update_xaxes(title_text=f"Value ({unit})" if unit else "Value", automargin=True)
    fig.update_layout(height=height, margin=dict(l=60, r=14, t=10, b=40),
                       paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)")
    return fig


# ─── Cached data loaders (reuse engine.py / launch.py unchanged) ───────────
@st.cache_data(ttl=86400, show_spinner="Loading platform data…")
def load_data(platform):
    conn = E.get_connection()
    df = L.load_platform_data(conn, platform)
    conn.close()
    return df


# Gaming-floor display names — the underlying value ("PFH"/"V2"/"V1"/"EdgeLabs")
# still drives every SQL query in engine.py/launch.py unchanged; this is purely
# a friendlier label so the sidebar reads like a casino floor plan, not a
# system name, without touching the ~30 queries keyed on the real values.
FLOOR_LABELS = {
    "V2": "Main Gaming Floor",
    "V1": "Heritage Floor",
    "PFH": "VIP Kiosk Network",
    "EdgeLabs": "Private Members' Club",
}

# Regulatory/technical product codes (p2p, hhr, gotskill, ...) -- display-only
# translation to plain game-category names. The underlying values are untouched
# (they're matched literally in ~15 SQL CASE WHEN clauses), this just relabels
# them wherever they're shown to a user. Covers both casings that show up in the
# data (GameCatalogView1.Product is lowercase, CrmLocationView.ConfigProduct and
# the What's New CASE WHEN output are Title Case).
GENERIC_PRODUCT_LABELS = {
    "p2p": "Classic", "P2P": "Classic",
    "sweeps": "Sweepstakes", "Sweeps": "Sweepstakes", "PFH + Sweeps": "Sweepstakes",
    "pulltabs": "Instant Win", "PullTabs": "Instant Win",
    "class2": "Bingo-Style", "Class 2": "Bingo-Style",
    "hhr": "Racing Game", "HHR": "Racing Game",
    "gotskill": "Skill Game", "Got Skill": "Skill Game",
    "pfh-edgelabs": "Digital Play",
    "Kiosk Only": "Kiosk", "PFH Only": "Standard",
    "PFH · Pong": "Kiosk", "PFH · Sweeps": "Sweepstakes",
    "Pull-Tabs": "Instant Win", "Pay to Play": "Classic",
}


def _generic_product(series):
    """Map a Series of raw product codes to generic display names, leaving any
    unmapped value (e.g. already-generic text, or 'Unknown') as-is."""
    return series.map(lambda v: GENERIC_PRODUCT_LABELS.get(v, v))


_AGS_PLATFORMS = ("PFH", "V1", "V2", "EdgeLabs")

@st.cache_data(ttl=86400, show_spinner="Loading all platforms…")
def _combined_platform_df():
    """Weekly-grain df across all 4 platforms, with a platform-qualified game_id so the
    same downstream groupby/merge-on-game_id logic All Games Health already uses stays
    correct without risking the cross-platform id-collision bug found in What's New."""
    parts = []
    for p in _AGS_PLATFORMS:
        try:
            d = load_data(p)
        except Exception:
            d = pd.DataFrame()
        if not d.empty:
            d = d.copy()
            d["game_id"] = p + "_" + d["game_id"].astype(str)
            parts.append(d)
    return pd.concat(parts, ignore_index=True) if parts else pd.DataFrame()

def _combined_catalog_from_data(cdf):
    """Same shape/logic as launch.catalog_from_data(), for a combined multi-platform df whose
    game_id is platform-qualified (e.g. "V2_9581") — can't call the original unchanged since
    its is_hr check (`int(game_id) >= 95000`) needs the raw numeric id, not the qualified string."""
    if cdf.empty:
        return pd.DataFrame()
    grp = cdf.groupby(["game_id", "game_name", "codebase", "launch_date"])
    cat = grp.agg(
        total_weeks=("launch_week", "max"),
        total_bet=("bet_handle", "sum"),
        total_net=("net_rev", "sum"),
        last_week_bet=("bet_handle", lambda x: x.iloc[-1] if len(x) else np.nan),
    ).reset_index()
    cat["is_hr"] = cat["game_id"].str.rsplit("_", n=1).str[-1].astype(int) >= 95000
    return cat.sort_values("game_name")

# ─── Sidebar ────────────────────────────────────────────────────────────────
# Le Grand Palais is one physical casino chain — there's no "which system" question
# for someone picking a game to look at, just the game itself. The dropdown below is
# built from every game across every underlying system, combined into one list by
# name. Internally, picking a specific game still resolves to the one system it
# actually runs on (needed to route the SQL underneath) via FLOOR_LABELS/_GC_PLATFORM
# — that routing is never shown to the user as a choice they have to make.
_DEFAULT_PLATFORM = "V2"  # fallback system for fleet-wide ("All Games") queries only

with st.sidebar:
    st.markdown(
        f'<div style="padding:6px 0 12px">'
        f'<div style="font-family:{SERIF};font-size:20px;font-weight:600;color:{GOLD};line-height:1.2;letter-spacing:.02em">LE GRAND PALAIS</div>'
        f'<div class="eyebrow" style="margin-top:2px">Casino Intelligence</div>'
        f'</div>', unsafe_allow_html=True)

    full_df = _combined_platform_df()
    full_cat = _combined_catalog_from_data(full_df)
    full_nonhr = full_cat[~full_cat["is_hr"]].copy()

    _ALL = "— All Games —"
    _sorted_full = full_nonhr.sort_values("launch_date", ascending=False)
    opts_labels = [_ALL] + [f"{r.game_name}  ·  launched {r.launch_date}" for r in _sorted_full.itertuples()]
    opts_qids = [None] + _sorted_full["game_id"].tolist()  # e.g. "V2_9581"
    sel_label = st.selectbox("Game", opts_labels, key="game_sel_v2")
    sel_qid = opts_qids[opts_labels.index(sel_label)]
    is_all = sel_qid is None

    if is_all:
        platform = _DEFAULT_PLATFORM
        sel_id = None
    else:
        platform, _raw_id = sel_qid.split("_", 1)
        sel_id = int(_raw_id)
    floor_label = FLOOR_LABELS.get(platform, platform)

    df = load_data(platform)
    if df.empty:
        st.error("No data returned.")
        st.stop()
    cat = L.catalog_from_data(df)
    nonhr = cat[~cat["is_hr"]].copy()

    st.divider()
    z = st.slider("Forecast Sensitivity", 1.0, 3.0, 1.64, 0.04, key="z_v2",
                  help="Lower = tighter expected range, flags more weeks. Higher = more relaxed.")
    horizon = st.slider("Weeks to forecast ahead", 4, 26, 13, key="hz_v2")
    scale_x = st.slider("Match games of similar size ×", 1.5, 5.0, 2.0, 0.5, key="scx_v2",
                        help="Only match games whose Week-0 wager is within this multiple of yours")
    n_match = st.slider("How many weeks to match on", 2, 16, 8, key="nm_v2")

    st.divider()
    st.caption(f"{len(full_nonhr)} games across the property")
    st.caption("Values in USD · week 0 = launch week")

# ─── Top bar ────────────────────────────────────────────────────────────────
if is_all:
    top_title, top_sub = "All Games", f"{len(full_nonhr)} games across the property"
else:
    meta = nonhr[nonhr["game_id"] == sel_id].iloc[0]
    top_title, top_sub = meta['game_name'], f"Launched {meta['launch_date']}"
    # Shared single-game shortcuts — every single-game tab reuses these instead of recomputing.
    gdf = df[df["game_id"] == sel_id].sort_values("launch_week").reset_index(drop=True)
    max_wk = int(gdf["launch_week"].max()) if not gdf.empty else 0
    w0_bet = float(gdf.loc[gdf["launch_week"] == 0, "bet_handle"].sum() or 0)

st.markdown(f"""
<div class="topbar">
  <div>
    <span class="eyebrow">Trends</span>
    <div class="topbar-title">{top_title}</div>
    <div class="topbar-sub">{top_sub}</div>
  </div>
</div>""", unsafe_allow_html=True)

# ─── Tabs (built incrementally) ─────────────────────────────────────────────
(tab1, tab_weekly, tab_new, tab_overview, tab_track, tab_similar, tab_full, tab_social,
 tab_compare, tab_clusters) = st.tabs(
    ["All Games Health", "Weekly Games", "What's New", "Location Overview", "Is It On Track?",
     "Similar Launches & What to Expect", "Full Breakdown", "Loyalty & Members", "Compare",
     "Game Clusters"])

_CLUSTER_META = {
    0: {"name": "Stars", "color": GREEN, "desc": "High hold %, strong retention, consistent revenue growth"},
    1: {"name": "Steady Earners", "color": T2, "desc": "Solid launch scale, stable hold, moderate decay"},
    2: {"name": "Quick Peaks", "color": AMBER, "desc": "Strong launch handle but faster decay — needs engagement strategy"},
    3: {"name": "Developing", "color": "#7A6A9C", "desc": "Newer or smaller launches, building performance over time"},
}

# ══════════════════════════════════════════════════════════════════
# TAB — BUSINESS OVERVIEW
# ══════════════════════════════════════════════════════════════════
_HS_PERIODS = ["Lifetime", "Last Week", "MTD", "Last Month", "QTD", "YTD"]
_LIFETIME_START = dt.date(2000, 1, 1)  # PFH/V1/V2 data doesn't predate this by years, so it's a safe "all of it" floor

def _hs_get_dates(period_label):
    today = dt.date.today()
    wd = today.weekday()
    tw_start = today - dt.timedelta(days=wd)
    lw_start = tw_start - dt.timedelta(days=7)
    lw_end = tw_start - dt.timedelta(days=1)
    if period_label == "Lifetime":
        cur, pri = (_LIFETIME_START, today), (_LIFETIME_START, today)
    elif period_label == "This Week":
        cur, pri = (tw_start, today), (lw_start, lw_end)
    elif period_label == "Last Week":
        cur = (lw_start, lw_end)
        pri = (lw_start - dt.timedelta(days=7), lw_start - dt.timedelta(days=1))
    elif period_label == "MTD":
        mtd_start = today.replace(day=1)
        lm_end = mtd_start - dt.timedelta(days=1)
        lm_start = lm_end.replace(day=1)
        days_in = (today - mtd_start).days
        cur, pri = (mtd_start, today), (lm_start, lm_start + dt.timedelta(days=days_in))
    elif period_label == "Last Month":
        lm_end = today.replace(day=1) - dt.timedelta(days=1)
        lm_start = lm_end.replace(day=1)
        llm_end = lm_start - dt.timedelta(days=1)
        llm_start = llm_end.replace(day=1)
        cur, pri = (lm_start, lm_end), (llm_start, llm_end)
    elif period_label == "QTD":
        q_month = ((today.month - 1) // 3) * 3 + 1
        qtd_start = today.replace(month=q_month, day=1)
        pq_end = qtd_start - dt.timedelta(days=1)
        pq_start = pq_end.replace(month=((pq_end.month - 1) // 3) * 3 + 1, day=1)
        days_in = (today - qtd_start).days
        cur, pri = (qtd_start, today), (pq_start, pq_start + dt.timedelta(days=days_in))
    else:  # YTD
        ytd_start = today.replace(month=1, day=1)
        ly_start = ytd_start.replace(year=today.year - 1)
        ly_equiv = ly_start + dt.timedelta(days=(today - ytd_start).days)
        cur, pri = (ytd_start, today), (ly_start, ly_equiv)
    return cur, pri

@st.cache_data(ttl=1800, show_spinner="Loading period data…")
def _load_period_pfh(start, end):
    sql = f"""
    SELECT TRY_CAST(b.GameId AS INT) AS game_id, MAX(gc.Name) AS game_name,
        SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
        SUM(CAST(b.TotalBet AS FLOAT)/100.0) - SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS net_rev,
        COUNT(DISTINCT b.StoreNumber) AS stores, SUM(b.Spins) AS spins
    FROM BetSpinSummaryCashView3Pong b
    LEFT JOIN GameCatalogView1 gc ON gc.Id = TRY_CAST(b.GameId AS INT)
    WHERE CAST(b."Date" AS DATE) BETWEEN '{start}' AND '{end}' AND TRY_CAST(b.GameId AS INT) IS NOT NULL
    GROUP BY TRY_CAST(b.GameId AS INT)
    """
    c = E.get_connection(); d = E.query_df(c, sql); c.close()
    return d

@st.cache_data(ttl=1800, show_spinner="Loading period data…")
def _load_period_v2v1(platform_code, start, end):
    gc_plat = L._GC_PLATFORM.get(platform_code, platform_code)
    sql = f"""
    SELECT gc.Id AS game_id, gc.Name AS game_name,
        SUM(CAST(g.TotalPlay AS FLOAT)/100.0) AS bet,
        SUM(CAST(g.TotalPlay AS FLOAT)/100.0) - SUM(CAST(g.TotalWin AS FLOAT)/100.0) AS net_rev,
        COUNT(DISTINCT g.SummaryLocationId) AS stores, SUM(g.PlayCount) AS spins
    FROM AnalyticsGameTerminalsGames g
    JOIN GameCatalogView1 gc ON gc.Id = g.Id
    WHERE gc.Platform = '{gc_plat}' AND CAST(g.SummaryDate AS DATE) BETWEEN '{start}' AND '{end}'
    GROUP BY gc.Id, gc.Name
    """
    c = E.get_connection(); d = E.query_df(c, sql); c.close()
    return d

@st.cache_data(ttl=1800, show_spinner="Loading period data…")
def _load_period_el(start, end):
    sql = f"""
    SELECT TRY_CAST(b.GameId AS INT) AS game_id, MAX(gc.Name) AS game_name,
        SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
        SUM(CAST(b.TotalBet AS FLOAT)/100.0) - SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS net_rev,
        COUNT(DISTINCT b.CasinoName) AS stores, SUM(b.Spins) AS spins
    FROM BetSpinSummaryCashView3 b
    LEFT JOIN GameCatalogView1 gc ON gc.Id = TRY_CAST(b.GameId AS INT)
    WHERE b.PlatformName = 'EdgeLabs' AND CAST(b."Date" AS DATE) BETWEEN '{start}' AND '{end}'
        AND TRY_CAST(b.GameId AS INT) IS NOT NULL
    GROUP BY TRY_CAST(b.GameId AS INT)
    """
    c = E.get_connection(); d = E.query_df(c, sql); c.close()
    return d

@st.cache_data(ttl=1800, show_spinner="Loading revenue since each release…")
def _load_release_window_perf(platform_code: str, game_dates: tuple) -> pd.DataFrame:
    """Bet/net revenue per game_id, summed ONLY over days on/after that specific
    row's own release date — never lifetime. A GameId that already existed for
    years before being re-enabled under a new CRM tag would otherwise drag its
    whole unrelated history into what's displayed as a brand-new release's
    numbers (see reference_whats_new_sql memory — confirmed on "Arctic Buffalo"
    and "Bank It")."""
    if not game_dates:
        return pd.DataFrame()
    values_sql = ",".join(f"({int(gid)}, CAST('{d}' AS DATE))" for gid, d in game_dates)
    if platform_code == "PFH":
        sql = f"""
        SELECT v.gid AS game_id,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) - SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS net_rev
        FROM (VALUES {values_sql}) v(gid, release_date)
        JOIN BetSpinSummaryCashView3Pong b
            ON TRY_CAST(b.GameId AS INT) = v.gid AND CAST(b."Date" AS DATE) >= v.release_date
        GROUP BY v.gid
        """
    elif platform_code in ("V1", "V2"):
        gc_plat = L._GC_PLATFORM.get(platform_code, platform_code)
        sql = f"""
        SELECT v.gid AS game_id,
            SUM(CAST(g.TotalPlay AS FLOAT)/100.0) AS bet,
            SUM(CAST(g.TotalPlay AS FLOAT)/100.0) - SUM(CAST(g.TotalWin AS FLOAT)/100.0) AS net_rev
        FROM (VALUES {values_sql}) v(gid, release_date)
        JOIN GameCatalogView1 gc ON gc.Id = v.gid AND gc.Platform = '{gc_plat}'
        JOIN AnalyticsGameTerminalsGames g
            ON g.Id = v.gid AND CAST(g.SummaryDate AS DATE) >= v.release_date
        GROUP BY v.gid
        """
    else:  # EdgeLabs
        sql = f"""
        SELECT v.gid AS game_id,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) - SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS net_rev
        FROM (VALUES {values_sql}) v(gid, release_date)
        JOIN BetSpinSummaryCashView3 b
            ON TRY_CAST(b.GameId AS INT) = v.gid AND b.PlatformName = 'EdgeLabs'
               AND CAST(b."Date" AS DATE) >= v.release_date
        GROUP BY v.gid
        """
    c = E.get_connection(); d = E.query_df(c, sql); c.close()
    return d

def _period_delta(cur, pri):
    if pri and pri != 0:
        return (cur - pri) / abs(pri) * 100
    return None


def _similarity_pct(distance, n_match_weeks=None, k=65.0):
    """DTW shape distance -> a 0-100 "similarity %" for display.

    dtw_distance() sums per-week absolute differences over the whole matched
    window, so raw distance scales with n_match_weeks (16 weeks compared racks
    up ~2x the distance of 8 weeks, for the same per-week closeness) — the old
    `100 / (1 + distance)` formula ignored that entirely and, worse, decayed so
    fast that even the algorithm's own top-ranked (closest) matches read as
    "0.4% similar", which just looks broken. This normalizes to a per-week
    distance first, then applies exponential decay tuned so that a solid
    peer-group match (the kind of candidate that actually gets returned as a
    top-5 peer) reads in the ~50-70% range instead of near-zero.
    """
    if distance is None or not np.isfinite(distance):
        return np.nan
    per_week = distance / max(float(n_match_weeks), 1.0) if n_match_weeks else distance
    return round(min(100.0 * np.exp(-per_week / k), 99.9), 1)

# ─── Weekly Games page — loaders ───────────────────────────────────────────
# Metric definitions here deliberately mirror the "Game Performance" Power BI
# semantic model (workspace "Game Performance Reports"), read directly from its
# DAX so this page and that report agree:
#   Bet ($)/Day = SUM(Total Bet $) / DISTINCTCOUNT(Date)
#   RTP%        = SUM(Total Win $) / SUM(Total Bet $)
#   Active Games= DISTINCTCOUNT(Game) where Date >= EDATE(TODAY(), -3)
# Panel splits come from real catalog fields: GameCatalogView1.ScreenOrientation
# (horizontal/vertical/responsive) for V2, and .Codebase (gen0/gen1/gen2) for PFH.
@st.cache_data(ttl=1800, show_spinner="Loading weekly game performance…")
def _load_weekly_games(start, end):
    """One row per (platform, product/codebase, orientation, game) for a date window.

    V2 = land-based terminals (AnalyticsGameTerminalsGames, has Product + orientation).
    PFH = the Pong/PFH online side of BetSpinSummaryCashView3, split by Codebase gen.
    """
    conn = E.get_connection()
    sql_v2 = f"""
    SELECT 'V2' AS platform, gc.Product AS segment, gc.ScreenOrientation AS orientation,
        gc.Id AS game_id, gc.Name AS game_name,
        SUM(CAST(g.TotalPlay AS FLOAT)/100.0) AS bet,
        SUM(CAST(g.TotalWin AS FLOAT)/100.0) AS win,
        COUNT(DISTINCT g.SummaryLocationId) AS stores,
        COUNT(DISTINCT CAST(g.SummaryDate AS DATE)) AS game_days
    FROM AnalyticsGameTerminalsGames g
    JOIN GameCatalogView1 gc ON gc.Id = g.Id AND gc.Platform = 'v2'
    WHERE CAST(g.SummaryDate AS DATE) BETWEEN '{start}' AND '{end}'
    GROUP BY gc.Product, gc.ScreenOrientation, gc.Id, gc.Name
    """
    sql_pfh = f"""
    SELECT 'PFH' AS platform, gc.Codebase AS segment, gc.ScreenOrientation AS orientation,
        gc.Id AS game_id, gc.Name AS game_name,
        SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
        SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS win,
        COUNT(DISTINCT b.StoreNumber) AS stores,
        COUNT(DISTINCT CAST(b."Date" AS DATE)) AS game_days
    FROM BetSpinSummaryCashView3 b
    JOIN GameCatalogView1 gc ON gc.Id = TRY_CAST(b.GameId AS INT)
    WHERE b.PlatformName = 'Pong' AND b.CasinoName = 'PFH'
      AND CAST(b."Date" AS DATE) BETWEEN '{start}' AND '{end}'
    GROUP BY gc.Codebase, gc.ScreenOrientation, gc.Id, gc.Name
    """
    try:
        d_v2 = E.query_df(conn, sql_v2)
    except Exception:
        d_v2 = pd.DataFrame()
    try:
        d_pfh = E.query_df(conn, sql_pfh)
    except Exception:
        d_pfh = pd.DataFrame()
    conn.close()
    d = pd.concat([x for x in (d_v2, d_pfh) if not x.empty], ignore_index=True) if (not d_v2.empty or not d_pfh.empty) else pd.DataFrame()
    if d.empty:
        return d
    for c in ("bet", "win", "stores", "game_days"):
        d[c] = pd.to_numeric(d[c], errors="coerce").fillna(0)
    d["segment"] = d["segment"].fillna("unknown").astype(str).str.lower()
    d["orientation"] = d["orientation"].fillna("unknown").astype(str).str.lower()
    return d


@st.cache_data(ttl=21600, show_spinner="Loading newly launched games…")
def _load_new_launches(as_of: str, days_back: int = 30):
    """Games whose first-ever recorded activity falls within days_back of as_of.

    Uses first-appearance in the fact data — the same basis as the Power BI model's
    "First Appearance" column — rather than CRM enable dates, so it stays consistent
    with the bet figures on this page. Store count is measured over the same window.

    ttl=21600 (6h), not the usual 1800 (30min): the PFH half of this query does an
    unfiltered MIN(Date) GROUP BY over BetSpinSummaryCashView3's entire history
    (~21s alone, profiled 2026-08-21) since there's no way to know a game's true
    first-ever date without scanning all of its rows. A "newly launched in the last
    30 days" list doesn't need fresher than same-business-day data, so this trades
    a few hours of staleness for paying that scan far less often.
    """
    conn = E.get_connection()
    sql_v2 = f"""
    SELECT 'V2' AS platform, gc.Name AS game_name, gc.ScreenOrientation AS orientation,
        fa.first_date AS launched, fa.stores
    FROM (
        SELECT g.Id AS gid, MIN(CAST(g.SummaryDate AS DATE)) AS first_date,
               COUNT(DISTINCT g.SummaryLocationId) AS stores
        FROM AnalyticsGameTerminalsGames g
        GROUP BY g.Id
    ) fa
    JOIN GameCatalogView1 gc ON gc.Id = fa.gid AND gc.Platform = 'v2'
    WHERE fa.first_date >= DATEADD('day', -{int(days_back)}, CAST('{as_of}' AS DATE))
      AND fa.first_date <= CAST('{as_of}' AS DATE)
    """
    sql_pfh = f"""
    SELECT 'PFH' AS platform, gc.Name AS game_name, gc.ScreenOrientation AS orientation,
        fa.first_date AS launched, fa.stores
    FROM (
        SELECT TRY_CAST(b.GameId AS INT) AS gid, MIN(CAST(b."Date" AS DATE)) AS first_date,
               COUNT(DISTINCT b.StoreNumber) AS stores
        FROM BetSpinSummaryCashView3 b
        WHERE b.PlatformName = 'Pong' AND b.CasinoName = 'PFH'
          AND TRY_CAST(b.GameId AS INT) IS NOT NULL
        GROUP BY TRY_CAST(b.GameId AS INT)
    ) fa
    JOIN GameCatalogView1 gc ON gc.Id = fa.gid
    WHERE fa.first_date >= DATEADD('day', -{int(days_back)}, CAST('{as_of}' AS DATE))
      AND fa.first_date <= CAST('{as_of}' AS DATE)
    """
    out = []
    for sql in (sql_v2, sql_pfh):
        try:
            out.append(E.query_df(conn, sql))
        except Exception:
            pass
    conn.close()
    out = [x for x in out if not x.empty]
    if not out:
        return pd.DataFrame()
    d = pd.concat(out, ignore_index=True)
    d["launched"] = pd.to_datetime(d["launched"], errors="coerce")
    d["orientation"] = d["orientation"].fillna("–")
    return d.sort_values("launched", ascending=False).reset_index(drop=True)


@st.cache_data(ttl=1800, show_spinner="Counting active games…")
def _load_active_games(as_of: str, months_back: int = 3):
    """Distinct games with any activity in the trailing months_back — matches the
    Power BI model's Active Games measure (Date >= EDATE(TODAY(), -3))."""
    conn = E.get_connection()
    n = 0
    sql_v2 = f"""
    SELECT COUNT(DISTINCT g.Id) AS n FROM AnalyticsGameTerminalsGames g
    WHERE CAST(g.SummaryDate AS DATE) >= DATEADD('month', -{int(months_back)}, CAST('{as_of}' AS DATE))
      AND CAST(g.SummaryDate AS DATE) <= CAST('{as_of}' AS DATE)
    """
    sql_pfh = f"""
    SELECT COUNT(DISTINCT TRY_CAST(b.GameId AS INT)) AS n FROM BetSpinSummaryCashView3 b
    WHERE b.PlatformName = 'Pong' AND b.CasinoName = 'PFH'
      AND CAST(b."Date" AS DATE) >= DATEADD('month', -{int(months_back)}, CAST('{as_of}' AS DATE))
      AND CAST(b."Date" AS DATE) <= CAST('{as_of}' AS DATE)
    """
    for sql in (sql_v2, sql_pfh):
        try:
            r = E.query_df(conn, sql)
            if not r.empty and pd.notna(r["n"].iloc[0]):
                n += int(r["n"].iloc[0])
        except Exception:
            pass
    conn.close()
    return n

# ─── Social Casino page — loaders ──────────────────────────────────────────
# Player-identity metrics (Stickiness/DAU/MAU/ARPDAU) only mean something where
# a returning player can actually be identified — the online/player-account side
# of the business, not anonymous land-based terminals. Two real sources, verified
# live against the "Social Casino Detailed Report" Power BI model (workspace
# "Social Casino Reports", table "Calculations", 91 measures):
#   PFH      -> BetSpinSummaryCashView3 WHERE PlatformName='Pong' — PFH's online
#               identity (see reference_platform_taxonomy: PFH online = "Pong").
#               ALL currencies, not just Sweeps: cross-checked against that
#               report's own visuals, whose "PFH 57K" players and per-game counts
#               (Diamond 7s 20.27K, Fiery 7s 18K, Ruby X 11.96K) match the
#               all-currency figures exactly and NOT the SC-only slice (16,465).
#   EdgeLabs -> BetSpinSummarySocialView2 (its own 20+ casino-brand roster)
# Formulas mirror that model's DAX exactly:
#   Stickiness % = DAU / MAU   (MAU = distinct accounts in that calendar month)
#   ARPDAU = Net Revenue / DAU
# Every query aggregates in SQL (COUNT DISTINCT / SUM) — never pulls row-level
# spin data into pandas, since EdgeLabs alone has millions of rows/year here.
def _social_src(platform: str):
    """(from_sql, base_where, date_expr, acct_expr, bet_expr, win_expr, game_expr,
    has_casino_dim) for the platform's spin-level source.

    PFH/EdgeLabs read BetSpinSummaryCashView3(Pong)/BetSpinSummarySocialView2 — a
    real player-account column (AccountNumber), with a Casino/Aggregator/Currency
    dimension. V1/V2 (land-based) read AnalyticsGameTerminalsGames joined to
    GameCatalogView1 for the platform filter — a different schema (PlayerAccountNumber,
    TotalPlay/TotalWin in cents like the rest, SummaryDate, no casino-brand or
    currency dimension since it's one property's own gaming floor, not an aggregator
    network), but the loyalty-card tap on each row is exactly the same kind of
    player-identity signal AccountNumber gives PFH/EdgeLabs.
    """
    if platform == "EdgeLabs":
        return ("BetSpinSummarySocialView2", "1=1", '"Date"', "AccountNumber",
                "TotalBet", "TotalWin", "GameId", True)
    if platform == "PFH":
        return ("BetSpinSummaryCashView3", "PlatformName = 'Pong'", '"Date"', "AccountNumber",
                "TotalBet", "TotalWin", "GameId", True)
    if platform in ("V1", "V2"):
        gc_plat = L._GC_PLATFORM.get(platform, platform)
        return (f"AnalyticsGameTerminalsGames g JOIN GameCatalogView1 gc ON gc.Id = g.Id",
                f"gc.Platform = '{gc_plat}'", "g.SummaryDate", "g.PlayerAccountNumber",
                "g.TotalPlay", "g.TotalWin", "g.Id", False)
    raise ValueError(f"Unknown platform: {platform}")


def _social_where(platform, casino, aggregator, game_id, start, end, currency=None):
    table, base, date_expr, _acct, _bet, _win, game_expr, _has_casino = _social_src(platform)
    clauses = [base, f"CAST({date_expr} AS DATE) BETWEEN '{start}' AND '{end}'"]
    # Currency must be pinned for EdgeLabs: that table mixes GC (Gold Coins,
    # play-money), SC (Sweeps Coins) and WOW in one column, and summing money
    # across them is meaningless — GC alone yields a 23,825% "RTP" because its
    # win/bet relationship isn't real economics. Verified against live SQL.
    # V1/V2 have no currency/casino/aggregator dimension at all, so these three
    # never apply there — the sidebar UI only ever sets them for EdgeLabs.
    if currency and currency != "All":
        clauses.append(f"CurrencyName = '{currency.replace(chr(39), chr(39)*2)}'")
    if casino and casino != "All":
        clauses.append(f"CasinoName = '{casino.replace(chr(39), chr(39)*2)}'")
    if aggregator and aggregator != "All":
        clauses.append(f"AggregatorName = '{aggregator.replace(chr(39), chr(39)*2)}'")
    if game_id:
        clauses.append(f"TRY_CAST({game_expr} AS INT) = {int(game_id)}")
    return table, " AND ".join(clauses)


@st.cache_data(ttl=1800, show_spinner="Loading social casino filter options…")
def _load_social_filters(platform):
    table, base, _d, _a, _b, _w, _g, has_casino = _social_src(platform)
    if not has_casino:
        return [], []
    conn = E.get_connection()
    try:
        casinos = E.query_df(conn, f"SELECT DISTINCT CasinoName FROM {table} WHERE {base} AND CasinoName IS NOT NULL ORDER BY CasinoName")
    except Exception:
        casinos = pd.DataFrame()
    try:
        aggs = E.query_df(conn, f"SELECT DISTINCT AggregatorName FROM {table} WHERE {base} AND AggregatorName IS NOT NULL AND AggregatorName <> '' ORDER BY AggregatorName")
    except Exception:
        aggs = pd.DataFrame()
    conn.close()
    return (casinos["CasinoName"].dropna().tolist() if not casinos.empty else [],
            aggs["AggregatorName"].dropna().tolist() if not aggs.empty else [])


@st.cache_data(ttl=1800, show_spinner="Loading player activity by date…")
def _load_social_daily(platform, casino, aggregator, start, end, currency=None, game_id=None):
    table, where = _social_where(platform, casino, aggregator, game_id, start, end, currency)
    _t, _base, date_expr, acct_expr, bet_expr, win_expr, _g, _hc = _social_src(platform)
    sql = f"""
    SELECT CAST({date_expr} AS DATE) AS d, COUNT(DISTINCT {acct_expr}) AS players,
        SUM(CAST({bet_expr} AS FLOAT)/100.0) AS bet, SUM(CAST({win_expr} AS FLOAT)/100.0) AS win
    FROM {table} WHERE {where}
    GROUP BY CAST({date_expr} AS DATE)
    """
    conn = E.get_connection()
    try:
        d = E.query_df(conn, sql)
    finally:
        conn.close()
    if d.empty:
        return d
    d["d"] = pd.to_datetime(d["d"])
    for c in ("players", "bet", "win"):
        d[c] = pd.to_numeric(d[c], errors="coerce").fillna(0)
    return d.sort_values("d").reset_index(drop=True)


@st.cache_data(ttl=1800, show_spinner="Loading monthly active players…")
def _load_social_mau(platform, casino, aggregator, start, end, currency=None, game_id=None):
    table, where = _social_where(platform, casino, aggregator, game_id, start, end, currency)
    _t, _base, date_expr, acct_expr, _b, _w, _g, _hc = _social_src(platform)
    sql = f"""
    SELECT YEAR(CAST({date_expr} AS DATE)) AS yr, MONTH(CAST({date_expr} AS DATE)) AS mo,
        COUNT(DISTINCT {acct_expr}) AS mau
    FROM {table} WHERE {where}
    GROUP BY YEAR(CAST({date_expr} AS DATE)), MONTH(CAST({date_expr} AS DATE))
    """
    conn = E.get_connection()
    try:
        d = E.query_df(conn, sql)
    finally:
        conn.close()
    return d


@st.cache_data(ttl=1800, show_spinner="Loading players by game…")
def _load_social_by_game(platform, casino, aggregator, start, end, currency=None):
    table, where = _social_where(platform, casino, aggregator, None, start, end, currency)
    _t, _base, _d, acct_expr, _b, _w, game_expr, _hc = _social_src(platform)
    sql = f"""
    SELECT TRY_CAST({game_expr} AS INT) AS game_id, COUNT(DISTINCT {acct_expr}) AS players
    FROM {table} WHERE {where} AND TRY_CAST({game_expr} AS INT) IS NOT NULL
    GROUP BY TRY_CAST({game_expr} AS INT)
    """
    conn = E.get_connection()
    try:
        d = E.query_df(conn, sql)
        names = E.query_df(conn, "SELECT Id, Name FROM GameCatalogView1")
    finally:
        conn.close()
    if d.empty:
        return d
    d = d.merge(names, left_on="game_id", right_on="Id", how="left")
    d["Name"] = d["Name"].fillna("Game " + d["game_id"].astype(str))
    return d.sort_values("players", ascending=False).reset_index(drop=True)


@st.cache_data(ttl=1800, show_spinner="Loading players by casino…")
def _load_social_by_casino(platform, aggregator, start, end, currency=None, game_id=None):
    _t0, _base0, _d0, _a0, _b0, _w0, _g0, has_casino = _social_src(platform)
    if not has_casino:
        return pd.DataFrame(columns=["CasinoName", "players"])
    table, where = _social_where(platform, None, aggregator, game_id, start, end, currency)
    _t, _base, _d, acct_expr, _b, _w, _g, _hc = _social_src(platform)
    sql = f"""
    SELECT CasinoName, COUNT(DISTINCT {acct_expr}) AS players
    FROM {table} WHERE {where}
    GROUP BY CasinoName
    """
    conn = E.get_connection()
    try:
        d = E.query_df(conn, sql)
    finally:
        conn.close()
    return d.sort_values("players", ascending=False).reset_index(drop=True) if not d.empty else d


@st.cache_data(ttl=1800, show_spinner=False)
def _load_social_game_currencies(platform, game_id):
    """Which currencies a specific game actually has rows in — used to pick a sane
    default Currency (e.g. some EdgeLabs games only ever ran in GC, never SC).
    V1/V2 have no currency dimension at all, so this is always empty there."""
    table, base, _d, _a, _b, _w, game_expr, has_casino = _social_src(platform)
    if not game_id or not has_casino:
        return []
    sql = (f"SELECT DISTINCT CurrencyName FROM {table} WHERE {base} "
           f"AND TRY_CAST({game_expr} AS INT) = {int(game_id)} AND CurrencyName IS NOT NULL")
    conn = E.get_connection()
    try:
        d = E.query_df(conn, sql)
    finally:
        conn.close()
    return d["CurrencyName"].dropna().tolist() if not d.empty else []


@st.cache_data(ttl=1800, show_spinner="Loading player totals…")
def _load_social_totals(platform, casino, aggregator, start, end, currency=None, game_id=None):
    table, where = _social_where(platform, casino, aggregator, game_id, start, end, currency)
    _t, _base, _d, acct_expr, bet_expr, win_expr, _g, _hc = _social_src(platform)
    sql = f"SELECT COUNT(DISTINCT {acct_expr}) AS players, SUM(CAST({bet_expr} AS FLOAT)/100.0) AS bet, SUM(CAST({win_expr} AS FLOAT)/100.0) AS win FROM {table} WHERE {where}"
    conn = E.get_connection()
    try:
        d = E.query_df(conn, sql)
    finally:
        conn.close()
    if d.empty or pd.isna(d["players"].iloc[0]):
        return {"players": 0, "bet": 0.0, "win": 0.0}
    r = d.iloc[0]
    return {"players": int(r["players"] or 0), "bet": float(r["bet"] or 0), "win": float(r["win"] or 0)}


@st.cache_data(ttl=1800, show_spinner="Loading player segments…")
def _load_player_bets(platform, casino, aggregator, start, end, currency=None, game_id=None):
    """Per-account total wagering for the window — the basis for tiering players
    into loyalty segments (New/Casual, Regular, Premium, VIP)."""
    table, where = _social_where(platform, casino, aggregator, game_id, start, end, currency)
    _t, _base, _d, acct_expr, bet_expr, _w, _g, _hc = _social_src(platform)
    sql = f"SELECT {acct_expr} AS player, SUM(CAST({bet_expr} AS FLOAT)/100.0) AS bet FROM {table} WHERE {where} GROUP BY {acct_expr}"
    conn = E.get_connection()
    try:
        d = E.query_df(conn, sql)
    finally:
        conn.close()
    if d.empty:
        return pd.DataFrame(columns=["player", "bet"])
    d["bet"] = pd.to_numeric(d["bet"], errors="coerce").fillna(0)
    return d


# Loyalty tiers by total wagering in the window, split at fixed percentiles
# (bottom 50% / next 30% / next 15% / top 5%) rather than fixed dollar amounts —
# a date range can be a week or two years, so a dollar cutoff either empties out
# or swallows every tier depending on the window. Percentiles always give a
# meaningful spread, and mirror how casino loyalty programs actually tier
# players (by relative value within the current population, not an absolute
# lifetime number).
_PLAYER_TIER_NAMES = ["New / Casual", "Regular", "Premium", "VIP / High Roller"]
_PLAYER_TIER_QUANTILES = [0.50, 0.80, 0.95]


def _segment_players(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame(columns=["segment", "players", "total_bet", "share_pct"]), []
    bets = df["bet"].to_numpy()
    thresholds = np.quantile(bets, _PLAYER_TIER_QUANTILES)
    tier_idx = np.searchsorted(thresholds, bets, side="right")
    rows = []
    for i, name in enumerate(_PLAYER_TIER_NAMES):
        mask = tier_idx == i
        rows.append({"segment": name, "players": int(mask.sum()), "total_bet": float(bets[mask].sum())})
    seg = pd.DataFrame(rows)
    tot = seg["total_bet"].sum()
    seg["share_pct"] = (seg["total_bet"] / tot * 100) if tot > 0 else 0.0
    return seg, thresholds.tolist()


@st.cache_data(ttl=1800, show_spinner="Loading revenue trend…")
def _load_calendar_daily(platform, start, end):
    """Whole-fleet bet/net revenue grouped by real calendar date (not launch week).

    Used only for the All Games Health revenue trend chart, which buckets these
    daily rows into day/week/month depending on the selected local period.
    """
    conn = E.get_connection()
    if platform == "PFH":
        sql = f"""
        SELECT CAST(b."Date" AS DATE) AS d,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) - SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS net_rev
        FROM BetSpinSummaryCashView3Pong b
        WHERE CAST(b."Date" AS DATE) BETWEEN '{start}' AND '{end}'
        GROUP BY CAST(b."Date" AS DATE)
        """
    elif platform in ("V1", "V2"):
        gc_plat = L._GC_PLATFORM.get(platform, platform)
        sql = f"""
        SELECT CAST(g.SummaryDate AS DATE) AS d,
            SUM(CAST(g.TotalPlay AS FLOAT)/100.0) AS bet,
            SUM(CAST(g.TotalPlay AS FLOAT)/100.0) - SUM(CAST(g.TotalWin AS FLOAT)/100.0) AS net_rev
        FROM AnalyticsGameTerminalsGames g
        JOIN GameCatalogView1 gc ON gc.Id = g.Id
        WHERE gc.Platform = '{gc_plat}' AND CAST(g.SummaryDate AS DATE) BETWEEN '{start}' AND '{end}'
        GROUP BY CAST(g.SummaryDate AS DATE)
        """
    elif platform == "EdgeLabs":
        sql = f"""
        SELECT CAST(b."Date" AS DATE) AS d,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) - SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS net_rev
        FROM BetSpinSummaryCashView3 b
        WHERE b.PlatformName = 'EdgeLabs' AND CAST(b."Date" AS DATE) BETWEEN '{start}' AND '{end}'
        GROUP BY CAST(b."Date" AS DATE)
        """
    else:
        conn.close()
        return pd.DataFrame(columns=["d", "bet", "net_rev"])
    d = E.query_df(conn, sql)
    conn.close()
    if d.empty:
        return pd.DataFrame(columns=["d", "bet", "net_rev"])
    for c in ("bet", "net_rev"):
        d[c] = pd.to_numeric(d[c], errors="coerce").fillna(0)
    d["d"] = pd.to_datetime(d["d"])
    return d.sort_values("d").reset_index(drop=True)

def _combined_calendar_daily(start, end):
    """Same shape as _load_calendar_daily, summed across all 4 platforms."""
    parts = []
    for p in _AGS_PLATFORMS:
        try:
            d = _load_calendar_daily(p, start, end)
            if not d.empty:
                parts.append(d)
        except Exception:
            pass
    if not parts:
        return pd.DataFrame(columns=["d", "bet", "net_rev"])
    allc = pd.concat(parts, ignore_index=True)
    return allc.groupby("d", as_index=False).agg(bet=("bet", "sum"), net_rev=("net_rev", "sum")).sort_values("d")

def _combined_period_pg(start, end):
    """Same shape as _load_period_pfh/_load_period_v2v1/_load_period_el, all 4 platforms
    stacked with a platform-qualified game_id (GameCatalogView1 ids are NOT unique across
    platforms — see the What's New fix — so a raw numeric id would silently collide here too)."""
    parts = []
    for p, loader in (("PFH", lambda: _load_period_pfh(start, end)),
                       ("V1", lambda: _load_period_v2v1("V1", start, end)),
                       ("V2", lambda: _load_period_v2v1("V2", start, end)),
                       ("EdgeLabs", lambda: _load_period_el(start, end))):
        try:
            d = loader()
        except Exception:
            d = pd.DataFrame()
        if not d.empty:
            d = d.copy()
            d["game_id"] = p + "_" + d["game_id"].astype(str)
            parts.append(d)
    return pd.concat(parts, ignore_index=True) if parts else pd.DataFrame(columns=["game_id", "game_name", "bet", "net_rev"])

# Approximate geographic centroids, USPS code -> (lat, lon). Static facts, not derived from
# any query — used only to place the state-abbreviation text labels on the choropleth map.
_STATE_CENTROIDS = {
    "AL": (32.8, -86.8), "AK": (64.2, -149.4), "AZ": (34.3, -111.1), "AR": (34.8, -92.2),
    "CA": (36.8, -119.4), "CO": (39.0, -105.5), "CT": (41.6, -72.7), "DE": (39.0, -75.5),
    "FL": (27.8, -81.7), "GA": (32.2, -83.4), "HI": (20.8, -156.3), "ID": (44.1, -114.5),
    "IL": (40.0, -89.2), "IN": (39.8, -86.1), "IA": (42.0, -93.5), "KS": (38.5, -98.3),
    "KY": (37.5, -85.3), "LA": (31.2, -91.8), "ME": (45.4, -69.0), "MD": (39.1, -76.8),
    "MA": (42.2, -71.5), "MI": (44.3, -85.4), "MN": (46.4, -93.1), "MS": (32.7, -89.7),
    "MO": (38.5, -92.5), "MT": (46.9, -110.4), "NE": (41.5, -99.9), "NV": (38.5, -117.1),
    "NH": (43.5, -71.6), "NJ": (40.1, -74.5), "NM": (34.5, -106.1), "NY": (42.9, -75.6),
    "NC": (35.5, -79.4), "ND": (47.5, -100.4), "OH": (40.4, -82.8), "OK": (35.6, -96.9),
    "OR": (44.1, -120.5), "PA": (40.6, -77.2), "RI": (41.7, -71.5), "SC": (33.8, -81.2),
    "SD": (44.4, -100.2), "TN": (35.9, -86.7), "TX": (31.1, -97.6), "UT": (39.3, -111.1),
    "VT": (44.0, -72.7), "VA": (37.8, -78.2), "WA": (47.4, -120.5), "WV": (38.6, -80.5),
    "WI": (44.3, -89.6), "WY": (43.0, -107.5), "DC": (38.9, -77.0),
}

# Centroids for location codes that show up in CrmLocationView.StateProv but aren't one of
# the 50 US states + DC — the choropleth (locationmode="USA-states") silently drops anything
# else with no error, so these get a dot marker instead (see the map code below).
# Confirmed against the live database on 2026-08-17 (not guessed):
#   - Canadian provinces (ON, QC, ...) — standard 2-letter codes.
#   - PR = Puerto Rico (BusinessName "Puerto Rico #1/2/3", City "San Juan") — a real US
#     territory, but Plotly's USA-states geometry doesn't include it either.
#   - PS = Trinidad (BusinessName/City show "Port of Spain", "Couva", "Curepe", "Freeport") —
#     not a real ISO code, just this org's internal shorthand.
#   - PH and SB both resolve to Sint Maarten locations (Philipsburg, Simpson Bay, Cole Bay) —
#     the SAME island tagged with two different internal codes. That's a real data-quality
#     inconsistency in CrmLocationView, not something to silently merge here — both codes are
#     plotted as-is (at the same real coordinates) so nothing is hidden, but a fleet-wide
#     "revenue by region" total will double-count/fragment Sint Maarten until the source data
#     is cleaned up.
_NON_US_CENTROIDS = {
    "ON": (50.0, -85.0), "QC": (52.0, -71.8), "BC": (53.7, -127.6), "AB": (54.0, -114.0),
    "MB": (55.0, -97.0), "SK": (54.0, -106.0), "NS": (45.0, -63.0), "NB": (46.5, -66.0),
    "PE": (46.4, -63.2), "NL": (53.1, -60.0), "YT": (64.0, -135.0), "NT": (64.8, -124.8),
    "NU": (70.0, -85.0),
    "PR": (18.47, -66.12),  # Puerto Rico — US territory, not a state
    "PS": (10.66, -61.52),  # Trinidad (internal code, not ISO)
    "PH": (18.02, -63.05),  # Sint Maarten — Philipsburg
    "SB": (18.04, -63.10),  # Sint Maarten — Simpson Bay / Cole Bay (same island as PH)
}

@st.cache_data(ttl=1800, show_spinner="Loading footprint…")
def load_game_footprint(platform, game_id):
    """Lifetime per-location bet/net for ONE game — footprint/concentration only
    ("is this a real title or one hot store"), not a period breakdown. EdgeLabs has
    no location dimension (casino-based) so this returns empty for it."""
    if platform == "EdgeLabs":
        return pd.DataFrame()
    conn = E.get_connection()
    if platform == "PFH":
        sql = f"""
        SELECT CAST(b.StoreNumber AS VARCHAR) AS loc_id,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) - SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS net_rev,
            MIN(CAST(b."Date" AS DATE)) AS first_seen, MAX(CAST(b."Date" AS DATE)) AS last_seen
        FROM BetSpinSummaryCashView3Pong b
        WHERE TRY_CAST(b.GameId AS INT) = {int(game_id)}
        GROUP BY CAST(b.StoreNumber AS VARCHAR)
        """
    else:
        sql = f"""
        SELECT CAST(g.SummaryLocationId AS VARCHAR) AS loc_id,
            SUM(CAST(g.TotalPlay AS FLOAT)/100.0) AS bet,
            SUM(CAST(g.TotalPlay AS FLOAT)/100.0) - SUM(CAST(g.TotalWin AS FLOAT)/100.0) AS net_rev,
            MIN(CAST(g.SummaryDate AS DATE)) AS first_seen, MAX(CAST(g.SummaryDate AS DATE)) AS last_seen
        FROM AnalyticsGameTerminalsGames g
        WHERE g.Id = {int(game_id)}
        GROUP BY CAST(g.SummaryLocationId AS VARCHAR)
        """
    d = E.query_df(conn, sql)
    conn.close()
    if d.empty:
        return d
    for c in ("bet", "net_rev"):
        d[c] = pd.to_numeric(d[c], errors="coerce").fillna(0)
    return d


@st.cache_data(ttl=1800, show_spinner="Loading lifetime store counts…")
def load_lifetime_stores(platform) -> pd.DataFrame:
    """True lifetime distinct-location count per game, one GROUP BY GameId pass over
    the whole unfiltered table, no date bound -- the correct denominator for any
    lifetime-scoped per-store metric. Mirrors load_game_footprint's per-game query
    but for every game on the platform at once. EdgeLabs has no location dimension
    (casino-based), so this returns empty for it.

    Added because Full Roster's "#Stores" was really the *last week's* store count
    sitting next to otherwise-lifetime Total Bet/Net Revenue columns -- correct
    individually, but reads as a lifetime figure it isn't. This gives the actually-
    lifetime number to show instead.
    """
    if platform == "EdgeLabs":
        return pd.DataFrame(columns=["game_id", "stores_lifetime"])
    conn = E.get_connection()
    if platform == "PFH":
        sql = """
        SELECT TRY_CAST(b.GameId AS INT) AS game_id, COUNT(DISTINCT b.StoreNumber) AS stores_lifetime
        FROM BetSpinSummaryCashView3Pong b
        WHERE TRY_CAST(b.GameId AS INT) IS NOT NULL
        GROUP BY TRY_CAST(b.GameId AS INT)
        """
    else:
        sql = """
        SELECT g.Id AS game_id, COUNT(DISTINCT g.SummaryLocationId) AS stores_lifetime
        FROM AnalyticsGameTerminalsGames g
        GROUP BY g.Id
        """
    try:
        d = E.query_df(conn, sql)
    finally:
        conn.close()
    if not d.empty:
        d["game_id"] = pd.to_numeric(d["game_id"], errors="coerce")
        d["stores_lifetime"] = pd.to_numeric(d["stores_lifetime"], errors="coerce")
    return d

# Real geography breakdown — joins each platform's real bet/win data to CrmLocationView
# (SQL Server, not Oracle). EdgeLabs is casino-based (CasinoName, not a location dimension),
# so it's not included here.
@st.cache_data(ttl=1800, show_spinner="Loading geography…")
def load_geo_detail(platform, start, end):
    """Per location x game rows with state / location / account-manager / game attached.

    One query per (platform, period); the four page filters are then applied in pandas so
    changing a filter is instant instead of re-querying. ~14k rows for PFH — small enough
    to hold in memory. Money columns are cents in source, divided by 100 here.
    """
    conn = E.get_connection()
    if platform == "PFH":
        sql = f"""
        SELECT loc.StateProv AS state, loc.Latitude AS lat, loc.Longitude AS lon,
            ISNULL(loc.BusinessName, 'Unknown') AS location_name,
            ISNULL(loc.AccountManager, 'Unassigned') AS account_manager,
            ISNULL(loc.ConfigProduct, 'Unknown') AS product,
            TRY_CAST(b.GameId AS INT) AS game_id,
            MAX(gc.Name) AS game_name,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) AS bet,
            SUM(CAST(b.TotalBet AS FLOAT)/100.0) - SUM(CAST(b.TotalWin AS FLOAT)/100.0) AS net_rev
        FROM BetSpinSummaryCashView3Pong b
        LEFT JOIN CrmLocationView loc ON CAST(b.StoreNumber AS VARCHAR) = CAST(loc.LocationId AS VARCHAR)
        LEFT JOIN GameCatalogView1 gc ON gc.Id = TRY_CAST(b.GameId AS INT)
        WHERE TRY_CAST(b.GameId AS INT) IS NOT NULL
          AND CAST(b."Date" AS DATE) BETWEEN '{start}' AND '{end}'
        GROUP BY loc.StateProv, loc.Latitude, loc.Longitude, loc.BusinessName, loc.AccountManager, loc.ConfigProduct, TRY_CAST(b.GameId AS INT)
        """
    elif platform in ("V1", "V2"):
        gc_plat = L._GC_PLATFORM.get(platform, platform)
        sql = f"""
        SELECT loc.StateProv AS state, loc.Latitude AS lat, loc.Longitude AS lon,
            ISNULL(loc.BusinessName, 'Unknown') AS location_name,
            ISNULL(loc.AccountManager, 'Unassigned') AS account_manager,
            ISNULL(loc.ConfigProduct, 'Unknown') AS product,
            gc.Id AS game_id, gc.Name AS game_name,
            SUM(CAST(g.TotalPlay AS FLOAT)/100.0) AS bet,
            SUM(CAST(g.TotalPlay AS FLOAT)/100.0) - SUM(CAST(g.TotalWin AS FLOAT)/100.0) AS net_rev
        FROM AnalyticsGameTerminalsGames g
        JOIN GameCatalogView1 gc ON gc.Id = g.Id
        LEFT JOIN CrmLocationView loc ON CAST(g.SummaryLocationId AS VARCHAR) = CAST(loc.LocationId AS VARCHAR)
        WHERE gc.Platform = '{gc_plat}'
          AND CAST(g.SummaryDate AS DATE) BETWEEN '{start}' AND '{end}'
        GROUP BY loc.StateProv, loc.Latitude, loc.Longitude, loc.BusinessName, loc.AccountManager, loc.ConfigProduct, gc.Id, gc.Name
        """
    else:
        conn.close()
        return pd.DataFrame()
    d = E.query_df(conn, sql)
    conn.close()
    if d.empty:
        return d
    for c in ("bet", "net_rev", "lat", "lon"):
        d[c] = pd.to_numeric(d[c], errors="coerce")
    d[["bet", "net_rev"]] = d[["bet", "net_rev"]].fillna(0)
    d["state"] = d["state"].fillna("Unknown").astype(str).str.strip().str.upper()
    for c in ("location_name", "account_manager", "game_name", "product"):
        d[c] = d[c].fillna("Unknown").astype(str).str.strip()
    return d


@st.cache_data(ttl=1800, show_spinner="Loading top players…")
def load_top_players(platform, start, end, limit=8):
    """Top loyalty accounts by wagering for the period. VIP Kiosk Network reads
    AccountNumber off its spin table directly; the gaming floors (V1/V2) use the
    PlayerAccountNumber loyalty-card tap recorded per terminal-day row."""
    conn = E.get_connection()
    try:
        if platform == "PFH":
            sql = f"""
            SELECT AccountNumber AS player, SUM(CAST(TotalBet AS FLOAT)/100.0) AS bet
            FROM BetSpinSummaryCashView3Pong
            WHERE CAST("Date" AS DATE) BETWEEN '{start}' AND '{end}'
            GROUP BY AccountNumber ORDER BY bet DESC LIMIT {limit}
            """
        elif platform in ("V1", "V2"):
            gc_plat = L._GC_PLATFORM.get(platform, platform)
            sql = f"""
            SELECT g.PlayerAccountNumber AS player, SUM(CAST(g.TotalPlay AS FLOAT)/100.0) AS bet
            FROM AnalyticsGameTerminalsGames g JOIN GameCatalogView1 gc ON gc.Id = g.Id
            WHERE gc.Platform = '{gc_plat}' AND CAST(g.SummaryDate AS DATE) BETWEEN '{start}' AND '{end}'
            GROUP BY g.PlayerAccountNumber ORDER BY bet DESC LIMIT {limit}
            """
        else:
            return pd.DataFrame(columns=["player", "bet"])
        d = E.query_df(conn, sql)
        if not d.empty:
            d["bet"] = pd.to_numeric(d["bet"], errors="coerce").fillna(0)
        return d
    except Exception:
        return pd.DataFrame(columns=["player", "bet"])
    finally:
        conn.close()


with tab_overview:
    ribbon("Location Overview", f"{floor_label} · revenue by region, property, account manager and game", T2, "OVERVIEW")

    if platform == "EdgeLabs":
        st.info("The Private Members' Club has no property/region dimension, so this geographic page "
                "doesn't apply. Use the other tabs for its analysis.")
    else:
        # ── Period + the four page-scoped filters ───────────────────────
        # Defaults to YTD: a partial current week makes for a near-empty map.
        period = st.radio("Period", _HS_PERIODS, index=0, horizontal=True,
                          label_visibility="collapsed", key="ov_period")
        cur_range, _pri_range = _hs_get_dates(period)
        s, e = cur_range

        geo_raw = load_geo_detail(platform, str(s), str(e))

        if geo_raw.empty:
            st.info("No data for this period.")
        else:
            # No in-page Game filter — the sidebar's Game selector (used by every other tab)
            # scopes this page too, so there's one game control for the whole app, not two.
            if not is_all:
                geo_raw = geo_raw[geo_raw["game_id"].astype("Int64") == int(sel_id)]

            f1, f2, f3 = st.columns(3)
            state_opts = sorted(geo_raw["state"].unique())
            sel_states = f1.multiselect("Region", state_opts, default=[],
                                        placeholder="All regions", key="ov_f_state")

            # Location / manager options cascade off the states already chosen, so the
            # dropdowns only ever offer combinations that actually exist.
            _scoped = geo_raw[geo_raw["state"].isin(sel_states)] if sel_states else geo_raw
            sel_locs = f2.multiselect("Location", sorted(_scoped["location_name"].unique()),
                                      default=[], placeholder="All locations", key="ov_f_loc")
            _scoped2 = _scoped[_scoped["location_name"].isin(sel_locs)] if sel_locs else _scoped
            sel_mgrs = f3.multiselect("Account Manager", sorted(_scoped2["account_manager"].unique()),
                                      default=[], placeholder="All managers", key="ov_f_mgr")

            geo_f = geo_raw.copy()
            if sel_states: geo_f = geo_f[geo_f["state"].isin(sel_states)]
            if sel_locs:   geo_f = geo_f[geo_f["location_name"].isin(sel_locs)]
            if sel_mgrs:   geo_f = geo_f[geo_f["account_manager"].isin(sel_mgrs)]

            if geo_f.empty:
                st.warning("No data matches the selected filters.")
            else:
                tot_bet = float(geo_f["bet"].sum())
                tot_net = float(geo_f["net_rev"].sum())
                hold = (tot_net / tot_bet) * 100 if tot_bet > 0 else 0
                n_states = geo_f["state"].nunique()
                n_locs = geo_f["location_name"].nunique()

                _active = [x for x in [
                    f"game: {meta['game_name']}" if not is_all else None,
                    f"{len(sel_states)} states" if sel_states else None,
                    f"{len(sel_locs)} locations" if sel_locs else None,
                    f"{len(sel_mgrs)} managers" if sel_mgrs else None] if x]
                st.caption(f"{s:%b %d, %Y} – {e:%b %d, %Y}"
                           + (f" · filtered to {', '.join(_active)}" if _active else " · no filters applied"))

                krow([
                    {"label": "Net Revenue", "value": _usd(tot_net)},
                    {"label": "Bet Handle", "value": _usd(tot_bet)},
                    {"label": "Hold %", "value": _pct(hold, 2)},
                    {"label": "Regions", "value": str(n_states)},
                    {"label": "Properties", "value": f"{n_locs:,}"},
                ])

                by_state = (geo_f.groupby("state", as_index=False)
                                 .agg(bet=("bet", "sum"), net_rev=("net_rev", "sum"),
                                      locations=("location_name", "nunique"))
                                 .sort_values("net_rev", ascending=False))
                by_state["hold_pct"] = (((by_state["net_rev"] / by_state["bet"].clip(lower=1)) * 100).clip(upper=99.9)).round(1)
                # One row per actual property, with its real lat/lon — this used to collapse
                # to one dot per province/state (so all of Québec's 15 properties showed as a
                # single point); now every property gets its own dot at its real city.
                by_loc = (geo_f.groupby(["location_name", "state"], as_index=False)
                               .agg(bet=("bet", "sum"), net_rev=("net_rev", "sum"),
                                    lat=("lat", "first"), lon=("lon", "first"))
                               .sort_values("net_rev", ascending=False))
                by_loc["hold_pct"] = (((by_loc["net_rev"] / by_loc["bet"].clip(lower=1)) * 100).clip(upper=99.9)).round(1)

                mcol, rcol = st.columns([7, 5])

                with mcol:
                    # One dot per actual property at its real coordinates (from
                    # CrmLocationView.Latitude/Longitude) — real OpenStreetMap tiles (free, no
                    # API key). Bubble size = Net Revenue. Properties with no coordinate on file
                    # (shouldn't happen for this data, but defensive) are dropped from the map,
                    # not silently placed at 0,0 — they're still in the bar charts below.
                    plotted = by_loc.dropna(subset=["lat", "lon"])
                    plotted_lats = plotted["lat"].tolist()
                    plotted_lons = plotted["lon"].tolist()

                    # Bubble area (not radius) should scale with value, or small properties read
                    # as falsely near-zero next to one 50x their size — sqrt keeps that
                    # proportional. net_rev can be negative, so scale on magnitude, not the raw
                    # clipped-at-zero value, or every losing property collapses to the same
                    # minimum size regardless of how large the loss is.
                    _max_abs_net = max(float(plotted["net_rev"].abs().max()), 1.0) if not plotted.empty else 1.0
                    _bubble_sizes = (np.sqrt(plotted["net_rev"].abs() / _max_abs_net) * 34 + 7).tolist()
                    _halo_sizes = [s + 4 for s in _bubble_sizes]

                    # go.Scattermap (MapLibre-based), not the older go.Scattermapbox -- same free
                    # open-street-map style with no token needed, but not deprecated.
                    # A tan/beige colorscale (fine on the old flat abstract map) all but vanishes
                    # against real OSM tiles' own green/tan terrain -- confirmed live, most bubbles
                    # were unreadable. Switched to a bold red<->green diverging scale (loss vs
                    # profit, cmid=0) plus a dark halo trace underneath every bubble for contrast,
                    # since Scattermap markers have no `line` (outline) sub-property of their own.
                    fig_map = go.Figure()
                    fig_map.add_trace(go.Scattermap(
                        lat=plotted_lats, lon=plotted_lons, mode="markers",
                        marker=dict(size=_halo_sizes, sizemode="diameter", color="rgba(30,26,20,0.55)"),
                        hoverinfo="skip", showlegend=False))
                    fig_map.add_trace(go.Scattermap(
                        lat=plotted_lats, lon=plotted_lons, mode="markers",
                        marker=dict(size=_bubble_sizes, sizemode="diameter",
                                    color=plotted["net_rev"], colorscale=[[0, MOVE_DOWN], [0.5, "#E8DCC8"], [1, MOVE_UP]],
                                    cmid=0, colorbar=dict(title="Net Rev", tickprefix="$", thickness=11, len=0.75),
                                    opacity=0.95),
                        customdata=plotted[["location_name", "state", "net_rev", "bet", "hold_pct"]].values,
                        hovertemplate="<b>%{customdata[0]}</b> (%{customdata[1]})<br>Net Revenue: $%{customdata[2]:,.0f}<br>"
                                     "Total Bet: $%{customdata[3]:,.0f}<br>"
                                     "Hold: %{customdata[4]:.1f}%<extra></extra>",
                        showlegend=False))

                    # No fitbounds equivalent for map subplots (unlike geo's fitbounds="locations") --
                    # center on whatever's actually plotted so the view stays correct as filters change,
                    # rather than a fixed North-America center that could leave a filtered subset off-screen.
                    _center_lat = sum(plotted_lats) / len(plotted_lats) if plotted_lats else 45.0
                    _center_lon = sum(plotted_lons) / len(plotted_lons) if plotted_lons else -90.0
                    fig_map.update_layout(
                        map=dict(style="open-street-map", center=dict(lat=_center_lat, lon=_center_lon), zoom=2.9),
                        height=645, margin=dict(l=0, r=0, t=4, b=0), paper_bgcolor="rgba(0,0,0,0)")
                    st.markdown('<div class="card"><div class="card-title">Net Revenue by Property · bubbles sized by Net Revenue</div>', unsafe_allow_html=True)
                    st.plotly_chart(fig_map, use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)

                with rcol:
                    t5s = by_state.head(5)
                    fig_t5s = go.Figure(go.Bar(x=t5s["net_rev"], y=t5s["state"], orientation="h",
                        marker_color=T2, text=[_usd(v) for v in t5s["net_rev"]], textposition="outside",
                        hovertemplate="<b>%{y}</b><br>Net: $%{x:,.0f}<extra></extra>"))
                    fig_t5s.update_layout(yaxis=dict(autorange="reversed"))
                    st.markdown('<div class="card"><div class="card-title">Top 5 Regions — Net Revenue</div>', unsafe_allow_html=True)
                    st.plotly_chart(plotly_base(fig_t5s, h=300, ml=50, mr=80, mt=4, mb=30), use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)

                    t5l = by_loc.head(5)
                    fig_t5l = go.Figure(go.Bar(x=t5l["net_rev"], y=t5l["location_name"].str[:26], orientation="h",
                        marker_color=MOVE_UP, text=[_usd(v) for v in t5l["net_rev"]], textposition="outside",
                        customdata=t5l[["location_name", "state"]].values,
                        hovertemplate="<b>%{customdata[0]}</b> (%{customdata[1]})<br>Net: $%{x:,.0f}<extra></extra>"))
                    fig_t5l.update_layout(yaxis=dict(autorange="reversed"))
                    st.markdown('<div class="card"><div class="card-title">Top 5 Properties — Net Revenue</div>', unsafe_allow_html=True)
                    st.plotly_chart(plotly_base(fig_t5l, h=300, ml=170, mr=80, mt=4, mb=30), use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)

                # Second row — All States (left) beside Revenue by Location Type (right),
                # so the wide blank strip below the map/top-5 row gets used too.
                bcol, lcol = st.columns([7, 5])
                with bcol:
                    fig_all = go.Figure(go.Bar(x=by_state["net_rev"], y=by_state["state"], orientation="h",
                        marker_color=T2, text=[_usd(v) for v in by_state["net_rev"]], textposition="outside",
                        hovertemplate="<b>%{y}</b><br>Net: $%{x:,.0f}<extra></extra>"))
                    fig_all.update_layout(yaxis=dict(autorange="reversed"))
                    fig_all.update_xaxes(title_text="Net revenue ($)")
                    st.markdown(f'<div class="card"><div class="card-title">All Regions ({len(by_state)})</div>', unsafe_allow_html=True)
                    st.plotly_chart(plotly_base(fig_all, h=max(415, len(by_state) * 30), ml=85, mr=70, mt=4, mb=55), use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)

                with lcol:
                    # Built from geo_f — the SAME period + state/location/manager/game filtered
                    # data as everything else on this page, not a separate always-lifetime query.
                    # So picking a period here changes this chart exactly like it changes the map.
                    _geo_generic = geo_f.assign(product=_generic_product(geo_f["product"]))
                    loc_type_df = (_geo_generic.groupby("product", as_index=False)
                                        .agg(bet=("bet", "sum"), net_rev=("net_rev", "sum"))
                                        .sort_values("net_rev", ascending=False))
                    if loc_type_df.empty:
                        st.info("No game-category breakdown available.")
                    else:
                        fig_lt = go.Figure(go.Bar(x=loc_type_df["net_rev"], y=loc_type_df["product"], orientation="h",
                            marker_color="#7A6A9C", text=[_usd(v) for v in loc_type_df["net_rev"]], textposition="outside",
                            hovertemplate="<b>%{y}</b><br>Net: $%{x:,.0f}<extra></extra>"))
                        fig_lt.update_layout(yaxis=dict(autorange="reversed"))
                        fig_lt.update_xaxes(title_text="Net revenue ($)")
                        period_note = "lifetime" if period == "Lifetime" else period
                        st.markdown(f'<div class="card"><div class="card-title">Revenue by Game Category ({period_note})</div>', unsafe_allow_html=True)
                        st.plotly_chart(plotly_base(fig_lt, h=max(415, len(loc_type_df) * 30), ml=100, mr=70, mt=4, mb=55), use_container_width=True, theme=None)
                        st.markdown('</div>', unsafe_allow_html=True)

                # Third row — Top Players, using this floor's loyalty-account data.
                top_players = load_top_players(platform, str(s), str(e))
                if not top_players.empty:
                    max_bet = float(top_players["bet"].max()) or 1.0
                    rows_html = "".join(
                        f'<div class="lrow"><span class="lrank">{i+1}</span>'
                        f'<span class="lname">{r.player}</span>'
                        f'<div class="lbar-w"><div class="lbar" style="width:{r.bet/max_bet*100:.0f}%;background:{GOLD}"></div></div>'
                        f'<span class="lval">{_usd(r.bet)}</span></div>'
                        for i, r in enumerate(top_players.itertuples())
                    )
                    st.markdown(
                        f'<div style="background:{SURFACE};border:1px solid {BORDER};border-radius:10px;padding:14px 16px;margin-top:14px">'
                        f'<div class="eyebrow">Top Players — {floor_label} loyalty accounts, by wagering</div>'
                        f'<div style="margin-top:8px">{rows_html}</div></div>',
                        unsafe_allow_html=True)

@st.cache_data(ttl=1800, show_spinner=False)
def load_active_players(platform, start, end):
    """Distinct loyalty-account count for the period — only floors with loyalty-card
    tracking (VIP Kiosk Network, Private Members' Club, and the gaming floors' own
    PlayerAccountNumber) can identify a returning player."""
    conn = E.get_connection()
    try:
        if platform == "PFH":
            sql = f"""SELECT COUNT(DISTINCT AccountNumber) AS n FROM BetSpinSummaryCashView3Pong
                      WHERE CAST("Date" AS DATE) BETWEEN '{start}' AND '{end}'"""
        elif platform == "EdgeLabs":
            sql = f"""SELECT COUNT(DISTINCT AccountNumber) AS n FROM BetSpinSummaryCashView3EdgeLabs
                      WHERE CAST("Date" AS DATE) BETWEEN '{start}' AND '{end}'"""
        else:
            gc_plat = L._GC_PLATFORM.get(platform, platform)
            sql = f"""SELECT COUNT(DISTINCT g.PlayerAccountNumber) AS n
                      FROM AnalyticsGameTerminalsGames g JOIN GameCatalogView1 gc ON gc.Id = g.Id
                      WHERE gc.Platform = '{gc_plat}' AND CAST(g.SummaryDate AS DATE) BETWEEN '{start}' AND '{end}'"""
        d = E.query_df(conn, sql)
        return int(d["n"].iloc[0]) if not d.empty and pd.notna(d["n"].iloc[0]) else 0
    except Exception:
        return 0
    finally:
        conn.close()


def _active_players_for(scope_label, platform_code, start, end):
    if scope_label == "All Floors":
        return sum(load_active_players(p, start, end) for p in ["PFH", "V2", "V1", "EdgeLabs"])
    return load_active_players(platform_code, start, end)


# ══════════════════════════════════════════════════════════════════
# TAB — ALL GAMES
# ══════════════════════════════════════════════════════════════════
with tab1:
    # Always the whole property, every system combined — one casino, one portfolio
    # view. No toggle: this used to offer a single-system option, but Le Grand
    # Palais doesn't have a "which system" question for the user to answer.
    ags_scope = "All Floors"
    h_label = "Le Grand Palais"
    h_df = full_df
    h_cat = full_cat
    h_nonhr = full_nonhr

    ribbon("All Games Health", "Full property portfolio, every game — always fleet-wide, independent of the sidebar's game selector", T2, "FLEET")

    # ── Row 1: revenue trend (adaptive granularity) + Top 5 / 5 Latest ──
    period2 = st.radio("Trend period", ["MTD", "YTD", "Last Week", "Lifetime"], index=0, horizontal=True,
                       label_visibility="collapsed", key="agh_period")
    cur2, pri2 = _hs_get_dates(period2)
    s2, e2 = cur2
    ps2, pe2 = pri2

    if ags_scope == "All Floors":
        cal = _combined_calendar_daily(str(s2), str(e2))
    else:
        cal = _load_calendar_daily(platform, str(s2), str(e2))
    cur_net = float(cal["net_rev"].sum()) if not cal.empty else 0.0
    cur_bet = float(cal["bet"].sum()) if not cal.empty else 0.0
    if period2 == "Lifetime":
        trend_delta = None  # comparing lifetime-to-date against itself is meaningless — no prior period to show
    else:
        if ags_scope == "All Floors":
            cal_pri = _combined_calendar_daily(str(ps2), str(pe2))
        else:
            cal_pri = _load_calendar_daily(platform, str(ps2), str(pe2))
        pri_net = float(cal_pri["net_rev"].sum()) if not cal_pri.empty else 0.0
        trend_delta = _period_delta(cur_net, pri_net)

    if ags_scope == "All Floors":
        per_game = _combined_period_pg(str(s2), str(e2))
    elif platform == "PFH":
        per_game = _load_period_pfh(str(s2), str(e2))
    elif platform in ("V1", "V2"):
        per_game = _load_period_v2v1(platform, str(s2), str(e2))
    else:
        per_game = _load_period_el(str(s2), str(e2))
    if per_game.empty:
        per_game = pd.DataFrame(columns=["game_id", "net_rev"])
    per_game["net_rev"] = pd.to_numeric(per_game["net_rev"], errors="coerce").fillna(0)
    pg = h_nonhr[["game_id", "game_name"]].drop_duplicates("game_id").merge(
        per_game[["game_id", "net_rev"]], on="game_id", how="left")
    pg["net_rev"] = pd.to_numeric(pg["net_rev"], errors="coerce").fillna(0)

    # ── Lifetime roster (needed here for the Bet Level mini-chart, and again below for the full table) ──
    # stores_lifetime = true distinct locations EVER (one GROUP BY GameId pass, no date
    # bound) -- NOT the last launch-week's store count. Full Roster used to show that
    # last-week snapshot next to otherwise-lifetime Total Bet/Net Revenue columns, which
    # read as "lifetime stores" but wasn't (confirmed live: "Bank It 2" showed 56 there
    # vs. its real 415 lifetime locations). net_store_day now divides by this too, for
    # the same reason -- a lifetime average shouldn't be denominated by a one-week count.
    if ags_scope == "All Floors":
        _lts_parts = []
        for _p in _AGS_PLATFORMS:
            _lts = load_lifetime_stores(_p)
            if not _lts.empty:
                _lts = _lts.copy()
                _lts["game_id"] = _p + "_" + _lts["game_id"].astype("Int64").astype(str)
                _lts_parts.append(_lts)
        stores_lifetime = (pd.concat(_lts_parts, ignore_index=True) if _lts_parts
                            else pd.DataFrame(columns=["game_id", "stores_lifetime"]))
    else:
        stores_lifetime = load_lifetime_stores(platform)

    # "Since Jan 2025" columns -- an ADDITIONAL window shown alongside lifetime, not a
    # replacement. Requested to compare against a specific external re-derivation that
    # used this exact start date; that write-up explicitly called this window's "lifetime
    # label false" and only used it to reproduce a known-wrong reference point, so the
    # verified-correct lifetime columns stay as the primary figures -- this just adds the
    # same comparison directly into the table instead of it living in a one-off SQL run.
    _SINCE_START = "2025-01-01"
    _since_end = str(dt.date.today())

    def _load_since_2025(_p):
        if _p == "PFH":
            _d = _load_period_pfh(_SINCE_START, _since_end)
        elif _p in ("V1", "V2"):
            _d = _load_period_v2v1(_p, _SINCE_START, _since_end)
        else:
            _d = _load_period_el(_SINCE_START, _since_end)
        if _d.empty:
            return pd.DataFrame(columns=["game_id", "bet_since2025", "net_since2025"])
        return _d[["game_id", "bet", "net_rev"]].rename(columns={"bet": "bet_since2025", "net_rev": "net_since2025"})

    if ags_scope == "All Floors":
        _since_parts = []
        for _p in _AGS_PLATFORMS:
            _sd = _load_since_2025(_p)
            if not _sd.empty:
                _sd = _sd.copy()
                _sd["game_id"] = _p + "_" + _sd["game_id"].astype("Int64").astype(str)
                _since_parts.append(_sd)
        since2025 = (pd.concat(_since_parts, ignore_index=True) if _since_parts
                     else pd.DataFrame(columns=["game_id", "bet_since2025", "net_since2025"]))
    else:
        since2025 = _load_since_2025(platform)

    spins_tot = h_df.groupby("game_id")["spins"].sum().reset_index().rename(columns={"spins": "total_spins"})
    life = (h_cat.merge(stores_lifetime, on="game_id", how="left")
                 .merge(spins_tot, on="game_id", how="left")
                 .merge(since2025, on="game_id", how="left"))
    fl_bet = max(float(life["total_bet"].sum()), 1.0)
    fl_net = float(life["total_net"].sum())
    life["bet_share"] = ((life["total_bet"] / fl_bet * 100).clip(upper=99.9)).round(1)
    life["hold_pct"] = (((life["total_net"] / life["total_bet"].clip(lower=1)) * 100).clip(upper=99.9)).round(1)
    life["days_live"] = (life["total_weeks"] + 1) * 7
    life["net_store_day"] = (life["total_net"] / life["stores_lifetime"].clip(lower=1) / life["days_live"]).round(2)
    life["hold_since2025"] = np.clip(np.where(life["bet_since2025"] > 0,
                                      (life["net_since2025"] / life["bet_since2025"]) * 100, np.nan), None, 99.9).round(1) \
        if "bet_since2025" in life.columns else np.nan
    life["tenure"] = np.select(
        [life["total_weeks"] < 13, life["total_weeks"] < 52],
        ["🌿 New Release", "🌳 Floor Regular"], default="🌲 House Classic")
    # Full Roster's default view/ranking is "since Jan 2025" (bet_since2025) -- lifetime
    # stays fully computed (Flag/guardrail below is still keyed on lifetime net_rev, never
    # this) so the historical-damage signal doesn't disappear, it's just not the primary
    # sort here anymore. Other panels above (Top 5 Games, etc.) explicitly re-sort by
    # total_bet themselves and are unaffected by this.
    life = life.sort_values("bet_since2025", ascending=False, na_position="last").reset_index(drop=True)

    # Modeled floor operating cost, as a share of Net Revenue -- not derived from a
    # real cost feed (none exists for this synthetic clone). Illustrative only.
    _opex_pct = 0.30
    _profit_dollars = cur_net * (1 - _opex_pct)
    _active_players = _active_players_for(h_label, platform, str(s2), str(e2))

    rcol1, rcol2 = st.columns([7, 5])
    with rcol1:
        st.caption(f"{s2:%b %d, %Y} – {e2:%b %d, %Y}")
        krow([
            {"label": "Net Revenue", "value": _usd(cur_net), "delta_pct": trend_delta},
            {"label": "Bet Handle", "value": _usd(cur_bet)},
            {"label": "Hold %", "value": _pct((cur_net / cur_bet) * 100, 2) if cur_bet else "–"},
            {"label": "Profit Margin", "value": _pct((1 - _opex_pct) * 100, 0), "sub": f"{_usd(_profit_dollars)} est. profit"},
            {"label": "Active Player Accounts", "value": f"{_active_players:,}"},
        ])
        if cal.empty:
            st.info("No revenue data for this period.")
        else:
            if period2 == "Last Week":
                cal["bucket"] = cal["d"].dt.strftime("%a %m/%d")
                gran = "day by day"
            elif period2 == "MTD":
                wk_num = ((cal["d"] - pd.Timestamp(s2)).dt.days // 7) + 1
                cal["bucket"] = "Week " + wk_num.astype(str)
                gran = "week by week"
            else:
                cal["bucket"] = cal["d"].dt.strftime("%b %Y")
                gran = "month by month"
            order = list(dict.fromkeys(cal.sort_values("d")["bucket"]))
            trend = cal.groupby("bucket", as_index=False).agg(net_rev=("net_rev", "sum"), bet=("bet", "sum"))
            trend["bucket"] = pd.Categorical(trend["bucket"], categories=order, ordered=True)
            trend = trend.sort_values("bucket")
            fig_trend = go.Figure(go.Bar(x=trend["bucket"], y=trend["net_rev"], marker_color=T2,
                hovertemplate="<b>%{x}</b><br>Net Rev: $%{y:,.0f}<extra></extra>"))
            fig_trend.update_yaxes(title_text="Net revenue ($)")
            st.markdown(f'<div class="card"><div class="card-title">All Games Revenue — {gran}</div>', unsafe_allow_html=True)
            st.plotly_chart(plotly_base(fig_trend, h=470), use_container_width=True, theme=None)
            st.markdown('</div>', unsafe_allow_html=True)

    with rcol2:
        top5 = pg.sort_values("net_rev", ascending=False).head(5)
        fig_top5 = go.Figure(go.Bar(x=top5["net_rev"], y=top5["game_name"].str[:22], orientation="h",
            marker_color=T2, text=[_usd(v) for v in top5["net_rev"]], textposition="outside",
            hovertemplate="<b>%{y}</b><br>Net: $%{x:,.0f}<extra></extra>"))
        fig_top5.update_layout(yaxis=dict(autorange="reversed"))
        st.markdown('<div class="card"><div class="card-title">Top 5 Games — Net Revenue</div>', unsafe_allow_html=True)
        st.plotly_chart(plotly_base(fig_top5, h=200, ml=130, mr=70, mt=6, mb=28), use_container_width=True, theme=None)
        st.markdown('</div>', unsafe_allow_html=True)

        latest5 = h_nonhr.sort_values("launch_date", ascending=False).head(5)[["game_id", "game_name"]].merge(
            pg[["game_id", "net_rev"]], on="game_id", how="left")
        latest5["net_rev"] = latest5["net_rev"].fillna(0)
        fig_l5 = go.Figure(go.Bar(x=latest5["net_rev"], y=latest5["game_name"].str[:22], orientation="h",
            marker_color=MOVE_UP, text=[_usd(v) for v in latest5["net_rev"]], textposition="outside",
            hovertemplate="<b>%{y}</b><br>Net: $%{x:,.0f}<extra></extra>"))
        fig_l5.update_layout(yaxis=dict(autorange="reversed"))
        st.markdown('<div class="card"><div class="card-title">5 Latest Games — Net Revenue</div>', unsafe_allow_html=True)
        st.plotly_chart(plotly_base(fig_l5, h=200, ml=130, mr=70, mt=6, mb=28), use_container_width=True, theme=None)
        st.markdown('</div>', unsafe_allow_html=True)

        top5bl = life.sort_values("total_bet", ascending=False).head(5)
        fig_bl_mini = go.Figure(go.Bar(x=top5bl["total_bet"], y=top5bl["game_name"].str[:22], orientation="h",
            marker_color=AMBER, text=[_usd(v) for v in top5bl["total_bet"]], textposition="outside",
            hovertemplate="<b>%{y}</b><br>Bet: $%{x:,.0f}<extra></extra>"))
        fig_bl_mini.update_layout(yaxis=dict(autorange="reversed"))
        st.markdown('<div class="card"><div class="card-title">Top 5 Games — Bet Level (Lifetime)</div>', unsafe_allow_html=True)
        st.plotly_chart(plotly_base(fig_bl_mini, h=200, ml=130, mr=70, mt=6, mb=28), use_container_width=True, theme=None)
        st.markdown('</div>', unsafe_allow_html=True)

    # ── Row 2: full games roster (fleet-wide) ──
    _n_flagged = int((life["total_net"] < 0).sum())
    ribbon("Full Roster", f"All {h_label} games ranked by wagering since Jan 2025 — {len(life)} games tracked"
           + (f" · {_n_flagged} flagged (negative lifetime net)" if _n_flagged else ""), T2, "DETAIL")

    # Guardrail, made visible in the table itself, not just in Is It On Track's status:
    # negative LIFETIME net means RTP > 100% lifetime -- real, confirmed money lost, never
    # auto-suppressed regardless of dollar size (see the "Classic Times" -$590 case). This
    # stays keyed on lifetime net_rev even though the table's default view/ranking below is
    # "since Jan 2025" -- the guardrail must not lose sight of historical damage just
    # because the primary display window changed.
    # Note: the Flag below is computed from life["total_net"] (true lifetime), even though
    # none of the lifetime columns are shown in this table anymore -- removing them from
    # display doesn't change what the guardrail checks, just what's visible alongside it.
    life["flag"] = np.where(life["total_net"] < 0, "⚠ NEEDS REVIEW", "")
    # game_id is platform-qualified here (e.g. "V2_9581") -- strip the system prefix for
    # display, since which internal system a game runs on isn't shown to the user anymore.
    life["_game_id_disp"] = life["game_id"].astype(str).str.rsplit("_", n=1).str[-1]
    life_tbl = life[["game_name", "_game_id_disp", "tenure",
                     "bet_since2025", "net_since2025", "hold_since2025",
                     "net_store_day", "bet_share", "flag"]].rename(columns={
        "game_name": "Game", "_game_id_disp": "Game ID", "tenure": "Tenure",
        "bet_since2025": "Bet", "net_since2025": "Net Revenue", "hold_since2025": "Hold %",
        "net_store_day": "Net/Casino/Day", "bet_share": "Bet Share %", "flag": "Flag",
    })

    def _highlight_flagged(row):
        return [f"background-color: {_hex_alpha(MOVE_DOWN, 0.12)}; color: {MOVE_DOWN}; font-weight: 600"] * len(row) \
            if row["Flag"] else [""] * len(row)

    st.caption("Bet/Net Revenue/Hold % are since Jan 1, 2025. The Flag is still based on full lifetime net "
               "revenue (not shown here) — a game can show healthy recent numbers and still be flagged for "
               "historical damage.")
    st.dataframe(
        life_tbl.style.format({
            "Bet": "${:,.0f}", "Net Revenue": "${:,.0f}",
            "Hold %": "{:.1f}%", "Net/Casino/Day": "${:.2f}", "Bet Share %": "{:.1f}%",
        }, na_rep="–").apply(_highlight_flagged, axis=1),
        use_container_width=True, hide_index=True, height=560,
    )

# ══════════════════════════════════════════════════════════════════
# TAB — WEEKLY GAMES
# ══════════════════════════════════════════════════════════════════
# Streamlit port of the "Game Performance Dashboard" Power BI page: a grid of
# Top-N-by-Bet($)/Day panels split by platform → product → screen orientation,
# plus newly-launched games, product share, and the three headline KPIs.
#
# Two deliberate departures from that report, both because the report is wrong:
#  1. Its "Horizontal" and "Vertical" panels show identical numbers — the
#     orientation filter isn't actually applied there. The underlying data does
#     differ (V2 P2P for 13–19 Jul: vertical $5.66M vs horizontal $5.63M), so
#     these panels really do filter on ScreenOrientation.
#  2. Its "Total Handle … M/day" is a weekly total labelled as a daily rate
#     (SUMX over every fact row of a per-day measure). Shown here as an explicit
#     week total with the true daily average beside it.
WG_SEGMENTS = [
    ("V2", "p2p", "horizontal", "Classic — Horizontal", "#1F6F52"),
    ("V2", "p2p", "vertical", "Classic — Vertical", "#1F6F52"),
    ("V2", "sweeps", "horizontal", "Sweepstakes — Horizontal", "#7B2B3B"),
    ("V2", "sweeps", "vertical", "Sweepstakes — Vertical", "#7B2B3B"),
    ("V2", "pulltabs", "horizontal", "Instant Win — Horizontal", "#2C5A82"),
    ("V2", "pulltabs", "vertical", "Instant Win — Vertical", "#2C5A82"),
    ("PFH", "gen2", None, "Kiosk Games — Current Gen", "#7A6524"),
    ("PFH", "gen1", None, "Kiosk Games — Classic Gen", "#4E3B2A"),
]


def _wg_panel(title, color, rows, total_rtp, total_betday, note=None, top_n=10):
    """One Top-N panel rendered as an HTML table — colored header, in-cell bars,
    and a sticky total row, matching the source Power BI visual."""
    max_bet = max([r["bet_day"] for r in rows], default=0) or 1
    body = ""
    for i, r in enumerate(rows[:top_n], start=1):
        pct = max(2.0, r["bet_day"] / max_bet * 100)
        rtp_txt = f"{r['rtp']:.2f}%" if r["rtp"] is not None else "–"
        body += (
            f'<tr>'
            f'<td style="text-align:right;padding:6px 8px;color:{T2};width:34px">{i}</td>'
            f'<td style="padding:6px 8px;color:{TEXT};white-space:nowrap;overflow:hidden;'
            f'text-overflow:ellipsis;max-width:220px" title="{r["game_name"]}">{r["game_name"]}</td>'
            f'<td style="text-align:right;padding:6px 8px;color:{T2};white-space:nowrap">{rtp_txt}</td>'
            f'<td style="padding:6px 8px;width:42%">'
            f'<div style="position:relative;background:{S2};border-radius:4px;height:22px">'
            f'<div style="width:{pct:.1f}%;background:{color};height:22px;border-radius:4px"></div>'
            f'<span style="position:absolute;right:6px;top:0;font-size:14px;line-height:22px;'
            f'font-weight:600;color:{TEXT}">{_usd(r["bet_day"])}</span></div></td>'
            f'</tr>'
        )
    if not rows:
        body = f'<tr><td colspan="4" style="padding:14px;color:{T3};font-size:15px">No data in this window.</td></tr>'
    tot_rtp_txt = f"{total_rtp:.2f}%" if total_rtp is not None else "–"
    return (
        f'<div class="wg-zoomable" style="border:1px solid {BORDER};border-radius:10px;overflow:hidden;background:{SURFACE};margin-bottom:14px">'
        f'<div style="background:{color};color:#fff;font-size:17px;font-weight:700;padding:9px 10px;text-align:center">{title}</div>'
        f'<table style="width:100%;border-collapse:collapse;font-size:15px">'
        f'<thead><tr style="border-bottom:1px solid {BORDER};background:{BG}">'
        f'<th style="text-align:right;padding:6px 8px;font-size:13px;color:{T2}">#</th>'
        f'<th style="text-align:left;padding:6px 8px;font-size:13px;color:{T2}">Game</th>'
        f'<th style="text-align:right;padding:6px 8px;font-size:13px;color:{T2}">Hold%</th>'
        f'<th style="text-align:left;padding:6px 8px;font-size:13px;color:{T2}">Bet ($)/Day</th>'
        f'</tr></thead><tbody>{body}</tbody>'
        f'<tfoot><tr style="border-top:2px solid {BORDER};background:{BG};font-weight:700">'
        f'<td></td><td style="padding:7px 8px;font-size:13px;color:{T2}">'
        f'{("Top " + str(min(top_n, len(rows))) + " of " + str(len(rows))) if rows else ""}</td>'
        f'<td style="text-align:right;padding:7px 8px;font-size:15px;color:{TEXT}">{tot_rtp_txt}</td>'
        f'<td style="text-align:right;padding:7px 8px;font-size:15px;color:{TEXT}">{_usd(total_betday)}</td>'
        f'</tr></tfoot></table>'
        + (f'<div style="padding:5px 10px;font-size:13px;color:{T3}">{note}</div>' if note else "")
        + '</div>'
    )


with tab_weekly:
    ribbon("Weekly Games", "Top games by Bet ($)/Day across the property — one week at a glance, "
                           "independent of the sidebar's game selector", T2, "WEEKLY")

    # ── Week picker: complete Mon–Sun weeks, most recent first ──
    _wg_today = dt.date.today()
    _wg_last_mon = _wg_today - dt.timedelta(days=_wg_today.weekday() + 7)
    _wg_weeks = [(_wg_last_mon - dt.timedelta(weeks=i)) for i in range(12)]
    _wg_labels = [f"{m.strftime('%d %b')} – {(m + dt.timedelta(days=6)).strftime('%d %b %Y')}" for m in _wg_weeks]
    wcol1, wcol2 = st.columns([2, 3])
    with wcol1:
        _wg_pick = st.selectbox("Week", _wg_labels, index=0, key="wg_week", label_visibility="collapsed")
    _wg_start = _wg_weeks[_wg_labels.index(_wg_pick)]
    _wg_end = _wg_start + dt.timedelta(days=6)
    with wcol2:
        st.markdown(f'<div style="padding-top:6px;font-size:17px;color:{T2}">'
                    f'<b style="color:{TEXT};font-size:19px">{_wg_start.strftime("%d %b")} – {_wg_end.strftime("%d %b %Y")}</b>'
                    f' · generated {_wg_today.strftime("%d %b %Y")}</div>', unsafe_allow_html=True)

    try:
        wg = _load_weekly_games(str(_wg_start), str(_wg_end))
    except Exception as e:
        wg = pd.DataFrame()
        st.error(f"Could not load weekly game data: {e}")

    if wg.empty:
        st.info(f"No V2 or PFH activity recorded for {_wg_start} – {_wg_end}.")
    else:
        # One shared per-day denominator per platform: the number of days in the
        # window that platform actually reported. Using each game's own active-day
        # count (as the Power BI measure does) makes the column stop summing to its
        # own total, which reads as broken; this keeps Bet($)/Day additive.
        _wg_days = {p: max(1, int(g["game_days"].max())) for p, g in wg.groupby("platform")}
        wg["bet_day"] = wg.apply(lambda r: r["bet"] / _wg_days.get(r["platform"], 7), axis=1)
        wg["rtp"] = np.clip(np.where(wg["bet"] > 0, (wg["bet"] - wg["win"]) / wg["bet"] * 100, np.nan), None, 99.9)

        def _wg_slice(plat, seg, orient):
            s = wg[(wg["platform"] == plat)]
            if plat == "PFH":
                s = s[s["segment"] == "gen2"] if seg == "gen2" else s[s["segment"] != "gen2"]
            else:
                s = s[s["segment"] == seg]
                if orient:
                    s = s[s["orientation"] == orient]
            return s.sort_values("bet_day", ascending=False)

        # ── Grid: 3 rows × 4 cols, mirroring the source report's layout ──
        _panels = []
        for plat, seg, orient, title, color in WG_SEGMENTS:
            s = _wg_slice(plat, seg, orient)
            rows = [{"game_name": str(r["game_name"]), "rtp": (float(r["rtp"]) if pd.notna(r["rtp"]) else None),
                     "bet_day": float(r["bet_day"])} for _, r in s.iterrows()]
            tb, tw = float(s["bet"].sum()), float(s["win"].sum())
            _panels.append({
                "title": title, "color": color, "rows": rows,
                "total_rtp": min((tb - tw) / tb * 100, 99.9) if tb > 0 else None,
                "total_betday": tb / _wg_days.get(plat, 7),
            })
        _pmap = {p["title"]: p for p in _panels}

        def _render(title, note=None):
            p = _pmap[title]
            st.markdown(_wg_panel(p["title"], p["color"], p["rows"], p["total_rtp"],
                                   p["total_betday"], note=note), unsafe_allow_html=True)

        r1c1, r1c2, r1c3, r1c4 = st.columns(4)
        with r1c1:
            _render("Classic — Horizontal")
        with r1c2:
            _render("Classic — Vertical")
        with r1c3:
            # Newly launched — last 30 days, first-ever activity date
            try:
                nl = _load_new_launches(str(_wg_end), 30)
            except Exception:
                nl = pd.DataFrame()
            _nl_body = ""
            if nl.empty:
                _nl_body = f'<tr><td colspan="5" style="padding:14px;color:{T3};font-size:15px">No new games in the 30 days to {_wg_end.strftime("%d %b")}.</td></tr>'
            else:
                for _, r in nl.head(12).iterrows():
                    _nl_body += (
                        f'<tr><td style="padding:6px 8px;color:{TEXT};white-space:nowrap;overflow:hidden;'
                        f'text-overflow:ellipsis;max-width:200px" title="{r["game_name"]}">{r["game_name"]}</td>'
                        f'<td style="padding:6px 8px;color:{T2}">{r["orientation"]}</td>'
                        f'<td style="padding:6px 8px;color:{T2}">{FLOOR_LABELS.get(r["platform"], r["platform"])}</td>'
                        f'<td style="padding:6px 8px;color:{T2};white-space:nowrap">{r["launched"].strftime("%d %b %Y") if pd.notna(r["launched"]) else "–"}</td>'
                        f'<td style="text-align:right;padding:6px 8px;color:{TEXT}">{int(r["stores"]):,}</td></tr>')
            st.markdown(
                f'<div class="wg-zoomable" style="border:1px solid {BORDER};border-radius:10px;overflow:hidden;background:{SURFACE};margin-bottom:14px">'
                f'<div style="background:#5B4B8A;color:#fff;font-size:17px;font-weight:700;padding:9px 10px;text-align:center">'
                f'★ Newly Launched Games — Last 30 Days ★</div>'
                f'<table style="width:100%;border-collapse:collapse;font-size:15px">'
                f'<thead><tr style="border-bottom:1px solid {BORDER};background:{BG}">'
                f'<th style="text-align:left;padding:6px 8px;font-size:13px;color:{T2}">Game</th>'
                f'<th style="text-align:left;padding:6px 8px;font-size:13px;color:{T2}">Orientation</th>'
                f'<th style="text-align:left;padding:6px 8px;font-size:13px;color:{T2}">System</th>'
                f'<th style="text-align:left;padding:6px 8px;font-size:13px;color:{T2}">Launched</th>'
                f'<th style="text-align:right;padding:6px 8px;font-size:13px;color:{T2}"># Casinos</th>'
                f'</tr></thead><tbody>{_nl_body}</tbody></table></div>', unsafe_allow_html=True)
        with r1c4:
            _share = []
            for p in _panels:
                if p["total_betday"] > 0:
                    _share.append((p["title"].replace(" — ", " "), p["total_betday"]))
            st.markdown(f'<div class="card"><div class="card-title">Product Share — Bet/Day</div>', unsafe_allow_html=True)
            if _share:
                fig_wg_share = go.Figure(go.Pie(
                    labels=[s[0] for s in _share], values=[s[1] for s in _share], hole=0.42,
                    marker=dict(colors=[PLT_COLORWAY[i % len(PLT_COLORWAY)] for i in range(len(_share))],
                                line=dict(color=SURFACE, width=1.5)),
                    textinfo="percent", textfont=dict(size=13),
                    hovertemplate="<b>%{label}</b><br>$%{value:,.0f}/day (%{percent})<extra></extra>"))
                fig_wg_share.update_layout(height=330, margin=dict(l=4, r=4, t=6, b=6),
                                           paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                                           legend=dict(font=dict(size=12), orientation="v",
                                                       yanchor="middle", y=0.5, x=1.0))
                st.plotly_chart(fig_wg_share, use_container_width=True, theme=None)
            else:
                st.caption("No share data.")
            st.markdown('</div>', unsafe_allow_html=True)

        r2c1, r2c2, r2c3, r2c4 = st.columns(4)
        with r2c1:
            _render("Sweepstakes — Horizontal")
        with r2c2:
            _render("Sweepstakes — Vertical")
        with r2c3:
            _render("Kiosk Games — Current Gen")
        with r2c4:
            _tot_bet_wk = float(wg["bet"].sum())
            _tot_win_wk = float(wg["win"].sum())
            _tot_betday = sum(p["total_betday"] for p in _panels)
            try:
                _n_active = _load_active_games(str(_wg_end), 3)
            except Exception:
                _n_active = None
            _kpi_style = (f"background:{SURFACE};border:1px solid {BORDER};border-radius:10px;"
                          f"padding:14px 16px;margin-bottom:14px;text-align:center")
            st.markdown(
                f'<div style="{_kpi_style}"><div class="eyebrow">Total Handle (week)</div>'
                f'<div style="font-size:34px;font-weight:700;color:{TEXT};font-family:{SERIF}">{_usd(_tot_bet_wk)}</div>'
                f'<div style="font-size:14px;color:{T2}">{_usd(_tot_betday)} / day</div></div>'
                f'<div style="{_kpi_style}"><div class="eyebrow">Overall Hold %</div>'
                f'<div style="font-size:34px;font-weight:700;color:{TEXT};font-family:{SERIF}">'
                f'{((_tot_bet_wk - _tot_win_wk) / _tot_bet_wk * 100) if _tot_bet_wk > 0 else 0:.2f}%</div></div>'
                f'<div style="{_kpi_style}"><div class="eyebrow">Total Active Games</div>'
                f'<div style="font-size:34px;font-weight:700;color:{TEXT};font-family:{SERIF}">'
                f'{_n_active if _n_active is not None else "–"}</div>'
                f'<div style="font-size:14px;color:{T2}">any activity in last 3 months</div></div>',
                unsafe_allow_html=True)

        r3c1, r3c2, r3c3, _r3c4 = st.columns(4)
        with r3c1:
            _render("Instant Win — Horizontal")
        with r3c2:
            _render("Instant Win — Vertical")
        with r3c3:
            _render("Kiosk Games — Classic Gen", note="Classic Gen = every kiosk codebase except the current generation.")

        _oth = wg[(wg["platform"] == "V2") & (~wg["segment"].isin(["p2p", "sweeps", "pulltabs"]))]
        _resp = wg[(wg["platform"] == "V2") & (wg["segment"].isin(["p2p", "sweeps", "pulltabs"]))
                   & (~wg["orientation"].isin(["horizontal", "vertical"]))]
        _skipped = []
        if not _oth.empty:
            _skipped.append(f"{_oth['game_id'].nunique()} V2 games in other products "
                            f"({', '.join(sorted(_oth['segment'].unique()))}) — {_usd(float(_oth['bet'].sum()))} bet")
        if not _resp.empty:
            _skipped.append(f"{_resp['game_id'].nunique()} V2 games with non-horizontal/vertical orientation "
                            f"— {_usd(float(_resp['bet'].sum()))} bet")
        st.caption(
            f"Bet ($)/Day = total bet ÷ {_wg_days.get('V2', 7)} reported days (V2) / "
            f"{_wg_days.get('PFH', 7)} (PFH). Hold % = (total bet − total win) ÷ total bet — the house's "
            "share of what's wagered. Panels filter on GameCatalogView1.ScreenOrientation "
            "(V2) and .Codebase (PFH)."
            + (" Not shown in the panels above: " + "; ".join(_skipped) + "." if _skipped else ""))


# ══════════════════════════════════════════════════════════════════
# TAB — WHAT'S NEW
# ══════════════════════════════════════════════════════════════════
with tab_new:
    ribbon("What's New", "New game releases across the property — independent of the sidebar's filters", T2, "NEW")
    fcol1, fcol2 = st.columns(2)
    with fcol1:
        win_label = st.radio("Window", ["30 days", "60 days", "90 days", "120 days"], index=3,
                              horizontal=True, label_visibility="collapsed", key="wn_days_sel")
        wn_days = int(win_label.split()[0])
    with fcol2:
        view_mode = st.radio("View", ["Game Releases", "All Updates"],
                              horizontal=True, key="wn_view_mode", label_visibility="collapsed")
    plat_filter = "All"  # every system combined -- no per-system filter in the UI anymore

    cutoff = pd.Timestamp.now() - pd.Timedelta(days=wn_days)
    try:
        wn_df = load_whats_new(wn_days)
    except Exception:
        st.warning("Could not load release data. Check database connection.")
        wn_df = pd.DataFrame()

    if view_mode == "Game Releases":
        ribbon("Game Releases", f"Last {wn_days} days — first enable date · launch-day location count", GREEN, "GAME")
        try:
            gr_land = load_game_releases()
            gr_land = gr_land[gr_land["LastDate"] >= cutoff]
        except Exception as e:
            gr_land = pd.DataFrame()
            st.error(f"Land-based load error: {e}")
        try:
            gr_el = load_edgelabs_releases(wn_days)
        except Exception as e:
            gr_el = pd.DataFrame()
            st.error(f"EdgeLabs load error: {e}")

        gr_df = pd.concat([gr_land, gr_el], ignore_index=True) if not (gr_land.empty and gr_el.empty) else pd.DataFrame()

        if not gr_df.empty:
            cls = gr_df["PlatformProduct"].apply(lambda x: pd.Series(_classify_game_platform(x)))
            gr_df["_tc"], gr_df["_vd"] = cls[0], cls[1]

            # De-dupe PFH/V1 double-tagging: PFH's land-based side runs ON V1, so one real
            # release gets logged under BOTH tags in CrmUpdateLogView (same Note/Studio/
            # release date — confirmed live on "ULTIMATE GOAL 10X": 3 rows, all with the
            # identical 144-location count and a null GameId, since the CRM's per-location
            # Platform string was recorded inconsistently across an otherwise single event).
            # Collapse any (Studio, Note, LastDate) group whose rows are ALL tagged PFH
            # and/or V1 into one row — but only when they agree on GameId (or all lack one).
            # A genuine GameId conflict means the rows resolved to different catalog games
            # and must NOT be guessed into one.
            def _dedupe_pfh_v1(df):
                out = []
                for _, grp in df.groupby(["Studio", "Note", "LastDate"], dropna=False):
                    if len(grp) == 1 or not set(grp["_tc"]).issubset({"PFH", "V1"}):
                        out.append(grp)
                        continue
                    gids = grp["GameId"].dropna().unique()
                    if len(gids) > 1:
                        out.append(grp)  # genuine conflict — don't guess, keep separate
                        continue
                    rep = (grp[grp["_tc"] == "PFH"].head(1) if (grp["_tc"] == "PFH").any() else grp.head(1)).copy()
                    rep["Locations"] = grp["Locations"].max()
                    out.append(rep)
                return pd.concat(out, ignore_index=True) if out else df

            gr_df = _dedupe_pfh_v1(gr_df)
            gr_df = gr_df.sort_values("LastDate", ascending=False).reset_index(drop=True)

            # Enrich with revenue/RTP SINCE EACH ROW'S OWN RELEASE DATE — never lifetime.
            # A GameId that already existed for years before being re-enabled under a new
            # CRM tag (e.g. an old P2P title later enabled under "V2 Pull-Tabs" with no
            # matching catalog row) would otherwise drag its whole unrelated history into
            # what's displayed as a brand-new release's numbers — confirmed live on
            # "Arctic Buffalo" (showed 2 days live / $1.75M net, actually a 2.5-year-old
            # game) and "Bank It" (showed 33 days live / -$22.7M net, actually a 4-year-old
            # jackpot game). Scoping to >= release date fixes both, and is a no-op for
            # genuinely new games since they have no data before their own release anyway.
            _perf_by_plat = {}
            for _pc in ("PFH", "V1", "V2", "EdgeLabs"):
                _rows = gr_df[(gr_df["_tc"] == _pc) & gr_df["GameId"].notna()]
                _pairs = tuple(sorted({(int(r["GameId"]), r["LastDate"].strftime("%Y-%m-%d"))
                                        for _, r in _rows.iterrows()}))
                try:
                    _pg = _load_release_window_perf(_pc, _pairs)
                    _perf_by_plat[_pc] = ({int(r["game_id"]): {"net_rev": float(r["net_rev"]), "total_bet": float(r["bet"])}
                                            for _, r in _pg.iterrows()} if not _pg.empty else {})
                except Exception:
                    _perf_by_plat[_pc] = {}

            def _wn_lookup(row):
                m = _perf_by_plat.get(row["_tc"], {})
                g = row["GameId"]
                if pd.notna(g) and int(g) in m:
                    return m[int(g)]
                return None
            _looked_up = gr_df.apply(_wn_lookup, axis=1)
            gr_df["_net_rev"] = _looked_up.apply(lambda v: v["net_rev"] if v else None)
            gr_df["_hold_pct"] = _looked_up.apply(
                lambda v: min((v["net_rev"] / v["total_bet"]) * 100, 99.9) if v and v["total_bet"] > 0 else None)

            # ── Platform comparison data — always all 3, regardless of the filter below ──
            cmp_df = gr_df[gr_df["_tc"].isin(["PFH", "V2", "EdgeLabs"])]
            cnt_by_plat = cmp_df.groupby("_tc").size().reindex(["PFH", "V2", "EdgeLabs"]).fillna(0).reset_index(name="n")

            if plat_filter != "All":
                gr_df = gr_df[gr_df["_tc"] == plat_filter]

            gr_rev = gr_df[gr_df["_net_rev"].notna()]
            krow([{"label": "New Games", "value": str(len(gr_df))}])

            # ── New Releases by Platform, beside Top New Games by Net Revenue ──
            top10 = gr_rev.sort_values("_net_rev", ascending=False).head(10)
            ccol1, ccol2 = st.columns(2)
            with ccol1:
                fig_cnt = go.Figure(go.Bar(x=cnt_by_plat["n"], y=cnt_by_plat["_tc"], orientation="h",
                    marker_color=T2, text=cnt_by_plat["n"].astype(int), textposition="outside",
                    hovertemplate="<b>%{y}</b><br>New releases: %{x:.0f}<extra></extra>"))
                fig_cnt.update_layout(yaxis=dict(autorange="reversed"))
                st.markdown(f'<div class="card"><div class="card-title">New Releases by Platform — last {wn_days} days</div>', unsafe_allow_html=True)
                st.plotly_chart(plotly_base(fig_cnt, h=380, ml=90, mr=50, mt=6, mb=28), use_container_width=True, theme=None)
                st.markdown('</div>', unsafe_allow_html=True)
            with ccol2:
                if not top10.empty:
                    fig_t10rev = go.Figure(go.Bar(x=top10["_net_rev"], y=top10["Note"].str[:22], orientation="h",
                        marker_color=T2, text=[_usd(v) for v in top10["_net_rev"]], textposition="outside",
                        hovertemplate="<b>%{y}</b><br>Net Rev: $%{x:,.0f}<extra></extra>"))
                    fig_t10rev.update_layout(yaxis=dict(autorange="reversed"))
                    _tlabel = "All Floors" if plat_filter == "All" else FLOOR_LABELS.get(plat_filter, plat_filter)
                    st.markdown(f'<div class="card"><div class="card-title">Top {len(top10)} New Games — Net Revenue since release ({_tlabel})</div>', unsafe_allow_html=True)
                    st.plotly_chart(plotly_base(fig_t10rev, h=380, ml=140, mr=70, mt=6, mb=34), use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.info("No revenue data yet for new releases in this window.")

            today = pd.Timestamp.now().normalize()
            disp = gr_df.copy()
            disp["Days Live"] = (today - disp["LastDate"]).dt.days
            disp["_tc"] = disp["_tc"].map(lambda v: FLOOR_LABELS.get(v, v))
            disp["_vd"] = _generic_product(disp["_vd"])
            disp = disp.rename(columns={
                "_tc": "System", "_vd": "Game Category", "Note": "Game Name", "GameId": "Game ID",
                "Config": "Config", "LastDate": "Released", "Locations": "Locs at Launch",
                "_net_rev": "Net Revenue", "_hold_pct": "Hold %",
            })
            cols = ["System", "Game Category", "Game Name", "Game ID", "Config", "Released",
                    "Locs at Launch", "Days Live", "Net Revenue", "Hold %"]
            _tbl_h = 38 + 35 * (len(disp) + 1)  # stretch to fit every row — no inner scrollbar
            st.dataframe(
                disp[cols].style.format({
                    "Game ID": "{:.0f}",
                    "Released": lambda d: d.strftime("%d %b %Y") if pd.notna(d) else "–",
                    "Locs at Launch": "{:.0f}", "Days Live": "{:.0f}",
                    "Net Revenue": "${:,.0f}", "Hold %": "{:.1f}%",
                }, na_rep="–"),
                use_container_width=True, hide_index=True, height=_tbl_h,
            )
        else:
            st.info(f"No new releases found in the last {wn_days} days.")

    else:  # All Updates
        ribbon("All Updates", "Every game-release update log entry in this window", T2, "LOG")
        # Math/Payout/Pool rows are excluded — that feature is gone, this is game
        # releases and rollouts only.
        wn_df_games = wn_df[~wn_df["Category"].isin(["Payout", "Pool", "Progressive"])] if not wn_df.empty else wn_df
        if not wn_df_games.empty:
            tbl = (wn_df_games.groupby(["Studio", "ConfigPlatform", "RequiredProduct", "Category", "Note"], as_index=False, dropna=False)
                        .agg(LastDate=("LastDate", "max"), Locations=("Locations", "sum"))
                        .sort_values("LastDate", ascending=False))
            st.dataframe(
                tbl.rename(columns={"ConfigPlatform": "Platform", "RequiredProduct": "Product", "LastDate": "Date"})
                   .style.format({"Date": lambda d: d.strftime("%d %b %Y") if pd.notna(d) else "–", "Locations": "{:.0f}"}),
                use_container_width=True, hide_index=True, height=600,
            )
            st.caption(f"{len(tbl):,} unique entries · {wn_df_games['Locations'].sum():,.0f} total location updates · last {wn_days} days")
        else:
            st.info("No release data found.")

# ══════════════════════════════════════════════════════════════════
# SHARED — 30/60/90 Launch Report (.docx), triggered from 3 tabs
# ══════════════════════════════════════════════════════════════════
def _report_ordinal(n):
    n = int(round(n))
    if 11 <= (n % 100) <= 13:
        return f"{n}th"
    return f"{n}{ {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th') }"


def _report_fig_png(fig, width_in=6.3, height_px=380):
    """Export a Plotly figure as a white-background PNG for embedding in the report.
    Charts use a transparent paper/plot background on-screen (styled via CSS instead) —
    export needs an explicit white background or kaleido's default could go either way."""
    f2 = go.Figure(fig)
    f2.update_layout(paper_bgcolor="white", plot_bgcolor="white", height=height_px,
                      margin=dict(l=50, r=20, t=30, b=45))
    img = f2.to_image(format="png", width=int(width_in * 140), height=height_px, scale=2)
    return io.BytesIO(img)


def build_launch_report_bytes():
    """
    Assemble the 30/60/90 Launch Report for the currently-selected game as a .docx,
    reusing the exact same peer-pool/quick-score/forecast/footprint/cannibalization
    functions the live tabs call — this function is independent of which tab triggered
    it (Is It On Track / Similar Launches & What to Expect / Full Breakdown all call
    this one shared builder, so there is only ever one report, not three).
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def hexrgb(h):
        h = h.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    # ── compute everything first ─────────────────────────────────────────
    max_wk_r = int(gdf["launch_week"].max())
    weeks_live_r = max_wk_r + 1
    tenure_r = _tenure_band(weeks_live_r)

    peer_catalog_r = load_sql_catalog()
    if not peer_catalog_r.empty:
        peer_pool_r = L.find_peer_pool(peer_catalog_r, df, sel_id, scale_tolerance=scale_x)
        peer_df_r = df[df["game_id"].isin(peer_pool_r["ids"])]
    else:
        peer_pool_r = {"family": "whole fleet (no catalog)", "pool_size": int(df["game_id"].nunique()), "scale_applied": False}
        peer_df_r = df

    cur_row_r = gdf[gdf["launch_week"] == max_wk_r]
    cv_bet_r = float(cur_row_r["bet_handle"].iloc[0] or 0) if not cur_row_r.empty else 0.0
    cv_net_r = float(cur_row_r["net_rev"].iloc[0] or 0) if not cur_row_r.empty else 0.0
    cv_hold_r = float(cur_row_r["hold_pct"].iloc[0]) if not cur_row_r.empty and pd.notna(cur_row_r["hold_pct"].iloc[0]) else None
    game_lifetime_net_r = float(gdf["net_rev"].sum())
    plat_total_net_r = float(cat["total_net"].sum()) if not cat.empty else 0.0
    share_lifetime_r = min((game_lifetime_net_r / plat_total_net_r * 100), 99.9) if plat_total_net_r > 0 else None
    rank_list_r = cat.sort_values("total_net", ascending=False)["game_id"].tolist() if not cat.empty else []
    rank_n_r = (rank_list_r.index(sel_id) + 1) if sel_id in rank_list_r else None
    n_ranked_r = len(rank_list_r)

    # Exec status — same thresholds as Is It On Track's own badge, computed independently
    # here rather than calling into that tab's inline code (keeps this report from ever
    # being able to break the already-verified live page).
    exec_kpis_r = [("bet_handle", "Bet Handle", True), ("hold_pct", "Hold %", True),
                    ("bet_decay", "Bet Decay", False), ("net_rev", "Net Revenue", True)]
    n_r_r = n_a_r = 0
    flag_detail_r = []
    for kpi_e, lbl_e, hi_e in exec_kpis_r:
        bnd_e = bkpi(peer_df_r, kpi_e)
        brow_e = bnd_e[bnd_e["launch_week"] == max_wk_r] if not bnd_e.empty else pd.DataFrame()
        val_e = gdf[gdf["launch_week"] == max_wk_r][kpi_e].values
        if len(val_e) and not brow_e.empty and pd.notna(val_e[0]):
            cls_e = _flag_cls(float(val_e[0]), float(brow_e["p25"].iloc[0]), float(brow_e["p75"].iloc[0]),
                               float(brow_e["p10"].iloc[0]), float(brow_e["p90"].iloc[0]), higher_is_better=hi_e)
            n_r_r += cls_e == "r"; n_a_r += cls_e == "a"
            if cls_e in ("r", "a"):
                flag_detail_r.append([lbl_e, "Flag — outside P10–P90" if cls_e == "r" else "Watch — outside P25–P75"])
    eh_r = "NEEDS REVIEW" if n_r_r >= 2 else ("WATCH" if (n_r_r >= 1 or n_a_r >= 2) else "ON TRACK")
    # Guardrail: lifetime RTP > 100% (win > bet) is real, confirmed money lost -- never
    # let percentile-band noise on unrelated metrics soften that to WATCH/ON TRACK.
    if game_lifetime_net_r < 0:
        eh_r = "NEEDS REVIEW"

    # Trend + archetype — same formulas as Is It On Track's inline block, recomputed here
    # (not called into) for the same reason: this report must never be able to regress
    # that already-verified page.
    peak_net_r = float(gdf["net_rev"].max()) if gdf["net_rev"].notna().any() else 0.0
    last4_net_r = float(gdf[gdf["launch_week"] > max_wk_r - 4]["net_rev"].mean()) if len(gdf) else 0.0
    pct_of_peak_r = (last4_net_r / peak_net_r * 100) if peak_net_r > 0 else None
    roll_r = gdf.sort_values("launch_week")["net_rev"].rolling(4, min_periods=4).mean().dropna()
    if weeks_live_r < 13 or len(roll_r) < 4:
        trend_status_r = "Insufficient data — too early to classify (<13 weeks live)"
    else:
        declines_r = int((roll_r.diff() < 0).sum())
        if (pct_of_peak_r is not None and pct_of_peak_r < 20) or declines_r >= len(roll_r) - 1:
            trend_status_r = "Declining"
        elif roll_r.iloc[-1] >= roll_r.iloc[0]:
            trend_status_r = "Growing"
        else:
            trend_status_r = "Stable"
    bh_series_r = gdf.sort_values("launch_week")["bet_handle"]
    peak_wk_r = int(gdf.loc[gdf["bet_handle"].idxmax(), "launch_week"]) if gdf["bet_handle"].notna().any() else 0
    w0_bh_r = float(bh_series_r.iloc[0]) if len(bh_series_r) else 0.0
    _early_wk_r = min(12, max_wk_r)
    _early_row_r = gdf[gdf["launch_week"] == _early_wk_r]
    _early_bh_r = float(_early_row_r["bet_handle"].iloc[0]) if not _early_row_r.empty and pd.notna(_early_row_r["bet_handle"].iloc[0]) else None
    _early_pct_r = (_early_bh_r / w0_bh_r * 100) if (_early_bh_r is not None and w0_bh_r > 0) else None
    if weeks_live_r < 8:
        archetype_r = "Still forming"
    elif pct_of_peak_r is not None and pct_of_peak_r >= 60:
        archetype_r = "Evergreen"
    elif peak_wk_r <= 1 and _early_pct_r is not None and _early_pct_r < 30:
        archetype_r = "Flash-in-the-pan"
    elif peak_wk_r >= 3:
        archetype_r = "Slow-burn"
    else:
        archetype_r = "Steady decliner"
    archetype_defs_r = {
        "Still forming": "Too early to classify (<8 weeks live)",
        "Evergreen": f"Still {pct_of_peak_r:.0f}% of its own peak" if pct_of_peak_r is not None else "Still near its own peak",
        "Flash-in-the-pan": "Peaked at launch, crashed within ~12 weeks",
        "Slow-burn": "Built up over several weeks before peaking",
        "Steady decliner": "Peaked early, faded gradually over time",
    }

    # Footprint / concentration
    foot_r = load_game_footprint(platform, sel_id)
    if foot_r.empty:
        n_locs_r, top_share_r, conc_label_r = 0, None, "–"
    else:
        n_locs_r = int(foot_r["loc_id"].nunique())
        total_loc_net_r = float(foot_r["net_rev"].sum())
        top_loc_net_r = float(foot_r.sort_values("net_rev", ascending=False)["net_rev"].iloc[0]) if total_loc_net_r != 0 else 0.0
        top_share_r = min((top_loc_net_r / total_loc_net_r * 100), 99.9) if total_loc_net_r > 0 else None
        conc_label_r = ("Concentrated" if (top_share_r is not None and top_share_r >= 40)
                        else "Watch" if (top_share_r is not None and top_share_r >= 25) else "Spread out")

    # Soft-launch check — reuses the same L.detect_ramp() already used by the live "Is It
    # On Track" tab (store/casino count growth in the first 4 weeks), not a new heuristic.
    # Paired with the real week-0 location count (read directly off gdf, not invented) so a
    # suspiciously thin week-0 (e.g. a 1-2 location pilot day) is visible next to the
    # recorded launch date rather than silently feeding a distorted bet_decay baseline.
    # NOT applied to EdgeLabs: launch_date there is defined as the first real-money spin
    # (see load_edgelabs_weekly_all) — there's no separate "soft launch" phase to detect,
    # and casino count naturally growing after launch (more operators picking the game up)
    # isn't the same thing as a land-based pilot rollout.
    ramp_r = L.detect_ramp(df, sel_id) if platform != "EdgeLabs" else {"is_ramping": False, "note": ""}
    w0_stores_r = None
    if "stores" in gdf.columns:
        _w0_rows_r = gdf[gdf["launch_week"] == 0]["stores"].dropna()
        if not _w0_rows_r.empty:
            w0_stores_r = int(_w0_rows_r.iloc[0])
    # Week-0 player count — a soft-launch/test period on Pong (PFH's online identity)
    # shows up as a handful of test accounts before the real launch weeks later.
    # NOT applied to EdgeLabs: there, launch = first spin by definition, so a thin
    # week-0 player count isn't a "soft launch" signal, just early ramp-up.
    w0_players_r = None
    if platform != "EdgeLabs" and "players" in gdf.columns:
        _w0_prows_r = gdf[gdf["launch_week"] == 0]["players"].dropna()
        if not _w0_prows_r.empty:
            w0_players_r = int(_w0_prows_r.iloc[0])
    _soft_launch_bits_r = []
    if w0_stores_r is not None:
        _soft_launch_bits_r.append(f"{w0_stores_r} location{'s' if w0_stores_r != 1 else ''}")
    if w0_players_r is not None:
        _soft_launch_bits_r.append(f"{w0_players_r} player{'s' if w0_players_r != 1 else ''}")
    _soft_launch_val_r = f"{meta['launch_date']}" + (f" ({', '.join(_soft_launch_bits_r)})" if _soft_launch_bits_r else "")

    # Peers (DTW top-5) + percentile radar
    with_peers_r = find_peers_scaled(df, sel_id, kpi="bet_decay", n_weeks=n_match, top_k=5, scale_tolerance=scale_x)
    rk_metrics_r = [("hold_pct", "Hold %"), ("bet_decay", "Decay"), ("arpu", "ARPU"),
                    ("avg_bet", "Avg bet"), ("spp", "Spins/P"), ("player_decay", "Pl. decay")]
    peer_pool_ids_incl_self_r = peer_pool_r["ids"] | {sel_id} if isinstance(peer_pool_r.get("ids"), set) else {sel_id}
    peer_df_for_rank_r = df[df["game_id"].isin(peer_pool_ids_incl_self_r)]
    rk_pcts_r, rk_widened_r = [], []
    for k, _ in rk_metrics_r:
        p = L.pct_rank(peer_df_for_rank_r, sel_id, k, max_wk_r)
        wid = False
        if pd.isna(p):
            p = L.pct_rank(df, sel_id, k, max_wk_r)
            wid = True
        rk_pcts_r.append(p); rk_widened_r.append(wid)

    # First 30/60/90 + peer-median expected
    band_net_r = bkpi(peer_df_r, "net_rev")
    milestones_r = []
    for lbl, days in (("First 30 Days", 30), ("First 60 Days", 60), ("First 90 Days", 90)):
        wk_cut = days // 7
        if max_wk_r < wk_cut:
            milestones_r.append({"label": lbl, "actual": None, "expected": None})
            continue
        actual = float(gdf[gdf["launch_week"] <= wk_cut]["net_rev"].sum())
        bnd_cut = band_net_r[band_net_r["launch_week"] <= wk_cut]
        expected = float(bnd_cut["p50"].sum()) if not bnd_cut.empty else None
        milestones_r.append({"label": lbl, "actual": actual, "expected": expected})

    # Forecast
    blended_r = L.fit_blended_forecast(df, sel_id, peer_catalog_r, n_forecast_weeks=horizon, kpi="bet_handle", scale_tolerance=scale_x)
    rh_r = blended_r["reliable_horizon"]
    avg_hold_r = float(gdf["hold_pct"].dropna().mean()) / 100.0 if gdf["hold_pct"].notna().any() else 0.05
    avg_hold_r = max(0.01, min(avg_hold_r, 0.5))
    fcast_r = blended_r["forecast"]
    wks_rel_r = [f["week"] for f in fcast_r if not f["directional"]]
    base_rel_r = [f["value"] * avg_hold_r for f in fcast_r if not f["directional"]]
    next_n_net_r = sum(f["value"] * avg_hold_r for f in fcast_r)

    # Cannibalization — only if the game has cleared the same 4-week gate the live check uses
    cannib_rows_r = []
    cannib_note_r = None
    if platform == "EdgeLabs":
        cannib_note_r = "Not available — EdgeLabs is casino-based (no location dimension)."
    elif weeks_live_r < 4:
        cannib_note_r = "Not enough live weeks yet (<4) to run a before/after location comparison."
    else:
        tgt_dt_r = pd.to_datetime(meta["launch_date"], errors="coerce")
        if pd.isna(tgt_dt_r):
            cannib_note_r = "Launch date unavailable — cannot compute cannibalization window."
        else:
            cw_r = 4
            before_start_r, before_end_r = tgt_dt_r - pd.Timedelta(weeks=cw_r), tgt_dt_r - pd.Timedelta(days=1)
            after_start_r, after_end_r = tgt_dt_r, min(tgt_dt_r + pd.Timedelta(weeks=cw_r), pd.Timestamp.today())
            geo_life_r = load_geo_detail(platform, str(_LIFETIME_START), str(dt.date.today()))
            geo_pre_r = load_geo_detail(platform, str(before_start_r.date()), str(before_end_r.date()))
            geo_post_r = load_geo_detail(platform, str(after_start_r.date()), str(after_end_r.date()))
            target_locs_r = set(geo_life_r[geo_life_r["game_id"].astype("Int64") == int(sel_id)]["location_name"]) if not geo_life_r.empty else set()
            if not target_locs_r:
                cannib_note_r = f"No location history found for {top_title} — cannot scope cannibalization to its footprint."
            elif geo_pre_r.empty or geo_post_r.empty:
                cannib_note_r = "Not enough location data both before and after launch to compute cannibalization."
            else:
                pre_s_r = geo_pre_r[geo_pre_r["location_name"].isin(target_locs_r) & (geo_pre_r["game_id"].astype("Int64") != int(sel_id)) & ~geo_pre_r["game_id"].apply(L._is_hr)]
                post_s_r = geo_post_r[geo_post_r["location_name"].isin(target_locs_r) & (geo_post_r["game_id"].astype("Int64") != int(sel_id)) & ~geo_post_r["game_id"].apply(L._is_hr)]
                pre_g_r = pre_s_r.groupby(["game_id", "game_name"])["bet"].sum().rename("bet_before")
                post_g_r = post_s_r.groupby(["game_id", "game_name"])["bet"].sum().rename("bet_after")
                cdf_r = pd.concat([pre_g_r, post_g_r], axis=1).reset_index().fillna(0)
                cdf_r = cdf_r[cdf_r["bet_before"] > 0]
                if cdf_r.empty:
                    cannib_note_r = f"No other games had bet activity at {top_title}'s {len(target_locs_r)} locations before its launch."
                else:
                    cdf_r["change"] = cdf_r["bet_after"] - cdf_r["bet_before"]
                    cdf_r["change_pct"] = cdf_r["change"] / cdf_r["bet_before"] * 100
                    cdf_r = cdf_r.sort_values("change")
                    n_drop_r = int((cdf_r["change"] < 0).sum())
                    cannib_note_r = (f"Scoped to the {len(target_locs_r)} locations where {top_title} is live. "
                                     f"{n_drop_r} of {len(cdf_r)} other games there saw a bet-handle drop in the "
                                     f"{cw_r} weeks after launch.")
                    cannib_rows_r = cdf_r.head(5).to_dict("records")

    # ── build the document ───────────────────────────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def h_title(text, size=18):
        p = doc.add_paragraph(); r = p.add_run(text)
        r.font.size = Pt(size); r.bold = True; r.font.color.rgb = hexrgb(TEXT)
        p.paragraph_format.space_after = Pt(2)
        return p

    def h_section(num, title):
        p = doc.add_paragraph(); r = p.add_run(f"{num}.  {title}")
        r.font.size = Pt(13); r.bold = True; r.font.color.rgb = hexrgb(TEXT)
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), BORDER.lstrip("#"))
        pbdr.append(bottom); pPr.append(pbdr)
        return p

    def p_body(text, size=10.5, bold=False, italic=False, color=T2):
        p = doc.add_paragraph(); r = p.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = hexrgb(color)
        p.paragraph_format.space_after = Pt(6)
        return p

    def kv_line(label, value, note=""):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}:  "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = hexrgb(TEXT)
        r2 = p.add_run(str(value)); r2.font.size = Pt(10.5); r2.font.color.rgb = hexrgb(TEXT)
        if note:
            r3 = p.add_run(f"   —  {note}"); r3.italic = True; r3.font.size = Pt(9.5); r3.font.color.rgb = hexrgb(T2)
        p.paragraph_format.space_after = Pt(3)
        return p

    def add_table(headers, rows):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Light Grid Accent 1"
        for i, htext in enumerate(headers):
            c = tbl.rows[0].cells[i]; c.text = ""
            r = c.paragraphs[0].add_run(htext); r.bold = True; r.font.size = Pt(9.5)
        for row in rows:
            cells = tbl.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                r = cells[i].paragraphs[0].add_run(str(val)); r.font.size = Pt(9.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_chart(fig, height_px=340):
        buf = _report_fig_png(fig, height_px=height_px)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(buf, width=Inches(6.3))
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # 0. Header ------------------------------------------------------------
    h_title(top_title)
    p = doc.add_paragraph(); r = p.add_run(
        f"{platform} · ID {sel_id} · Launched {meta['launch_date']} · {weeks_live_r} weeks live ({tenure_r}) · "
        f"Generated {dt.date.today().strftime('%d %b %Y')}")
    r.font.size = Pt(10); r.font.color.rgb = hexrgb(T2)
    p.paragraph_format.space_after = Pt(10)
    kv_line("Overall Status", eh_r)
    kv_line("Peer basis", f"{peer_pool_r['pool_size']} games ({peer_pool_r['family']})")

    # 1. Executive Summary --------------------------------------------------
    h_section(1, "Executive Summary")
    _rank_bit = f"#{rank_n_r} of {n_ranked_r} {platform} games by lifetime Game Net" if rank_n_r else "unranked (no lifetime net data)"
    _share_bit = f"{share_lifetime_r:.1f}% of {platform}'s lifetime Game Net" if share_lifetime_r is not None else "an unknown share of platform net"
    summary_bits = [
        f"{top_title} is {weeks_live_r} weeks into launch and currently rated {eh_r}.",
        f"It ranks {_rank_bit}, contributing {_share_bit}.",
        f"Trend: {trend_status_r}. Archetype: {archetype_r} — {archetype_defs_r.get(archetype_r, '')}.",
    ]
    if cannib_note_r and cannib_rows_r:
        summary_bits.append("A cannibalization check at this game's own locations found measurable drops in other titles' bet handle — see Section 8.")
    p_body(" ".join(summary_bits), color=TEXT)

    # 2. Reach & Footprint ---------------------------------------------------
    h_section(2, "Reach & Footprint")
    kv_line("Soft Launch Date", _soft_launch_val_r)
    if ramp_r.get("is_ramping"):
        p_body(ramp_r["note"], color=TEXT)
    kv_line("Casinos live (lifetime)", f"{n_locs_r:,}" if n_locs_r else "–")
    kv_line("Concentration", conc_label_r,
            f"top location is {top_share_r:.1f}% of lifetime Game Net" if top_share_r is not None else "")
    if conc_label_r == "Concentrated":
        p_body("Most of this game's earnings come from one location — that's a real risk: a single site closing or "
              "swapping the terminal would remove most of the game's revenue, not just a slice of it.")
    elif conc_label_r == "Watch":
        p_body("Earnings are moderately concentrated — worth keeping an eye on, not yet a red flag.")
    if not foot_r.empty and n_locs_r >= 2:
        _foot_top_r = foot_r[foot_r["net_rev"] > 0].sort_values("net_rev", ascending=False)
        if len(_foot_top_r) > 12:
            _other_sum_r = float(_foot_top_r["net_rev"].iloc[12:].sum())
            _tm_labels_r = [f"Loc {lid}" for lid in _foot_top_r["loc_id"].iloc[:12]] + (["Other locations"] if _other_sum_r > 0 else [])
            _tm_values_r = _foot_top_r["net_rev"].iloc[:12].tolist() + ([_other_sum_r] if _other_sum_r > 0 else [])
        else:
            _tm_labels_r = [f"Loc {lid}" for lid in _foot_top_r["loc_id"]]
            _tm_values_r = _foot_top_r["net_rev"].tolist()
        if _tm_values_r:
            add_chart(_treemap_fig(_tm_labels_r, _tm_values_r, height=280), height_px=280)
            p_body("Lifetime Game Net by location — tile size = share of this game's total", size=9, italic=True)
        _loc_total_r = float(foot_r["net_rev"].sum())
        loc_rows = [[f"Loc {r['loc_id']}", _usd(r["net_rev"]),
                     f"{min(r['net_rev'] / _loc_total_r * 100, 99.9):.1f}%" if _loc_total_r > 0 else "–"]
                    for _, r in _foot_top_r.head(5).iterrows()]
        if loc_rows:
            add_table(["Top Locations", "Lifetime Game Net", "% of Total"], loc_rows)

    # 3. The Money ------------------------------------------------------------
    h_section(3, "The Money")
    p_body("Game Net is the house hold (Bet − Win) — not Pong's royalty. A confirmed royalty figure isn't wired into "
          "this report yet, so every dollar figure below is Game Net.", italic=True, size=9.5)
    kv_line(f"Week {max_wk_r} Bet Handle", _usd(cv_bet_r))
    kv_line(f"Week {max_wk_r} Net Revenue", _usd(cv_net_r))
    kv_line("Hold % (latest)", _pct(cv_hold_r) if cv_hold_r is not None else "–")
    kv_line("Lifetime Game Net", _usd(game_lifetime_net_r))

    wks_x_r = sorted(gdf["launch_week"].dropna().unique().astype(int).tolist())
    band_bh_r = bkpi(peer_df_r, "bet_handle")

    def _band_at(bd, col, w):
        rrow = bd[bd["launch_week"] == w] if not bd.empty else pd.DataFrame()
        return float(rrow[col].iloc[0]) if not rrow.empty and col in rrow.columns and pd.notna(rrow[col].iloc[0]) else None

    def _series_at(col, w):
        rrow = gdf[gdf["launch_week"] == w]
        return float(rrow[col].iloc[0]) if not rrow.empty and pd.notna(rrow[col].iloc[0]) else None

    fig_bh_r = go.Figure()
    fill_band(fig_bh_r, wks_x_r, [_band_at(band_bh_r, "p10", w) for w in wks_x_r], [_band_at(band_bh_r, "p90", w) for w in wks_x_r],
              "rgba(91,89,78,0.10)", "Peer range (P10–P90)")
    fig_bh_r.add_trace(go.Scatter(x=wks_x_r, y=[_band_at(band_bh_r, "p50", w) for w in wks_x_r], mode="lines",
                                   line=dict(color=T3, dash="dot", width=1.5), name="Peer median"))
    fig_bh_r.add_trace(go.Scatter(x=wks_x_r, y=[_series_at("bet_handle", w) for w in wks_x_r], mode="lines+markers",
                                   line=dict(color=MOVE_DOWN, width=2.5), marker=dict(size=5), name=top_title[:22]))
    fig_bh_r.update_xaxes(title_text="Launch week", automargin=True)
    fig_bh_r.update_yaxes(title_text="Bet handle ($)", automargin=True)
    add_chart(fig_bh_r)
    p_body("Bet Handle vs. Peer Range — week by week", size=9, italic=True)

    _recent_r = gdf.sort_values("launch_week").tail(6)
    rw_rows = [[f"Wk {int(r['launch_week'])}",
                _usd(r["bet_handle"]) if pd.notna(r.get("bet_handle")) else "–",
                _usd(r["net_rev"]) if pd.notna(r.get("net_rev")) else "–",
                _pct(r["hold_pct"]) if pd.notna(r.get("hold_pct")) else "–"]
               for _, r in _recent_r.iterrows()]
    if rw_rows:
        add_table(["Week", "Bet Handle", "Net Revenue", "Hold %"], rw_rows)
        p_body(f"Actual results for the last {len(rw_rows)} weeks.", size=9, italic=True)

    # 4. First 30/60/90 Days ---------------------------------------------------
    h_section(4, "First 30 / 60 / 90 Days")
    ms_rows = []
    for m in milestones_r:
        if m["actual"] is None:
            ms_rows.append([m["label"], "Not reached yet", "–", "–"])
            continue
        delta = (m["actual"] - m["expected"]) / m["expected"] * 100 if m["expected"] not in (None, 0) else None
        delta_str = f"{'▲' if delta >= 0 else '▼'} {abs(delta):.0f}%" if delta is not None else "–"
        ms_rows.append([m["label"], _usd(m["actual"]), _usd(m["expected"]) if m["expected"] is not None else "–", delta_str])
    add_table(["Milestone", "Actual Net Revenue", "Expected (peer median)", "vs. Expected"], ms_rows)
    p_body("\"Expected\" is the peer group's median Net Revenue accumulated over the same weeks — not a target, "
          "just what similar games typically did by that point.", size=9.5, italic=True)

    # 5. Engagement & Player Behavior -------------------------------------------
    h_section(5, "Engagement & Player Behavior")
    band_decay_r = bkpi(peer_df_r, "bet_decay")
    fig_bd_r = go.Figure()
    fill_band(fig_bd_r, wks_x_r, [_band_at(band_decay_r, "p10", w) for w in wks_x_r], [_band_at(band_decay_r, "p90", w) for w in wks_x_r],
              "rgba(91,89,78,0.10)", "Peer range (P10–P90)")
    fig_bd_r.add_trace(go.Scatter(x=wks_x_r, y=[_band_at(band_decay_r, "p50", w) for w in wks_x_r], mode="lines",
                                   line=dict(color=T3, dash="dot", width=1.5), name="Peer median"))
    fig_bd_r.add_trace(go.Scatter(x=wks_x_r, y=[_series_at("bet_decay", w) for w in wks_x_r], mode="lines+markers",
                                   line=dict(color=MOVE_DOWN, width=2.5), marker=dict(size=5), name=top_title[:22]))
    fig_bd_r.update_xaxes(title_text="Launch week", automargin=True)
    fig_bd_r.update_yaxes(title_text="Bet decay % (week 0 = 100)", automargin=True)
    add_chart(fig_bd_r)
    p_body("Bet Decay vs. Peer Range — is engagement holding up as well as similar games?", size=9, italic=True)
    latest_hold_r = float(gdf["hold_pct"].dropna().iloc[-1]) if gdf["hold_pct"].notna().any() else None
    latest_decay_r = float(gdf["bet_decay"].dropna().iloc[-1]) if gdf["bet_decay"].notna().any() else None
    if weeks_live_r < 13:
        p_body("Too early (under 13 weeks live) to separate a math problem from an appeal problem.")
    elif latest_hold_r is not None and latest_hold_r < 6 and (latest_decay_r is None or latest_decay_r > 50):
        p_body("Thin hold (<6%) with engagement still reasonably strong — looks like a math/config issue, not a demand problem.")
    elif latest_decay_r is not None and latest_decay_r < 30:
        p_body("Engagement has faded sharply — looks like an appeal problem, not a math one.")

    # 6. How It Compares to Peers ------------------------------------------------
    h_section(6, "How It Compares to Peers")
    if not with_peers_r:
        p_body("No DTW-matched peers found for this game within the current match window/scale settings.")
    else:
        peer_rows = [[p["game_name"][:30], f"{_similarity_pct(p['distance'], p.get('n_match_weeks')):.0f}%", p.get("total_weeks", "–")] for p in with_peers_r]
        add_table(["Closest Peer", "Shape Similarity", "Total Weeks Live"], peer_rows)
        peer_snap_rows = []
        for p in with_peers_r:
            pgid = p.get("game_id")
            pdf_r = df[df["game_id"] == pgid].sort_values("launch_week") if pgid is not None else pd.DataFrame()
            if pdf_r.empty:
                continue
            plast = pdf_r.iloc[-1]
            peer_snap_rows.append([
                p["game_name"][:26], f"Wk {int(plast['launch_week'])}",
                _usd(plast["bet_handle"]) if pd.notna(plast.get("bet_handle")) else "–",
                _pct(plast["hold_pct"]) if pd.notna(plast.get("hold_pct")) else "–",
            ])
        if peer_snap_rows:
            add_table(["Peer", "At Week", "Bet Handle", "Hold %"], peer_snap_rows)
            p_body("Each peer's own most recent week — a quick check they're a comparable scale, not a copy of this game's numbers.",
                  size=9, italic=True)
    rk_bits = [f"{lbl} {_report_ordinal(p)} pct" + (" (fleet-wide)" if wid else "")
              for (_, lbl), p, wid in zip(rk_metrics_r, rk_pcts_r, rk_widened_r) if pd.notna(p)]
    if rk_bits:
        p_body("Percentile rank vs. peer pool at the current week — " + ", ".join(rk_bits) + ". "
              "Player Decay is not direction-flipped: a high percentile there means high decay (worse), same as every other metric here.",
              size=9.5, italic=True)

    # 7. Trend & Where It's Heading -----------------------------------------------
    h_section(7, "Trend & Where It's Heading")
    kv_line("Trend", trend_status_r)
    kv_line("Archetype", archetype_r, archetype_defs_r.get(archetype_r, ""))
    if blended_r.get("self_weight") is not None:
        kv_line("Forecast blend", f"self-trend {blended_r['self_weight']*100:.0f}% · peer-trend {blended_r['peer_weight']*100:.0f}% "
                f"({len(blended_r['peers_used'])} peers)")
    if rh_r < horizon:
        p_body(f"Reliable through week +{rh_r} only — figures beyond that are directional, not a confident forecast.", size=9.5, italic=True)
    kv_line(f"Projected Net Revenue, next {horizon} weeks", _usd(next_n_net_r))
    if wks_rel_r:
        fig_fc_r = go.Figure()
        net_hist_r = gdf["net_rev"].fillna(0).values
        wks_hist_r = list(range(len(net_hist_r)))
        p10h = [_band_at(band_net_r, "p10", w) for w in wks_hist_r]
        p90h = [_band_at(band_net_r, "p90", w) for w in wks_hist_r]
        p50h = [_band_at(band_net_r, "p50", w) for w in wks_hist_r]
        fill_band(fig_fc_r, wks_hist_r, p10h, p90h, "rgba(91,89,78,0.10)", "Peer range (P10–P90)")
        fig_fc_r.add_trace(go.Scatter(x=wks_hist_r, y=p50h, mode="lines", line=dict(color=T3, dash="dot", width=1.2), name="Peer median"))
        fig_fc_r.add_trace(go.Scatter(x=wks_hist_r, y=list(net_hist_r), mode="lines+markers",
                                       line=dict(color=TEXT, width=2.2), marker=dict(size=4), name="Actual"))
        fig_fc_r.add_trace(go.Scatter(x=wks_rel_r, y=base_rel_r, mode="lines", line=dict(color=MOVE_UP, width=2, dash="dash"), name="Forecast"))
        fig_fc_r.update_xaxes(title_text="Launch week", automargin=True)
        fig_fc_r.update_yaxes(title_text="Net revenue ($)", automargin=True)
        add_chart(fig_fc_r)
        p_body("Net Revenue — actual vs. peer range, with a blended self-trend/peer-trend forecast", size=9, italic=True)
        fc_rows = []
        for f in fcast_r:
            tag = " (directional)" if f.get("directional") else ""
            fc_rows.append([f"Week +{f['week']}{tag}", _usd(f["value"]), _usd(f["value"] * avg_hold_r)])
        if fc_rows:
            add_table(["Week", "Projected Bet Handle", "Projected Net Revenue"], fc_rows)

    # 8. Cannibalization Check --------------------------------------------------
    h_section(8, "Cannibalization Check")
    p_body(cannib_note_r or "Not run.", color=TEXT)
    if cannib_rows_r:
        cb_rows = [[r_["game_name"][:28], _usd(r_["bet_before"]), _usd(r_["bet_after"]),
                   f"{'▲' if r_['change'] >= 0 else '▼'} {_usd(abs(r_['change']))}"] for r_ in cannib_rows_r]
        add_table(["Game (same locations)", "Bet Before", "Bet After", "Change"], cb_rows)
        p_body("A drop in another game's bet handle at the same locations does not prove cannibalization — "
              "seasonal patterns or terminal swaps can also explain it. Treat this as a starting point, not proof.",
              size=9.5, italic=True)

    # 9. Verdict & Next Check-in --------------------------------------------------
    h_section(9, "Verdict & Next Check-in")
    kv_line("Status", eh_r)
    verdict_map = {
        "ON TRACK": "Performing in line with or better than peers. No action needed beyond the next scheduled check-in.",
        "WATCH": "One or more metrics are outside the normal peer range. Worth a closer look, not yet urgent.",
        "NEEDS REVIEW": "Multiple metrics are outside the normal peer range. Recommend review before the next milestone.",
    }
    p_body(verdict_map.get(eh_r, ""), color=TEXT)
    if flag_detail_r:
        add_table(["Metric", "Status"], flag_detail_r)
    else:
        p_body("No metrics currently outside the normal peer range.", size=9.5, italic=True)
    next_check_r = dt.date.today() + dt.timedelta(days=30 if weeks_live_r < 13 else 60)
    kv_line("Suggested next check-in", next_check_r.strftime("%d %b %Y"))

    buf_doc = io.BytesIO()
    doc.save(buf_doc)
    buf_doc.seek(0)
    return buf_doc.getvalue()


def render_generate_report_button(widget_key: str):
    """Shared 'Generate Report' block — dropped at the top of Is It On Track, Similar
    Launches & What to Expect, and Full Breakdown. All three call this same function so
    there's exactly one report definition, not three drifting copies. Caches the built
    bytes in session_state per game so clicking Generate on one tab makes the download
    immediately available on the other two without rebuilding."""
    if is_all:
        return
    cache_key = f"_report_bytes_{sel_id}"
    rc1, rc2 = st.columns([1, 3])
    with rc1:
        clicked = st.button("📄 Generate 30/60/90 Report", key=f"gen_report_{widget_key}", use_container_width=True)
    with rc2:
        if clicked:
            with st.spinner("Generating report — pulling peers, forecast, and footprint data…"):
                try:
                    st.session_state[cache_key] = build_launch_report_bytes()
                except Exception as e:
                    st.session_state.pop(cache_key, None)
                    st.error(f"Report generation failed: {e}")
        if cache_key in st.session_state:
            fname = f"{str(top_title).replace(' ', '_').replace('/', '-')}_ID{sel_id}_LaunchReport_{dt.date.today().isoformat()}.docx"
            st.download_button("⬇ Download Report (.docx)", data=st.session_state[cache_key],
                               file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               key=f"dl_report_{widget_key}", use_container_width=True)


# ══════════════════════════════════════════════════════════════════
# TAB — IS IT ON TRACK?
# ══════════════════════════════════════════════════════════════════
with tab_track:
    render_generate_report_button("track")
    peer_catalog = load_sql_catalog()
    if is_all:
        ribbon("All Games Ranked — Is Each One On Track?",
               f"{len(nonhr)} {platform} launches, each compared to peers of similar launch size", T2, "STATUS")

        # ── Fleet financial health — period-scoped $, CEO-style ─────────────────────
        ceo_period = st.radio("Period", _HS_PERIODS, index=0, horizontal=True,
                              label_visibility="collapsed", key="track_ceo_period")
        _cp_cur, _cp_pri = _hs_get_dates(ceo_period)
        _cp_s, _cp_e = _cp_cur
        _cp_ps, _cp_pe = _cp_pri
        _cp_cal = _load_calendar_daily(platform, str(_cp_s), str(_cp_e))
        _cp_net = float(_cp_cal["net_rev"].sum()) if not _cp_cal.empty else 0.0
        _cp_bet = float(_cp_cal["bet"].sum()) if not _cp_cal.empty else 0.0
        _cp_hold = (_cp_net / _cp_bet) * 100 if _cp_bet else None
        if ceo_period == "Lifetime":
            _cp_net_d = _cp_bet_d = _cp_hold_d = None
        else:
            _cp_cal_pri = _load_calendar_daily(platform, str(_cp_ps), str(_cp_pe))
            _cp_net_pri = float(_cp_cal_pri["net_rev"].sum()) if not _cp_cal_pri.empty else 0.0
            _cp_bet_pri = float(_cp_cal_pri["bet"].sum()) if not _cp_cal_pri.empty else 0.0
            _cp_hold_pri = (_cp_net_pri / _cp_bet_pri) * 100 if _cp_bet_pri else None
            _cp_net_d = _period_delta(_cp_net, _cp_net_pri)
            _cp_bet_d = _period_delta(_cp_bet, _cp_bet_pri)
            _cp_hold_d = _period_delta(_cp_hold, _cp_hold_pri) if (_cp_hold is not None and _cp_hold_pri is not None) else None
        krow([
            {"label": "Game Net ($)", "value": _usd(_cp_net), "delta_pct": _cp_net_d,
             "sub": "house hold, not Pong royalty"},
            {"label": "Bet ($)", "value": _usd(_cp_bet), "delta_pct": _cp_bet_d},
            {"label": "Hold %", "value": _pct(_cp_hold, 2) if _cp_hold is not None else "–", "delta_pct": _cp_hold_d},
        ])

        rows = []
        for _, gc_ in nonhr.iterrows():
            gid2 = gc_["game_id"]
            g2 = df[df["game_id"] == gid2].sort_values("launch_week")
            if g2.empty:
                continue
            mwk2 = int(g2["launch_week"].max())
            hl2 = float(g2[g2["hold_pct"].notna()].iloc[-1]["hold_pct"]) if g2["hold_pct"].notna().any() else np.nan
            bd2 = float(g2[g2["bet_decay"].notna()].iloc[-1]["bet_decay"]) if g2["bet_decay"].notna().any() else np.nan
            nr2 = float(g2["net_rev"].sum())

            # Peer pool for THIS game (SkinOf > Platform+Product+Orientation > relax > loose)
            # instead of the whole-fleet df — a Pull-Tab title shouldn't be graded against
            # P2P/Sweeps giants just because they're on the same platform.
            if not peer_catalog.empty:
                pool2 = L.find_peer_pool(peer_catalog, df, gid2, scale_tolerance=scale_x)
                peer_df2 = df[df["game_id"].isin(pool2["ids"])]
            else:
                peer_df2 = df

            flags = []
            for kp2, hi2 in [("bet_handle", True), ("hold_pct", True), ("bet_decay", True), ("net_rev", True)]:
                bnd2 = bkpi(peer_df2, kp2)
                br2 = bnd2[bnd2["launch_week"] == mwk2] if not bnd2.empty else pd.DataFrame()
                vv2 = g2[g2["launch_week"] == mwk2][kp2].values
                if len(vv2) and not br2.empty and pd.notna(vv2[0]):
                    flags.append(_flag_cls(float(vv2[0]), float(br2["p25"].iloc[0]), float(br2["p75"].iloc[0]),
                                            float(br2["p10"].iloc[0]), float(br2["p90"].iloc[0]), higher_is_better=hi2))
            n_r, n_a = flags.count("r"), flags.count("a")
            status = "NEEDS REVIEW" if n_r >= 2 else ("WATCH" if (n_r >= 1 or n_a >= 2) else "ON TRACK")
            # Guardrail: lifetime RTP > 100% (win > bet) is real, confirmed money lost --
            # never let percentile-band noise on unrelated metrics soften that to WATCH/ON
            # TRACK. Caught live on V2's "Bank It"/"Bank It 2"/"Real Reels" family (5 GameIds,
            # RTP 220-390% lifetime) -- percentile-only flagging treated a $500 miss and a
            # $26M miss as the same-weight signal, which is exactly what this closes.
            if nr2 < 0:
                status = "NEEDS REVIEW"

            rows.append({"game_id": gid2, "Game Name": gc_["game_name"], "Launch Date": str(gc_["launch_date"]),
                         "Tenure": _tenure_band(mwk2 + 1), "Weeks Live": mwk2 + 1, "Hold % (latest)": hl2,
                         "Bet Decay (latest)": bd2, "Net Revenue (lifetime)": nr2, "Overall Status": status})

        t2 = pd.DataFrame(rows)
        if not t2.empty:
            _fleet_lifetime_net = float(t2["Net Revenue (lifetime)"].sum())
            t2["Share of Fleet %"] = ((t2["Net Revenue (lifetime)"] / _fleet_lifetime_net * 100).clip(upper=99.9)) if _fleet_lifetime_net > 0 else 0.0
            n_ok, n_wa, n_nr = (t2["Overall Status"] == "ON TRACK").sum(), (t2["Overall Status"] == "WATCH").sum(), (t2["Overall Status"] == "NEEDS REVIEW").sum()

            # ── Platform mix — always all 3 "other platforms", plus V2's own product mix.
            # Two separate donuts instead of one combined cross-platform pie — V1/PFH/EdgeLabs
            # don't have a useful Product breakdown (PFH's catalog Product is just one coarse
            # value, "pfh-edgelabs"/"Unknown" — confirmed earlier), so they're compared to each
            # other as whole platforms; V2 is the one platform with real Product diversity
            # (P2P/PullTabs/Sweeps/Class2/HHR), so it gets its own breakdown.
            _op_vals = {}
            for _op in ("V1", "PFH", "EdgeLabs"):
                _op_cal = _load_calendar_daily(_op, str(_cp_s), str(_cp_e))
                _op_vals[_op] = float(_op_cal["net_rev"].sum()) if not _op_cal.empty else 0.0

            _v2_prod_cur = _load_period_v2v1("V2", str(_cp_s), str(_cp_e))
            _v2_cat = peer_catalog[peer_catalog["Platform"].astype(str).str.lower() == "v2"] if not peer_catalog.empty else pd.DataFrame()
            _v2_mix = pd.DataFrame()
            if not _v2_prod_cur.empty and not _v2_cat.empty:
                _v2_cat_prod = (_v2_cat[["Id", "Product"]].rename(columns={"Id": "game_id"})
                                .assign(game_id=lambda d: pd.to_numeric(d["game_id"], errors="coerce"))
                                .drop_duplicates("game_id"))
                _v2_pg = _v2_prod_cur.merge(_v2_cat_prod, on="game_id", how="left")
                _v2_pg["Product"] = _generic_product(_v2_pg["Product"].fillna("Unknown"))
                _v2_mix = _v2_pg.groupby("Product", as_index=False)["net_rev"].sum().sort_values("net_rev", ascending=False)

            pcol1, pcol2 = st.columns(2)
            with pcol1:
                _op_labels_raw = [k for k in _op_vals if _op_vals[k] > 0]
                _op_labels = [FLOOR_LABELS.get(k, k) for k in _op_labels_raw]
                _op_values = [_op_vals[k] for k in _op_labels_raw]
                if _op_values:
                    fig_op = _treemap_fig(_op_labels, _op_values)
                    st.markdown(f'<div class="card"><div class="card-title">Rest of the Property — {ceo_period} Net Revenue</div>', unsafe_allow_html=True)
                    st.plotly_chart(fig_op, use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.info("No revenue elsewhere on the property in this period.")
            with pcol2:
                if not _v2_mix.empty:
                    fig_v2 = _treemap_fig(_v2_mix["Product"], _v2_mix["net_rev"])
                    st.markdown(f'<div class="card"><div class="card-title">Main Floor Game Category Mix — {ceo_period} Net Revenue</div>', unsafe_allow_html=True)
                    st.plotly_chart(fig_v2, use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.info("No game-category data for this period.")

            # ── Top/Bottom 5 by Net + Biggest Movers — period-scoped, compact (one page) ──
            if platform == "PFH":
                _prod_pg_cur = _load_period_pfh(str(_cp_s), str(_cp_e))
                _prod_pg_pri = _load_period_pfh(str(_cp_ps), str(_cp_pe)) if ceo_period != "Lifetime" else pd.DataFrame()
            elif platform in ("V1", "V2"):
                _prod_pg_cur = _load_period_v2v1(platform, str(_cp_s), str(_cp_e))
                _prod_pg_pri = _load_period_v2v1(platform, str(_cp_ps), str(_cp_pe)) if ceo_period != "Lifetime" else pd.DataFrame()
            else:
                _prod_pg_cur = _load_period_el(str(_cp_s), str(_cp_e))
                _prod_pg_pri = _load_period_el(str(_cp_ps), str(_cp_pe)) if ceo_period != "Lifetime" else pd.DataFrame()

            _pf = (_prod_pg_cur.drop(columns=["game_name"], errors="ignore")
                   .merge(nonhr[["game_id", "game_name"]].drop_duplicates("game_id"), on="game_id", how="inner")
                   if not _prod_pg_cur.empty else pd.DataFrame())
            if not _pf.empty:
                _fleet_net_total = float(_pf["net_rev"].sum())
                _pf["share_pct"] = ((_pf["net_rev"] / _fleet_net_total * 100).clip(upper=99.9)) if _fleet_net_total > 0 else 0.0

                pm1, pm2 = st.columns(2)
                with pm1:
                    top5n = _pf.sort_values("net_rev", ascending=False).head(5)
                    fig_t5n = go.Figure(go.Bar(x=top5n["net_rev"], y=top5n["game_name"].str[:22], orientation="h",
                        marker_color=T2, text=[f"{_usd(v)} ({s:.1f}%)" for v, s in zip(top5n["net_rev"], top5n["share_pct"])],
                        textposition="outside", hovertemplate="<b>%{y}</b><br>Net: $%{x:,.0f}<extra></extra>"))
                    fig_t5n.update_layout(yaxis=dict(autorange="reversed"))
                    st.markdown(f'<div class="card"><div class="card-title">Top 5 Games — {ceo_period} Net (share of fleet)</div>', unsafe_allow_html=True)
                    st.plotly_chart(plotly_base(fig_t5n, h=220, ml=130, mr=100, mt=6, mb=28), use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)
                with pm2:
                    bot5n = _pf.sort_values("net_rev", ascending=True).head(5)
                    fig_b5n = go.Figure(go.Bar(x=bot5n["net_rev"], y=bot5n["game_name"].str[:22], orientation="h",
                        marker_color=RED, text=[f"{_usd(v)} ({s:.1f}%)" for v, s in zip(bot5n["net_rev"], bot5n["share_pct"])],
                        textposition="outside", hovertemplate="<b>%{y}</b><br>Net: $%{x:,.0f}<extra></extra>"))
                    fig_b5n.update_layout(yaxis=dict(autorange="reversed"))
                    st.markdown(f'<div class="card"><div class="card-title">Bottom 5 Games — {ceo_period} Net (share of fleet)</div>', unsafe_allow_html=True)
                    st.plotly_chart(plotly_base(fig_b5n, h=220, ml=130, mr=100, mt=6, mb=28), use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)

                if ceo_period != "Lifetime" and not _prod_pg_pri.empty:
                    _mv = _prod_pg_cur[["game_id", "net_rev"]].rename(columns={"net_rev": "net_cur"}).merge(
                        _prod_pg_pri[["game_id", "net_rev"]].rename(columns={"net_rev": "net_pri"}), on="game_id", how="left")
                    _mv["net_pri"] = _mv["net_pri"].fillna(0)
                    _mv = _mv.merge(nonhr[["game_id", "game_name"]].drop_duplicates("game_id"), on="game_id", how="inner")
                    _floor = max(float(_mv["net_pri"].median()), 50.0)  # avoid absurd % off a near-zero base
                    _mv_elig = _mv[_mv["net_pri"] >= _floor].copy()
                    _mv_elig["pct_change"] = (_mv_elig["net_cur"] - _mv_elig["net_pri"]) / _mv_elig["net_pri"] * 100

                    if not _mv_elig.empty:
                        mcol1, mcol2 = st.columns(2)
                        with mcol1:
                            top_up = _mv_elig.sort_values("pct_change", ascending=False).head(5)
                            fig_up = go.Figure(go.Bar(x=top_up["pct_change"], y=top_up["game_name"].str[:22], orientation="h",
                                marker_color=MOVE_UP, text=[f"+{v:.0f}%" for v in top_up["pct_change"]], textposition="outside",
                                hovertemplate="<b>%{y}</b><br>%{x:.1f}%<extra></extra>"))
                            fig_up.update_layout(yaxis=dict(autorange="reversed"))
                            st.markdown(f'<div class="card"><div class="card-title">Biggest Gainers — {ceo_period} vs prior</div>', unsafe_allow_html=True)
                            st.plotly_chart(plotly_base(fig_up, h=220, ml=130, mr=60, mt=6, mb=28), use_container_width=True, theme=None)
                            st.markdown('</div>', unsafe_allow_html=True)
                        with mcol2:
                            top_dn = _mv_elig.sort_values("pct_change", ascending=True).head(5)
                            fig_dn = go.Figure(go.Bar(x=top_dn["pct_change"], y=top_dn["game_name"].str[:22], orientation="h",
                                marker_color=MOVE_DOWN, text=[f"{v:.0f}%" for v in top_dn["pct_change"]], textposition="outside",
                                hovertemplate="<b>%{y}</b><br>%{x:.1f}%<extra></extra>"))
                            fig_dn.update_layout(yaxis=dict(autorange="reversed"))
                            st.markdown(f'<div class="card"><div class="card-title">Biggest Decliners — {ceo_period} vs prior</div>', unsafe_allow_html=True)
                            st.plotly_chart(plotly_base(fig_dn, h=220, ml=130, mr=60, mt=6, mb=28), use_container_width=True, theme=None)
                            st.markdown('</div>', unsafe_allow_html=True)
                elif ceo_period == "Lifetime":
                    st.caption("Biggest movers isn't shown for Lifetime — there's no comparable prior window.")

            fc1, fc2 = st.columns(2)
            sts_filt = fc1.selectbox("Filter by Status", ["All", "ON TRACK", "WATCH", "NEEDS REVIEW"], key="t2f_sts")
            wks_filt = fc2.selectbox("Filter by Weeks Live", ["All", "0–4 weeks", "4–13 weeks", "13+ weeks"], key="t2f_wks")
            filt2 = t2.copy()
            if sts_filt != "All": filt2 = filt2[filt2["Overall Status"] == sts_filt]
            if wks_filt == "0–4 weeks": filt2 = filt2[filt2["Weeks Live"] <= 4]
            elif wks_filt == "4–13 weeks": filt2 = filt2[(filt2["Weeks Live"] > 4) & (filt2["Weeks Live"] <= 13)]
            elif wks_filt == "13+ weeks": filt2 = filt2[filt2["Weeks Live"] > 13]

            # Order by severity first (worst status at top), then by lifetime net revenue —
            # so the biggest problems, and the biggest games among them, surface first.
            # Without this the table inherits catalog_from_data()'s alphabetical order and
            # the "#" column would just mean "starts with A", not "worst performer".
            _sev = {"NEEDS REVIEW": 0, "WATCH": 1, "ON TRACK": 2}
            filt2 = (filt2.assign(_sev=filt2["Overall Status"].map(_sev).fillna(9))
                          .sort_values(["_sev", "Net Revenue (lifetime)"], ascending=[True, False])
                          .drop(columns=["_sev"]))
            disp2 = filt2.drop(columns=["game_id"])[["Game Name", "Tenure", "Launch Date", "Weeks Live",
                "Hold % (latest)", "Bet Decay (latest)", "Net Revenue (lifetime)", "Share of Fleet %",
                "Overall Status"]].reset_index(drop=True)
            disp2.insert(0, "#", range(1, len(disp2) + 1))

            def _sts_style(val):
                bg = {"ON TRACK": GREEN_LT, "WATCH": AMBER_LT, "NEEDS REVIEW": RED_LT}.get(val, "")
                fg = {"ON TRACK": GREEN, "WATCH": AMBER, "NEEDS REVIEW": RED}.get(val, TEXT)
                return f"background-color:{bg};color:{fg};font-weight:600"
            st.dataframe(
                disp2.style.map(_sts_style, subset=["Overall Status"]).format({
                    "Hold % (latest)": "{:.1f}%", "Bet Decay (latest)": "{:.1f}%", "Net Revenue (lifetime)": "${:,.0f}",
                    "Share of Fleet %": "{:.2f}%",
                }, na_rep="–"),
                use_container_width=True, hide_index=True, height=600,
            )
        else:
            st.info("No game data available.")
    else:
        if not peer_catalog.empty:
            peer_pool = L.find_peer_pool(peer_catalog, df, sel_id, scale_tolerance=scale_x)
            peer_df = df[df["game_id"].isin(peer_pool["ids"])]
        else:
            peer_pool = {"family": "whole fleet (no catalog)", "pool_size": int(df["game_id"].nunique()), "scale_applied": False}
            peer_df = df

        _rc1, _rc2 = st.columns([7, 5])
        with _rc1:
            ribbon("Is It On Track?",
                   f"{top_title} · Week 0–{int(gdf['launch_week'].max())} · vs. {peer_pool['pool_size']} peers "
                   f"({peer_pool['family']}{', scale-matched' if peer_pool['scale_applied'] else ''})", T2, "STATUS")
        with _rc2:
            st.markdown('<div class="eyebrow" style="margin:10px 0 2px">SINCE-LAUNCH WINDOW</div>', unsafe_allow_html=True)
            _window_choice = st.radio(
                "Since-launch window", ["Full History", "First 30 Days", "First 60 Days", "First 90 Days"],
                horizontal=True, key="track_window", label_visibility="collapsed")
        max_wk_t = int(gdf["launch_week"].max())
        weeks_live = max_wk_t + 1

        # ══════════════════════ compute everything first, render as one compact block ═══
        exec_kpis = [("bet_handle", True), ("hold_pct", True), ("bet_decay", False), ("net_rev", True)]
        exec_flags = {}
        for kpi_e, hi_e in exec_kpis:
            bnd_e = bkpi(peer_df, kpi_e)
            brow_e = bnd_e[bnd_e["launch_week"] == max_wk_t] if not bnd_e.empty else pd.DataFrame()
            val_e = gdf[gdf["launch_week"] == max_wk_t][kpi_e].values
            if len(val_e) and not brow_e.empty and pd.notna(val_e[0]):
                exec_flags[kpi_e] = _flag_cls(float(val_e[0]), float(brow_e["p25"].iloc[0]), float(brow_e["p75"].iloc[0]),
                                               float(brow_e["p10"].iloc[0]), float(brow_e["p90"].iloc[0]), higher_is_better=hi_e)
        n_r_e = sum(1 for v in exec_flags.values() if v == "r")
        n_a_e = sum(1 for v in exec_flags.values() if v == "a")
        eh = "NEEDS REVIEW" if n_r_e >= 2 else ("WATCH" if (n_r_e >= 1 or n_a_e >= 2) else "ON TRACK")
        # Guardrail: lifetime RTP > 100% (win > bet) is real, confirmed money lost -- never
        # let percentile-band noise on unrelated metrics soften that to WATCH/ON TRACK.
        if float(gdf["net_rev"].sum()) < 0:
            eh = "NEEDS REVIEW"
        eh_kind = {"ON TRACK": "g", "WATCH": "a", "NEEDS REVIEW": "r"}[eh]

        # NOT applied to EdgeLabs — see the report generator's identical guard above for why
        # (launch_date there = first real-money spin; no separate "soft launch" phase exists).
        ramp = L.detect_ramp(df, sel_id) if platform != "EdgeLabs" else {"is_ramping": False, "note": ""}
        w0_stores = None
        if "stores" in gdf.columns:
            _w0_rows = gdf[gdf["launch_week"] == 0]["stores"].dropna()
            if not _w0_rows.empty:
                w0_stores = int(_w0_rows.iloc[0])
        # Week-0 player count — a soft-launch/test period on Pong (PFH's online identity)
        # shows up as a handful of test accounts before the real launch weeks later.
        # NOT applied to EdgeLabs: there, launch = first spin by definition, so a thin
        # week-0 player count isn't a "soft launch" signal, just early ramp-up.
        w0_players = None
        if platform != "EdgeLabs" and "players" in gdf.columns:
            _w0_prows = gdf[gdf["launch_week"] == 0]["players"].dropna()
            if not _w0_prows.empty:
                w0_players = int(_w0_prows.iloc[0])
        _soft_launch_bits = []
        if w0_stores is not None:
            _soft_launch_bits.append(f"{w0_stores} location{'s' if w0_stores != 1 else ''}")
        if w0_players is not None:
            _soft_launch_bits.append(f"{w0_players} player{'s' if w0_players != 1 else ''}")
        _soft_launch_val = f"{meta['launch_date']}" + (f" ({', '.join(_soft_launch_bits)})" if _soft_launch_bits else "")

        # best_match_score/ms_data params are accepted but never referenced inside
        # compute_quick_score() (confirmed by reading launch.py) — passed as inert defaults.
        qs_player_df = pd.DataFrame()
        if platform == "EdgeLabs":
            try:
                qs_player_df = load_edgelabs_player_weeks(int(sel_id), str(meta["launch_date"]), "EdgeLabs")
            except Exception:
                pass
        elif "players" in gdf.columns and gdf["players"].notna().any():
            qs_player_df = gdf[["launch_week", "players"]].dropna(subset=["players"]).rename(columns={"players": "unique_players"})
        qs = L.compute_quick_score(gdf, {}, peer_df, 0.0, player_df=qs_player_df if not qs_player_df.empty else None)

        cur_row = gdf[gdf["launch_week"] == max_wk_t]
        cv_bet = float(cur_row["bet_handle"].iloc[0] or 0) if not cur_row.empty else 0.0
        cv_net = float(cur_row["net_rev"].iloc[0] or 0) if not cur_row.empty else 0.0
        cv_hold = float(cur_row["hold_pct"].iloc[0]) if not cur_row.empty and pd.notna(cur_row["hold_pct"].iloc[0]) else None

        game_lifetime_net = float(gdf["net_rev"].sum())
        plat_total_net = float(cat["total_net"].sum()) if not cat.empty else 0.0
        share_lifetime = (game_lifetime_net / plat_total_net * 100) if plat_total_net > 0 else None
        rank_list = cat.sort_values("total_net", ascending=False)["game_id"].tolist() if not cat.empty else []
        rank_n = (rank_list.index(sel_id) + 1) if sel_id in rank_list else None

        _r_end, _r_start = dt.date.today(), dt.date.today() - dt.timedelta(weeks=8)
        if platform == "PFH":
            _recent_pg = _load_period_pfh(str(_r_start), str(_r_end))
        elif platform in ("V1", "V2"):
            _recent_pg = _load_period_v2v1(platform, str(_r_start), str(_r_end))
        else:
            _recent_pg = _load_period_el(str(_r_start), str(_r_end))
        recent_total = float(_recent_pg["net_rev"].sum()) if not _recent_pg.empty else 0.0
        recent_game = float(_recent_pg[_recent_pg["game_id"] == sel_id]["net_rev"].sum()) if not _recent_pg.empty else 0.0
        share_recent = (recent_game / recent_total * 100) if recent_total > 0 else None
        losing_ground = share_recent is not None and share_lifetime is not None and share_recent < share_lifetime

        tenure = _tenure_band(weeks_live)
        peak_net = float(gdf["net_rev"].max()) if gdf["net_rev"].notna().any() else 0.0
        last4_net = float(gdf[gdf["launch_week"] > max_wk_t - 4]["net_rev"].mean()) if len(gdf) else 0.0
        pct_of_peak = (last4_net / peak_net * 100) if peak_net > 0 else None
        roll = gdf.sort_values("launch_week")["net_rev"].rolling(4, min_periods=4).mean().dropna()
        if weeks_live < 13 or len(roll) < 4:
            trend_status, trend_kind = "Insufficient data", "n"
            trend_note = (f"Only {weeks_live} weeks live — % of peak and trend status aren't meaningful yet "
                          "(a Recent-tenure game is at its own peak by definition; that's arithmetic, not a good sign).")
        else:
            declines = int((roll.diff() < 0).sum())
            if (pct_of_peak is not None and pct_of_peak < 20) or declines >= len(roll) - 1:
                trend_status, trend_kind = "Declining", "r"
            elif roll.iloc[-1] >= roll.iloc[0]:
                trend_status, trend_kind = "Growing", "g"
            else:
                trend_status, trend_kind = "Stable", "a"
            trend_note = f"{pct_of_peak:.0f}% of its own peak week, based on the trailing 4 weeks." if pct_of_peak is not None else ""

        bh_series = gdf.sort_values("launch_week")["bet_handle"]
        peak_wk = int(gdf.loc[gdf["bet_handle"].idxmax(), "launch_week"]) if gdf["bet_handle"].notna().any() else 0
        w0_bh = float(bh_series.iloc[0]) if len(bh_series) else 0.0
        # "Flash-in-the-pan" means it crashed FAST, not just "eventually ended up low" — a title
        # that peaked at launch and gradually faded over a full year is a long tail, not a flash.
        # Check the crash at an early checkpoint (~wk 12) rather than the current/latest week,
        # so a 50-week-old game that only recently dipped below 30% doesn't get mislabeled.
        _early_wk = min(12, max_wk_t)
        _early_row = gdf[gdf["launch_week"] == _early_wk]
        _early_bh = float(_early_row["bet_handle"].iloc[0]) if not _early_row.empty and pd.notna(_early_row["bet_handle"].iloc[0]) else None
        _early_pct = (_early_bh / w0_bh * 100) if (_early_bh is not None and w0_bh > 0) else None
        if weeks_live < 8:
            archetype = "Still forming"
        elif pct_of_peak is not None and pct_of_peak >= 60:
            archetype = "Evergreen"
        elif peak_wk <= 1 and _early_pct is not None and _early_pct < 30:
            archetype = "Flash-in-the-pan"
        elif peak_wk >= 3:
            archetype = "Slow-burn"
        else:
            archetype = "Steady decliner"
        _archetype_defs = {
            "Still forming": "Too early to classify (<8 weeks live)",
            "Evergreen": f"Still {pct_of_peak:.0f}% of its own peak" if pct_of_peak is not None else "Still near its own peak",
            "Flash-in-the-pan": "Peaked at launch, crashed within ~12 weeks",
            "Slow-burn": "Built up over several weeks before peaking",
            "Steady decliner": "Peaked early, faded gradually over time",
        }

        foot = load_game_footprint(platform, sel_id)
        if foot.empty:
            n_locs, top_share, conc_kind, conc_label = 0, None, "n", "–"
        else:
            n_locs = int(foot["loc_id"].nunique())
            foot_sorted = foot.sort_values("net_rev", ascending=False)
            top_loc_net = float(foot_sorted["net_rev"].iloc[0]) if len(foot_sorted) else 0.0
            total_loc_net = float(foot["net_rev"].sum())
            top_share = (top_loc_net / total_loc_net * 100) if total_loc_net > 0 else None
            conc_kind = "r" if (top_share is not None and top_share >= 40) else ("a" if (top_share is not None and top_share >= 25) else "g")
            conc_label = "Concentrated" if conc_kind == "r" else ("Watch" if conc_kind == "a" else "Spread out")

        gdf_s = gdf.sort_values("launch_week")
        spd = gdf_s["spins"] / gdf_s["stores"].clip(lower=1) / 7
        latest_hold = float(gdf_s["hold_pct"].dropna().iloc[-1]) if gdf_s["hold_pct"].notna().any() else None
        latest_decay = float(gdf_s["bet_decay"].dropna().iloc[-1]) if gdf_s["bet_decay"].notna().any() else None

        # First 30 / 60 / 90 days — cumulative through the nearest launch week. Kept to just
        # Net Revenue + Hold % (dropped Bet Handle, redundant with these two) so the table
        # actually fits in a half-width column instead of truncating every header/cell.
        _mstone_rows = []
        for _label, _days in (("First 30 Days", 30), ("First 60 Days", 60), ("First 90 Days", 90)):
            _wk_cut = _days // 7
            if weeks_live - 1 < _wk_cut:
                _mstone_rows.append({"Milestone": _label, "Days": _days, "Net Revenue": np.nan, "Hold %": np.nan})
                continue
            _sub = gdf[gdf["launch_week"] <= _wk_cut]
            _bet, _net = float(_sub["bet_handle"].sum()), float(_sub["net_rev"].sum())
            _hold = (_net / _bet * 100) if _bet > 0 else np.nan
            _mstone_rows.append({"Milestone": _label, "Days": _days, "Net Revenue": _net,
                                  "Hold %": _hold})
        mstone_df = pd.DataFrame(_mstone_rows)

        try:
            _analog = find_peers_scaled(df, sel_id, kpi="bet_decay", n_weeks=n_match, top_k=1, scale_tolerance=scale_x)
        except Exception:
            _analog = []

        # ══════════════════════ render — everything compact, at the top ════════════════
        st.markdown(f'<div style="margin:4px 0 10px">{badge(eh, eh_kind)}</div>', unsafe_allow_html=True)
        if ramp["is_ramping"]:
            st.markdown(f'<div class="ann ann-a"><span class="ann-tag">WATCH</span>'
                        f'<div class="ann-body"><strong>Soft launch detected.</strong> {ramp["note"]}</div></div>',
                        unsafe_allow_html=True)

        # Quick Score + Trend + Archetype on one bordered row — .card has no CSS rule of its
        # own in this file (only used elsewhere to wrap a plotly chart, which supplies its own
        # border via [data-testid="stPlotlyChart"]) so these text-only cards rendered borderless.
        _card_style = f"background:{SURFACE};border:1px solid {BORDER};border-radius:10px;padding:14px 16px"
        _qs_kind_map = {"green": "g", "yellow": "a", "red": "r"}
        qs_meta = [
            ("Game Performance", badge(qs["game_performance"]["label"], _qs_kind_map.get(qs["game_performance"]["color"], "n")), qs["game_performance"]["reason"]),
            ("Impact on Overall Sales", badge(qs["sales_impact"]["label"], _qs_kind_map.get(qs["sales_impact"]["color"], "n")), qs["sales_impact"]["reason"]),
            ("Held Player's Interest", badge(qs["player_interest"]["label"], _qs_kind_map.get(qs["player_interest"]["color"], "n")), qs["player_interest"]["reason"]),
            ("Trend", badge(trend_status, trend_kind), trend_note),
            ("Archetype", archetype, _archetype_defs.get(archetype, "")),
        ]
        qcols = st.columns(5)
        for col, (title, value_html, reason) in zip(qcols, qs_meta):
            col.markdown(
                f'<div style="{_card_style}"><div class="eyebrow">{title}</div>'
                f'<div style="margin:6px 0">{value_html}</div>'
                f'<div style="font-size:12px;color:{T2};line-height:1.4">{reason}</div></div>',
                unsafe_allow_html=True)

        krow([
            {"label": f"Week {max_wk_t} Bet", "value": _usd(cv_bet)},
            {"label": f"Week {max_wk_t} Net Rev", "value": _usd(cv_net)},
            {"label": "Hold %", "value": _pct(cv_hold) if cv_hold is not None else "–"},
            {"label": "Lifetime Game Net", "value": _usd(game_lifetime_net), "sub": "house hold, not Pong royalty"},
            {"label": "Tenure", "value": tenure, "sub": f"{weeks_live} weeks live"},
            {"label": "Locations", "value": f"{n_locs:,}" if n_locs else "–"},
            {"label": "Top Location Share", "value": f'{_pct(top_share) if top_share is not None else "–"}'
                                                       + (f" {badge(conc_label, conc_kind)}" if top_share is not None else "")},
            {"label": "Soft Launch Date", "value": _soft_launch_val},
        ])

        _mid_caps = []
        if weeks_live < 13:
            _mid_caps.append("Too early (Recent tenure) to split a math problem from an appeal problem.")
        elif latest_hold is not None and latest_hold < 6 and (latest_decay is None or latest_decay > 50):
            _mid_caps.append("Thin hold (<6%) with engagement still reasonably strong — looks like a math/config issue, not a demand problem.")
        elif latest_decay is not None and latest_decay < 30:
            _mid_caps.append("Engagement has faded sharply — looks like an appeal problem, not a math one.")
        if _mid_caps:
            st.caption(" · ".join(_mid_caps))

        # ── Since-launch window filter — widget itself now lives up top beside the ribbon
        # (rendered via _rc2 above); _window_choice was already captured from that widget.
        # Every chart and the milestone table below re-slice to it; nothing else reads it.
        _window_days = {"Full History": None, "First 30 Days": 30, "First 60 Days": 60, "First 90 Days": 90}[_window_choice]
        _window_wk_cut = (_window_days // 7) if _window_days is not None else None
        _window_suffix = "" if _window_days is None else f" — {_window_choice}"
        gdf_win = gdf if _window_wk_cut is None else gdf[gdf["launch_week"] <= _window_wk_cut]
        gdf_s_win = gdf_s if _window_wk_cut is None else gdf_s[gdf_s["launch_week"] <= _window_wk_cut]
        spd_win = spd if _window_wk_cut is None else spd[gdf_s["launch_week"] <= _window_wk_cut]
        mstone_df_win = mstone_df if _window_days is None else mstone_df[mstone_df["Days"] <= _window_days]
        if gdf_win.empty:
            st.caption(f"{top_title} hasn't reached {_window_choice.lower()} of data yet — showing full history instead.")
            gdf_win, gdf_s_win, spd_win, mstone_df_win = gdf, gdf_s, spd, mstone_df
            _window_suffix = ""

        # Shared "vs. peer range" plumbing — reused across Bet Handle, Net Revenue, and RTP %
        # below, so the peer comparison is visible on every trend chart, not just one.
        wks_x = sorted(gdf_win["launch_week"].dropna().unique().astype(int).tolist())

        def _series_at(col):
            return [float(gdf_win[gdf_win["launch_week"] == w][col].iloc[0])
                    if not gdf_win[gdf_win["launch_week"] == w].empty and pd.notna(gdf_win[gdf_win["launch_week"] == w][col].iloc[0]) else None
                    for w in wks_x]

        def _band_arr(band_df, col):
            out = []
            for w in wks_x:
                br = band_df[band_df["launch_week"] == w] if not band_df.empty else pd.DataFrame()
                out.append(float(br[col].iloc[0]) if not br.empty and col in br.columns and pd.notna(br[col].iloc[0]) else None)
            return out

        # ── Row: Bet Handle (line+band) | Net Revenue (bars+band) — bumped size ────
        ch1, ch2 = st.columns(2)
        with ch1:
            band_bh = bkpi(peer_df, "bet_handle")
            fig_bh = go.Figure()
            fill_band(fig_bh, wks_x, _band_arr(band_bh, "p10"), _band_arr(band_bh, "p90"), "rgba(91,89,78,0.06)", "P10–P90 peer range")
            fill_band(fig_bh, wks_x, _band_arr(band_bh, "p25"), _band_arr(band_bh, "p75"), "rgba(91,89,78,0.12)", "P25–P75 peer range")
            fig_bh.add_trace(go.Scatter(x=wks_x, y=_band_arr(band_bh, "p50"), mode="lines", line=dict(color=T3, dash="dot", width=1.5), name="Peer median"))
            fig_bh.add_trace(go.Scatter(x=wks_x, y=_series_at("bet_handle"), mode="lines+markers", line=dict(color=MOVE_DOWN, width=2.5),
                                         marker=dict(size=5), name=top_title[:22]))
            fig_bh.update_xaxes(title_text="Launch week", automargin=True)
            fig_bh.update_yaxes(title_text="Bet handle ($), weekly total", automargin=True)
            st.markdown(f'<div class="card"><div class="card-title">Bet Handle vs. Peer Range{_window_suffix}</div>'
                        f'<div style="font-size:12px;color:{T2};margin:-4px 0 6px">Weekly total bet handle vs. the {peer_pool["pool_size"]}-game peer band ({peer_pool["family"]})</div>', unsafe_allow_html=True)
            st.plotly_chart(plotly_base(fig_bh, h=420), use_container_width=True, theme=None)
            st.markdown('</div>', unsafe_allow_html=True)
        with ch2:
            band_net = bkpi(peer_df, "net_rev")
            fig_nr = go.Figure()
            fill_band(fig_nr, wks_x, _band_arr(band_net, "p10"), _band_arr(band_net, "p90"), "rgba(122,106,156,0.08)", "P10–P90 peer range")
            fill_band(fig_nr, wks_x, _band_arr(band_net, "p25"), _band_arr(band_net, "p75"), "rgba(122,106,156,0.18)", "P25–P75 peer range")
            fig_nr.add_trace(go.Scatter(x=wks_x, y=_band_arr(band_net, "p50"), mode="lines", line=dict(color=T3, dash="dot", width=1.5), name="Peer median"))
            fig_nr.add_trace(go.Bar(x=wks_x, y=_series_at("net_rev"), marker_color="#7A6A9C", name=top_title[:22]))
            fig_nr.update_xaxes(title_text="Launch week", automargin=True)
            fig_nr.update_yaxes(title_text="Net revenue ($), weekly total", automargin=True)
            st.markdown(f'<div class="card"><div class="card-title">Net Revenue vs. Peer Range{_window_suffix}</div>'
                        f'<div style="font-size:12px;color:{T2};margin:-4px 0 6px">Weekly Game Net (house hold) vs. the same peer band</div>', unsafe_allow_html=True)
            st.plotly_chart(plotly_base(fig_nr, h=420), use_container_width=True, theme=None)
            st.markdown('</div>', unsafe_allow_html=True)

        # ── Row: Hold % (line+band) | Engagement — clarified as store-level, not player-level ──
        pc1, pc2 = st.columns(2)
        with pc1:
            band_hold = bkpi(peer_df, "hold_pct")
            fig_rtp_t = go.Figure()
            fill_band(fig_rtp_t, wks_x, _band_arr(band_hold, "p10"), _band_arr(band_hold, "p90"),
                      "rgba(91,89,78,0.08)", "P10–P90 peer range")
            fig_rtp_t.add_trace(go.Scatter(x=wks_x, y=_band_arr(band_hold, "p50"),
                                            mode="lines", line=dict(color=T3, dash="dot", width=1.5), name="Peer median"))
            fig_rtp_t.add_trace(go.Scatter(x=wks_x, y=_series_at("hold_pct"),
                                            mode="lines+markers", line=dict(color=T2, width=2.5), marker=dict(size=5), name="Hold %"))
            fig_rtp_t.update_yaxes(title_text="Hold %", automargin=True)
            fig_rtp_t.update_xaxes(title_text="Launch week", automargin=True)
            st.markdown(f'<div class="card"><div class="card-title">Hold % vs. Peer Range{_window_suffix}</div>'
                        f'<div style="font-size:12px;color:{T2};margin:-4px 0 6px">Weekly Hold % (house share of wagers) vs. the same peer band</div>', unsafe_allow_html=True)
            st.plotly_chart(plotly_base(fig_rtp_t, h=300), use_container_width=True, theme=None)
            st.markdown('</div>', unsafe_allow_html=True)
        with pc2:
            fig_spd = go.Figure(go.Scatter(x=gdf_s_win["launch_week"], y=spd_win, mode="lines+markers", line=dict(color=AMBER, width=2)))
            fig_spd.update_yaxes(title_text="Spins / store / day", automargin=True)
            fig_spd.update_xaxes(title_text="Launch week", automargin=True)
            st.markdown(f'<div class="card"><div class="card-title">Engagement — Spins/Store/Day{_window_suffix}</div>'
                        f'<div style="font-size:12px;color:{T2};margin:-4px 0 6px">Store/terminal-level activity rate — not a player-level metric</div>', unsafe_allow_html=True)
            st.plotly_chart(plotly_base(fig_spd, h=300), use_container_width=True, theme=None)
            st.markdown('</div>', unsafe_allow_html=True)

        # ── Row: Avg Bet/Spin | First 30/60/90 Days ─────────────────────────────────
        ac1, ac2 = st.columns(2)
        with ac1:
            fig_ab = go.Figure(go.Scatter(x=gdf_s_win["launch_week"], y=gdf_s_win["avg_bet"], mode="lines+markers", line=dict(color=MOVE_UP, width=2)))
            fig_ab.update_yaxes(title_text="Avg bet / spin ($)", automargin=True)
            fig_ab.update_xaxes(title_text="Launch week", automargin=True)
            st.markdown(f'<div class="card"><div class="card-title">Avg Bet / Spin{_window_suffix}</div>'
                        f'<div style="font-size:12px;color:{T2};margin:-4px 0 6px">Average wager size per spin, weekly</div>', unsafe_allow_html=True)
            st.plotly_chart(plotly_base(fig_ab, h=280), use_container_width=True, theme=None)
            st.markdown('</div>', unsafe_allow_html=True)
        with ac2:
            # Plain HTML rows instead of st.dataframe — Streamlit's canvas grid doesn't
            # respect use_container_width for a table this small (confirmed live: rendered
            # at 63px wide inside a 364px column), so it truncated every header and cell.
            _mstone_html = []
            for _, r in mstone_df_win.iterrows():
                _net_str = _usd(r["Net Revenue"]) if pd.notna(r["Net Revenue"]) else "Not reached yet"
                _hold_str = f"{r['Hold %']:.1f}% Hold" if pd.notna(r["Hold %"]) else "–"
                _mstone_html.append(
                    f'<div style="display:flex;justify-content:space-between;align-items:baseline;'
                    f'padding:9px 2px;border-bottom:1px solid {BORDER}">'
                    f'<span style="font-weight:600">{r["Milestone"]}</span>'
                    f'<span>{_net_str}</span><span style="color:{T2}">{_hold_str}</span></div>')
            st.markdown(f'<div style="{_card_style}"><div class="card-title">First 30 / 60 / 90 Days — cumulative Net Revenue & Hold</div>'
                        + "".join(_mstone_html) + '</div>', unsafe_allow_html=True)

        _footer_bits = ["Efficiency (dev-hours, payback) intentionally omitted — no Jira/worklog data connected; a payback "
                        "figure off Game Net instead of real royalty would read ~7× too optimistic."]
        if _analog:
            _ap = _analog[0]
            _sim = _similarity_pct(_ap["distance"], _ap.get("n_match_weeks"))
            _footer_bits.append(f"Closest analog: **{_ap['game_name']}** ({_sim:.0f}% match, {_ap.get('peer_family', 'unknown')} tier).")
        st.caption(" · ".join(_footer_bits))

# ══════════════════════════════════════════════════════════════════
# TAB — SIMILAR LAUNCHES
# ══════════════════════════════════════════════════════════════════
with tab_similar:
    # Double-click ANY chart anywhere in the dashboard to pop it out full-size; a "✕" appears
    # top-right to close it back to normal. Also wires any element carrying the .wg-zoomable
    # class (plain-HTML panels that aren't a Plotly figure, e.g. the Weekly Games top-N
    # tables) to the same overlay via a cloned/scaled-up copy instead of a Plotly re-render.
    # Injected once here, but scans the whole document (not just this tab) — Streamlit keeps
    # every tab's panel mounted in the DOM simultaneously, so a document-wide scan reaches
    # every chart/panel on every tab, not only this page's.
    st.components.v1.html("""
    <script>
    (function() {
      const PDOC = window.parent.document;
      const PWIN = window.parent;

      // Traces with no x/y axes (pie/donut, sunburst, treemap, funnelarea) have no
      // zoom/autoscale gesture for Plotly to own, so `plotly_doubleclick` never fires for
      // them — confirmed live (a real double-click on a donut chart did nothing). Those get
      // a plain DOM 'dblclick' listener instead; everything else uses Plotly's own event.
      const NON_CARTESIAN_TYPES = ['pie', 'sunburst', 'treemap', 'funnelarea', 'indicator'];

      function wireChart(wrapper) {
        if (wrapper.dataset.dblclickWired) return;
        const gd = wrapper.querySelector('.js-plotly-plot');
        if (!gd || typeof gd.on !== 'function' || !gd.data) return;
        wrapper.dataset.dblclickWired = "1";
        wrapper.style.cursor = "zoom-in";
        wrapper.title = "Double-click to enlarge";

        function openOverlay() {
          if (!PWIN.Plotly) return;
          const overlay = PDOC.createElement('div');
          overlay.style.cssText = 'position:fixed;top:4vh;left:4vw;width:92vw;height:92vh;z-index:999999;' +
            'background:white;box-shadow:0 10px 50px rgba(0,0,0,0.45);border-radius:10px;padding:20px;';
          const closeBtn = PDOC.createElement('div');
          closeBtn.textContent = '\\u2715';
          closeBtn.title = 'Close';
          closeBtn.style.cssText = 'position:absolute;top:14px;right:20px;font-size:26px;cursor:pointer;' +
            'color:#333;font-weight:bold;z-index:1000000;line-height:1;';
          function closeOverlay() {
            overlay.remove();
            PDOC.removeEventListener('keydown', escHandler);
          }
          function escHandler(ev) { if (ev.key === 'Escape') closeOverlay(); }
          closeBtn.onclick = closeOverlay;
          PDOC.addEventListener('keydown', escHandler);
          overlay.appendChild(closeBtn);
          const plotHost = PDOC.createElement('div');
          plotHost.style.cssText = 'width:100%;height:100%;';
          overlay.appendChild(plotHost);
          PDOC.body.appendChild(overlay);
          const clonedLayout = JSON.parse(JSON.stringify(gd.layout));
          delete clonedLayout.width; delete clonedLayout.height;
          clonedLayout.autosize = true;
          // Preserve any explicit axis range set by the original chart (e.g. the Trajectory
          // overlay's 100-week cap) — autosize:true + responsive:true otherwise lets Plotly
          // recompute autorange on the fresh newPlot call and the cap gets lost in the popout.
          Object.keys(clonedLayout).forEach(function(k) {
            if (/^[xy]axis/.test(k) && clonedLayout[k] && Array.isArray(clonedLayout[k].range)) {
              clonedLayout[k].autorange = false;
            }
          });
          PWIN.Plotly.newPlot(plotHost, JSON.parse(JSON.stringify(gd.data)), clonedLayout,
                               {responsive: true, displayModeBar: true});
        }

        const isNonCartesian = gd.data.length > 0 && gd.data.every(function(t) {
          return NON_CARTESIAN_TYPES.indexOf(t.type) !== -1;
        });
        if (isNonCartesian) {
          wrapper.addEventListener('dblclick', openOverlay);
        } else {
          gd.on('plotly_doubleclick', openOverlay);
        }
      }

      // Plain HTML panels (tables, KPI-style cards built with raw markdown, not a Plotly
      // figure) can't fire plotly_doubleclick — they get their own dblclick listener that
      // pops a scaled-up clone into the same overlay style, so the interaction feels the
      // same as the chart zoom above even though there's no Plotly figure to re-render.
      function wireCard(el) {
        if (el.dataset.dblclickWired) return;
        el.dataset.dblclickWired = "1";
        el.style.cursor = "zoom-in";
        if (!el.title) el.title = "Double-click to enlarge";
        el.addEventListener('dblclick', function() {
          const overlay = PDOC.createElement('div');
          overlay.style.cssText = 'position:fixed;top:4vh;left:4vw;width:92vw;height:92vh;z-index:999999;' +
            'background:white;box-shadow:0 10px 50px rgba(0,0,0,0.45);border-radius:10px;padding:24px;overflow:auto;';
          const closeBtn = PDOC.createElement('div');
          closeBtn.textContent = '✕';
          closeBtn.title = 'Close';
          closeBtn.style.cssText = 'position:absolute;top:14px;right:20px;font-size:26px;cursor:pointer;' +
            'color:#333;font-weight:bold;z-index:1000000;line-height:1;';
          function closeOverlay() {
            overlay.remove();
            PDOC.removeEventListener('keydown', escHandler);
          }
          function escHandler(ev) { if (ev.key === 'Escape') closeOverlay(); }
          closeBtn.onclick = closeOverlay;
          PDOC.addEventListener('keydown', escHandler);
          overlay.appendChild(closeBtn);
          const clone = el.cloneNode(true);
          clone.style.cursor = 'default';
          clone.style.transformOrigin = 'top left';
          clone.style.transform = 'scale(1.5)';
          clone.style.marginTop = '20px';
          clone.removeAttribute('data-dblclick-wired');
          overlay.appendChild(clone);
          PDOC.body.appendChild(overlay);
        });
      }

      function scan() {
        PDOC.querySelectorAll('[data-testid="stPlotlyChart"]').forEach(wireChart);
        PDOC.querySelectorAll('.wg-zoomable').forEach(wireCard);
      }
      scan();
      const obs = new PWIN.MutationObserver(scan);
      obs.observe(PDOC.body, {childList: true, subtree: true});
    })();
    </script>
    """, height=0)

    render_generate_report_button("similar")
    if is_all:
        ribbon("Player Retention by Game — Week by Week",
               "Green = players kept wagering well · Red = sharp drop-off · Blank = not yet reached", T2, "RETENTION")
        st.caption("Each cell shows bet decay % at that launch week (week 0 = 100%). Sort order: fastest-fading games at top.")

        hm_df = df[~df["game_id"].apply(L._is_hr)].copy()
        max_wk_hm = min(13, int(hm_df["launch_week"].max())) if not hm_df.empty else 0
        hm_df = hm_df[hm_df["launch_week"] <= max_wk_hm]
        nm_map = hm_df.drop_duplicates("game_id").set_index("game_id")["game_name"].to_dict()
        pivot = hm_df.pivot_table(index="game_id", columns="launch_week", values="bet_decay", aggfunc="first")
        pivot.index = [str(nm_map.get(g, g))[:32] for g in pivot.index]
        last_col = pivot.columns[-1] if len(pivot.columns) else 0
        pivot = pivot.sort_values(last_col, ascending=True, na_position="last")

        if not pivot.empty:
            fig_hm = go.Figure(go.Heatmap(
                z=pivot.values, x=[f"Wk {c}" for c in pivot.columns], y=pivot.index.tolist(),
                colorscale=[[0, MOVE_DOWN], [0.5, "#E8DCC8"], [1, MOVE_UP]], zmid=50, zmin=0, zmax=100,
                colorbar=dict(title="Bet Decay %", thickness=14, len=0.7),
                hovertemplate="<b>%{y}</b><br>%{x}<br>Decay: %{z:.1f}%<extra></extra>"))
            plotly_base(fig_hm, h=max(560, len(pivot) * 30), ml=180, mr=50, mt=20, mb=40)
            st.plotly_chart(fig_hm, use_container_width=True, theme=None)
        else:
            st.info("Not enough data to build heatmap.")

        ribbon("What to Expect — All Active Games, Next 4 Weeks",
               f"{len(nonhr)} {platform} games, projected from each game's recent wagering trend", T2, "FORECAST")

        fcst_rows = []
        for _, gc4 in nonhr.iterrows():
            gid4 = gc4["game_id"]
            g4 = df[df["game_id"] == gid4].sort_values("launch_week")
            if len(g4) < 2:
                continue
            last4_bh = g4.tail(4)["bet_handle"].values
            if len(last4_bh) >= 2 and float(last4_bh[0]) > 0:
                dr = (float(last4_bh[-1]) / float(last4_bh[0])) ** (1.0 / max(len(last4_bh) - 1, 1))
                dr = max(0.5, min(1.5, dr))
            else:
                dr = 1.0
            mwk4, lb4 = int(g4["launch_week"].max()), float(g4.iloc[-1]["bet_handle"])
            for fw in range(1, 5):
                fcst_rows.append({"game_id": gid4, "game_name": gc4["game_name"], "proj_week": mwk4 + fw,
                                  "projected_bet": lb4 * (dr ** fw)})

        if fcst_rows:
            t4 = pd.DataFrame(fcst_rows)
            n_act4 = int(t4["game_id"].nunique())
            nxt_wk4 = float(t4[t4["proj_week"] == t4["proj_week"].min()]["projected_bet"].sum())
            tot4w = float(t4["projected_bet"].sum())

            krow([
                {"label": "Active Games in Forecast", "value": str(n_act4), "sub": "with ≥2 weeks data"},
                {"label": "Projected Next Week", "value": _usd(nxt_wk4), "sub": "combined fleet bet handle"},
                {"label": "4-Week Combined Total", "value": _usd(tot4w)},
                {"label": "Avg per Game (Next Wk)", "value": _usd(nxt_wk4 / n_act4) if n_act4 else "–"},
            ])

            t4_gids = t4["game_id"].unique()
            t4_gnames = t4.drop_duplicates("game_id").set_index("game_id")["game_name"].to_dict()
            fig_stack = go.Figure()
            for gid4 in t4_gids:
                gf = t4[t4["game_id"] == gid4].sort_values("proj_week")
                fig_stack.add_trace(go.Bar(x=[f"Wk {w}" for w in gf["proj_week"]], y=gf["projected_bet"],
                                            name=str(t4_gnames.get(gid4, gid4))[:24],
                                            hovertemplate="<b>%{name}</b><br>%{x}<br>Projected: $%{y:,.0f}<extra></extra>"))
            fig_stack.update_layout(barmode="stack")
            fig_stack.update_yaxes(title_text="Projected bet handle ($)")
            fig_stack.update_xaxes(title_text="Forecast week")
            st.markdown('<div class="card"><div class="card-title">Projected Bet Handle by Game — Next 4 Weeks (Stacked)</div>', unsafe_allow_html=True)
            st.plotly_chart(plotly_base(fig_stack, h=530), use_container_width=True, theme=None)
            st.markdown('</div>', unsafe_allow_html=True)
            st.caption("Projection uses each game's recent 4-week trend — simple exponential smoothing, not a statistical model.")

            ribbon("Games That Need Attention", "Projected to fall into the bottom quarter of performers within 4 weeks", T2, "RISK")
            bnd_p25_t4 = bkpi(df, "bet_handle")
            risk_rows4 = []
            for gid4 in t4_gids:
                gf4 = t4[t4["game_id"] == gid4].sort_values("proj_week")
                ga4 = df[df["game_id"] == gid4].sort_values("launch_week")
                if ga4.empty:
                    continue
                cb4, mwk4 = float(ga4.iloc[-1]["bet_handle"]), int(ga4["launch_week"].max())
                for _, rf4 in gf4.iterrows():
                    fw4 = int(rf4["proj_week"])
                    bnd_at = bnd_p25_t4[bnd_p25_t4["launch_week"] == min(fw4, int(bnd_p25_t4["launch_week"].max()))]
                    if bnd_at.empty:
                        continue
                    p25_at = float(bnd_at["p25"].iloc[0])
                    if float(rf4["projected_bet"]) < p25_at:
                        risk_rows4.append({"Game": str(t4_gnames.get(gid4, gid4)), "Current Week": mwk4, "Current Bet": cb4,
                                           "At-Risk Week": fw4, "Projected Bet": float(rf4["projected_bet"]),
                                           "Low-End Threshold": p25_at, "Gap to Threshold": float(rf4["projected_bet"]) - p25_at})
                        break
            if risk_rows4:
                risk_df4 = pd.DataFrame(risk_rows4).sort_values("Gap to Threshold")
                st.dataframe(risk_df4.style.format({"Current Bet": "${:,.0f}", "Projected Bet": "${:,.0f}",
                                                    "Low-End Threshold": "${:,.0f}", "Gap to Threshold": "${:+,.0f}"}),
                            use_container_width=True, hide_index=True)
            else:
                st.markdown(f'<div class="ann ann-g"><span class="ann-tag">OK</span><div class="ann-body">'
                            f'All games are projected to remain in the normal performance range over the next 4 weeks.</div></div>',
                            unsafe_allow_html=True)
        else:
            st.info("Not enough data for fleet forecast — games need at least 2 weeks of history.")
    else:
        ribbon("Similar Launches & What to Expect", f"{top_title} · matched peers, trajectory, and where it's heading", T2, "PEERS")
        st.caption("The bold line is this game; dashed lines are matched peers. Every comparison on this page is against "
                  "this game's own peer group, not the whole fleet.")

        peer_catalog_s = load_sql_catalog()
        if not peer_catalog_s.empty:
            peer_pool_s = L.find_peer_pool(peer_catalog_s, df, sel_id, scale_tolerance=scale_x)
            peer_df_s = df[df["game_id"].isin(peer_pool_s["ids"])]
        else:
            peer_pool_s = {"family": "whole fleet (no catalog)", "pool_size": int(df["game_id"].nunique()), "scale_applied": False}
            peer_df_s = df
        _card_style_s = f"background:{SURFACE};border:1px solid {BORDER};border-radius:10px;padding:14px 16px"

        def chart_head(title, help_text):
            st.markdown(f'<div class="card-title">{title}</div>'
                        f'<div style="font-size:12px;color:{T2};margin:-2px 0 6px">{help_text}</div>',
                        unsafe_allow_html=True)

        def _band3(bd, col, w):
            r = bd[bd["launch_week"] == w] if not bd.empty else pd.DataFrame()
            return float(r[col].iloc[0]) if not r.empty and col in r.columns and pd.notna(r[col].iloc[0]) else None

        def _at3(dfin, col, w):
            r = dfin[dfin["launch_week"] == w]
            return float(r[col].iloc[0]) if not r.empty and pd.notna(r[col].iloc[0]) else None

        with st.spinner("Finding similar games…"):
            peers = find_peers_scaled(df, sel_id, kpi="bet_decay", n_weeks=n_match, top_k=5, scale_tolerance=scale_x)

        peer_ids, peer_names = [], {}
        if not peers:
            st.info("No peers found. Try a wider match window or scale tolerance.")
        else:
            scale_used = peers[0].get("scale_constrained", True)
            if not scale_used:
                st.markdown(f'<div class="ann ann-a"><span class="ann-tag">WATCH</span>'
                            f'<div class="ann-body">Not enough scale-matched peers (within {scale_x}×). '
                            f'Showing best shape-matches regardless of scale.</div></div>', unsafe_allow_html=True)

            peer_ids = [p["game_id"] for p in peers]
            peer_names = {p["game_id"]: p["game_name"] for p in peers}

            # Peer table (Hold %/Retention at week 4 for this game + its 5 closest matches —
            # replaces the old pair of fleet-wide ranked bar charts) now lives in a collapsed
            # expander instead of always-on, so it doesn't eat vertical space by default.
            _peer_basis_label = (f"Peer basis: {peers[0].get('peer_family', 'unknown')} · {len(peers)} peers shown"
                                 + (f" · {peers[0]['peer_fallback_reason']}" if peers[0].get("peer_fallback_reason") else ""))
            with st.expander(_peer_basis_label):
                wk4_map = df[df["launch_week"] == 4].set_index("game_id")
                pt = pd.DataFrame(peers)
                pt["similarity"] = pt.apply(lambda r: _similarity_pct(r["distance"], r.get("n_match_weeks")), axis=1)
                pt["w0_bet"] = pt["game_id"].map(df[df["launch_week"] == 0].set_index("game_id")["bet_handle"].to_dict())
                pt["hold_wk4"] = pt["game_id"].map(wk4_map["hold_pct"].to_dict())
                pt["decay_wk4"] = pt["game_id"].map(wk4_map["bet_decay"].to_dict())
                show_cols = [c for c in ["game_name", "similarity", "distance", "n_match_weeks", "total_weeks", "w0_bet", "hold_wk4", "decay_wk4"] if c in pt.columns]
                pt_disp = pt[show_cols].copy()
                pt_disp.columns = ["Game", "Similarity %", "Shape Distance", "Match Wks", "Total Wks", "W0 Bet ($)", "Hold % Wk4", "Retention % Wk4"][:len(show_cols)]

                self_wk4 = gdf[gdf["launch_week"] == 4]
                self_row = pd.DataFrame([{
                    "Game": f"{top_title} (this game)",
                    "Similarity %": np.nan, "Shape Distance": np.nan, "Match Wks": np.nan,
                    "Total Wks": int(gdf["launch_week"].max()) + 1,
                    "W0 Bet ($)": float(gdf[gdf["launch_week"] == 0]["bet_handle"].iloc[0]) if not gdf[gdf["launch_week"] == 0].empty else np.nan,
                    "Hold % Wk4": float(self_wk4["hold_pct"].iloc[0]) if not self_wk4.empty and pd.notna(self_wk4["hold_pct"].iloc[0]) else np.nan,
                    "Retention % Wk4": float(self_wk4["bet_decay"].iloc[0]) if not self_wk4.empty and pd.notna(self_wk4["bet_decay"].iloc[0]) else np.nan,
                }])
                pt_disp = pd.concat([self_row, pt_disp], ignore_index=True)
                st.dataframe(
                    pt_disp.style.format({"Similarity %": "{:.1f}%", "Shape Distance": "{:.2f}", "W0 Bet ($)": "${:,.0f}",
                                         "Hold % Wk4": "{:.1f}%", "Retention % Wk4": "{:.1f}%"}, na_rep="–"),
                    use_container_width=True, hide_index=True)

        # ── Trajectory overlay | Bet Decay vs. Peer Range — side by side to save vertical space ──
        col_traj, col_decay = st.columns(2)
        with col_traj:
            if not peers:
                st.info("No matched-peer trajectory available.")
            else:
                bnd_decay = bkpi(df, "bet_decay")
                fp = go.Figure()
                if not bnd_decay.empty:
                    fill_band(fp, bnd_decay["launch_week"], bnd_decay["p25"], bnd_decay["p75"], "rgba(91,89,78,0.12)", "Middle 50% range")
                    fp.add_trace(go.Scatter(x=bnd_decay["launch_week"], y=bnd_decay["p50"], line=dict(color=T3, width=1, dash="dot"), name="Platform median"))
                peer_colors = [MOVE_DOWN, "#9A7A3D", MOVE_UP, "#5B7A9C", "#9C5B7A"]
                for i, pid in enumerate(peer_ids):
                    pd_ = df[df["game_id"] == pid].sort_values("launch_week")
                    fp.add_trace(go.Scatter(x=pd_["launch_week"], y=pd_["bet_decay"], mode="lines",
                                            line=dict(color=peer_colors[i % 5], width=1.5, dash="dash"), name=peer_names[pid][:26], opacity=0.8))
                fp.add_trace(go.Scatter(x=gdf["launch_week"], y=gdf["bet_decay"], mode="lines+markers",
                                        line=dict(color=TEXT, width=2.5), marker=dict(size=6), name=top_title[:26]))
                fp.add_vline(x=n_match - 0.5, line=dict(color=T3, width=1.5, dash="dash"), annotation_text=f"Match window ({n_match}w)")
                # Cap the view at 100 weeks (or this game's own length if longer) — some matched
                # peers run 250+ weeks, which was stretching the axis and squeezing this game's
                # own trajectory into a sliver at the left edge. Peer data itself isn't touched,
                # just cropped from view; this game's own line is never clipped.
                _traj_x_max = max(100, int(gdf["launch_week"].max()) + 5)
                fp.update_xaxes(title_text="Launch week", automargin=True, range=[-2, _traj_x_max], autorange=False)
                fp.update_yaxes(title_text="Bet decay % (week 0 = 100)", automargin=True)
                chart_head("Trajectory overlay — this game vs. matched peers",
                           "Bold line = this game; dashed = closest shape-matched peers; band = fleet middle-50%. "
                           "Double-click the chart to enlarge.")
                st.plotly_chart(plotly_base(fp, h=430), use_container_width=True, theme=None)
        with col_decay:
            # Bet Decay vs. Peer Range — is this game losing engagement faster than peers?
            band_decay_s = bkpi(peer_df_s, "bet_decay")
            wks_x3 = sorted(gdf["launch_week"].dropna().unique().astype(int).tolist())
            fig_bd = go.Figure()
            fill_band(fig_bd, wks_x3, [_band3(band_decay_s, "p10", w) for w in wks_x3], [_band3(band_decay_s, "p90", w) for w in wks_x3],
                      "rgba(91,89,78,0.10)", "Peer range (P10–P90)")
            fig_bd.add_trace(go.Scatter(x=wks_x3, y=[_band3(band_decay_s, "p50", w) for w in wks_x3], mode="lines",
                                         line=dict(color=T3, dash="dot", width=1.5), name="Peer median"))
            fig_bd.add_trace(go.Scatter(x=wks_x3, y=[_at3(gdf, "bet_decay", w) for w in wks_x3], mode="lines+markers",
                                         line=dict(color=MOVE_DOWN, width=2.5), marker=dict(size=5), name=top_title[:22]))
            fig_bd.update_xaxes(title_text="Launch week", automargin=True)
            fig_bd.update_yaxes(title_text="Bet decay % (week 0 = 100)", automargin=True)
            chart_head("Bet Decay vs. Peer Range",
                       f"Bet handle as a % of week 0, vs. the {peer_pool_s['pool_size']}-game peer range "
                       f"({peer_pool_s['family']}). Below the band = fading faster than similar games.")
            st.plotly_chart(plotly_base(fig_bd, h=430), use_container_width=True, theme=None)

        # ── Forecast: ONE actual line through a continuous peer band (history) + ONE
        # forecast cone (future) — no separate Best/Worst/Base scenario lines. Weeks outside
        # the peer's P10–P90 range get a highlighted marker, same idea as a clean "actual vs.
        # expectation band" chart — just one case, one line, dots where it breaks the band.
        band_net_s = bkpi(peer_df_s, "net_rev")
        catalog_fc = load_sql_catalog()
        blended = L.fit_blended_forecast(df, sel_id, catalog_fc, n_forecast_weeks=horizon, kpi="bet_handle", scale_tolerance=scale_x)
        rh = blended["reliable_horizon"]
        game_portfolio = L.compare_game_to_portfolio(df, sel_id, kpi="bet_handle")
        retro = L.compute_retrospective_standing(df, sel_id, kpi="bet_handle")

        fcast = blended["forecast"]
        wks_rel = [f["week"] for f in fcast if not f["directional"]]
        base_rel = [f["value"] for f in fcast if not f["directional"]]
        best_rel = [f["upper"] for f in fcast if not f["directional"]]
        wrst_rel = [f["lower"] for f in fcast if not f["directional"]]
        wks_dir = [f["week"] for f in fcast if f["directional"]]
        base_dir = [f["value"] for f in fcast if f["directional"]]
        best_dir = [f["upper"] for f in fcast if f["directional"]]
        wrst_dir = [f["lower"] for f in fcast if f["directional"]]

        avg_hold = float(gdf["hold_pct"].dropna().mean()) / 100.0 if gdf["hold_pct"].notna().any() else 0.05
        avg_hold = max(0.01, min(avg_hold, 0.5))
        nr_base_rel = [v * avg_hold for v in base_rel]
        nr_best_rel = [v * avg_hold for v in best_rel]
        nr_wrst_rel = [v * avg_hold for v in wrst_rel]
        nr_base_dir = [v * avg_hold for v in base_dir]
        nr_best_dir = [v * avg_hold for v in best_dir]
        nr_wrst_dir = [v * avg_hold for v in wrst_dir]

        s_best_nr, s_base_nr, s_worst_nr = sum(nr_best_rel + nr_best_dir), sum(nr_base_rel + nr_base_dir), sum(nr_wrst_rel + nr_wrst_dir)
        s_range = s_best_nr - s_worst_nr
        trailing_wks = min(horizon, int(max_wk) + 1)
        last_q_nr = float(gdf[gdf["launch_week"] > (max_wk - trailing_wks)]["net_rev"].sum())

        def _vs_last(v):
            if last_q_nr <= 0:
                return f"Last {trailing_wks}w: n/a"
            chg = (v - last_q_nr) / last_q_nr * 100
            return f"vs last {trailing_wks}w {_usd(last_q_nr)}  {'▲' if chg >= 0 else '▼'} {abs(chg):.0f}%"

        blend_parts = []
        if blended.get("self_weight") is not None:
            blend_parts.append(f"self-trend {blended['self_weight']*100:.0f}% · peer-trend {blended['peer_weight']*100:.0f}% "
                               f"({len(blended['peers_used'])} {blended['family']} peers)")
        if rh < horizon:
            blend_parts.append(f"reliable through week +{rh} — dashed = directional only")
        if blended.get("message"):
            blend_parts.append(blended["message"])

        pt_verdict = game_portfolio.get("verdict", "insufficient data")
        if pt_verdict != "insufficient data" and game_portfolio.get("target_slope_pct") is not None:
            pt_target, pt_median = game_portfolio["target_slope_pct"], game_portfolio["portfolio_median_pct"]
            kind = "g" if pt_verdict == "outperforming" else ("r" if pt_verdict == "underperforming" else "a")
            st.markdown(f'{badge(pt_verdict.upper(), kind)} <span style="font-size:12px;color:{T2}">Portfolio (last 12w): '
                       f'this game {pt_target:+.1f}%/wk vs median {pt_median:+.1f}%/wk</span>', unsafe_allow_html=True)

        if blend_parts:
            st.caption("Forecast: " + " · ".join(blend_parts))

        RETRO_CLR = {"outperforming": "rgba(61,122,106,0.10)", "underperforming": "rgba(179,74,58,0.10)"}
        retro_bands = []
        rp = None
        for rpt in (retro or []):
            vd = rpt["verdict"]
            if vd == "in line":
                if rp: retro_bands.append(rp); rp = None
                continue
            if rp and rp["verdict"] == vd:
                rp["end"] = rpt["launch_week"]
            else:
                if rp: retro_bands.append(rp)
                rp = {"start": rpt["launch_week"], "end": rpt["launch_week"], "verdict": vd}
        if rp: retro_bands.append(rp)

        net_arr = gdf["net_rev"].fillna(0).values
        wks_hist = list(range(len(net_arr)))
        p10_h = [_band3(band_net_s, "p10", w) for w in wks_hist]
        p90_h = [_band3(band_net_s, "p90", w) for w in wks_hist]
        p50_h = [_band3(band_net_s, "p50", w) for w in wks_hist]
        out_x, out_y = [], []
        for w, v, lo, hi in zip(wks_hist, net_arr, p10_h, p90_h):
            if lo is not None and hi is not None and (v < lo or v > hi):
                out_x.append(w); out_y.append(v)

        ff = go.Figure()
        for rb in retro_bands:
            ff.add_vrect(x0=rb["start"] - 0.5, x1=rb["end"] + 0.5, fillcolor=RETRO_CLR[rb["verdict"]], line_width=0, layer="below")
        fill_band(ff, wks_hist, p10_h, p90_h, "rgba(91,89,78,0.10)", "Peer range (P10–P90)")
        ff.add_trace(go.Scatter(x=wks_hist, y=p50_h, mode="lines", line=dict(color=T3, width=1.2, dash="dot"), name="Peer median"))
        if wks_rel:
            fill_band(ff, wks_rel, nr_wrst_rel, nr_best_rel, "rgba(47,125,110,0.16)", "Forecast 80% CI")
            ff.add_trace(go.Scatter(x=wks_rel, y=nr_base_rel, mode="lines", line=dict(color=MOVE_UP, width=2.2, dash="dash"), name="Forecast"))
        if wks_dir:
            sw = wks_rel[-1] if wks_rel else max_wk
            sb = nr_base_rel[-1] if wks_rel else float(net_arr[-1])
            ff.add_trace(go.Scatter(x=[sw] + wks_dir, y=[sb] + nr_base_dir, mode="lines", line=dict(color=MOVE_UP, width=1.5, dash="dot"), name="Forecast (directional)", showlegend=False))
        ff.add_trace(go.Scatter(x=wks_hist, y=list(net_arr), mode="lines+markers", line=dict(color=TEXT, width=2.5),
                                 marker=dict(size=5, color=TEXT), name="Actual"))
        if out_x:
            ff.add_trace(go.Scatter(x=out_x, y=out_y, mode="markers",
                                     marker=dict(size=9, color=MOVE_DOWN, line=dict(width=1.5, color=SURFACE)), name="Outside peer range"))
        ff.add_vline(x=len(net_arr) - 0.5, line=dict(color=T3, width=1))
        ff.update_xaxes(title_text="Launch week", automargin=True)
        ff.update_yaxes(title_text="Net revenue ($)", automargin=True)

        # ── Net Revenue chart | First 30/60/90 Days table — side by side ────────────
        col_nr, col_ms = st.columns([7, 5])
        with col_nr:
            chart_head("Net Revenue — Actual vs. Peer Range + Forecast",
                       "Black line = weekly Game Net; grey band = peer P10–P90; red dots = weeks outside that range. "
                       "Past the divider, the green cone is the blended self + peer forecast.")
            st.plotly_chart(plotly_base(ff, h=480), use_container_width=True, theme=None)
            if rh < horizon:
                st.caption(f"Forecast band exceeds reliable range after week +{rh}.")
        with col_ms:
            _exp_rows = []
            for _label, _days in (("First 30 Days", 30), ("First 60 Days", 60), ("First 90 Days", 90)):
                _wk_cut = _days // 7
                if int(gdf["launch_week"].max()) < _wk_cut:
                    _exp_rows.append({"Milestone": _label, "Actual": np.nan, "Expected": np.nan})
                    continue
                _actual = float(gdf[gdf["launch_week"] <= _wk_cut]["net_rev"].sum())
                _bnd_cut = band_net_s[band_net_s["launch_week"] <= _wk_cut]
                _expected = float(_bnd_cut["p50"].sum()) if not _bnd_cut.empty else np.nan
                _exp_rows.append({"Milestone": _label, "Actual": _actual, "Expected": _expected})
            exp_df = pd.DataFrame(_exp_rows)
            _exp_html = []
            for _, r in exp_df.iterrows():
                if pd.isna(r["Actual"]):
                    _exp_html.append(f'<div style="display:flex;justify-content:space-between;padding:9px 2px;border-bottom:1px solid {BORDER}">'
                                      f'<span style="font-weight:600">{r["Milestone"]}</span><span style="color:{T2}">Not reached yet</span></div>')
                    continue
                _delta = (r["Actual"] - r["Expected"]) / r["Expected"] * 100 if pd.notna(r["Expected"]) and r["Expected"] != 0 else None
                _delta_str = f'{"▲" if _delta >= 0 else "▼"} {abs(_delta):.0f}%' if _delta is not None else "–"
                _delta_clr = MOVE_UP if (_delta is not None and _delta >= 0) else MOVE_DOWN
                _exp_html.append(
                    f'<div style="padding:9px 2px;border-bottom:1px solid {BORDER}">'
                    f'<div style="font-weight:600">{r["Milestone"]}</div>'
                    f'<div style="display:flex;justify-content:space-between;font-size:13px;margin-top:2px">'
                    f'<span>Actual {_usd(r["Actual"])}</span>'
                    f'<span style="color:{T2}">Expected {_usd(r["Expected"]) if pd.notna(r["Expected"]) else "–"}</span>'
                    f'<span style="color:{_delta_clr};font-weight:600">{_delta_str}</span></div></div>')
            st.markdown(f'<div style="{_card_style_s}"><div class="card-title">First 30 / 60 / 90 Days — Actual vs. Expected (peer median)</div>'
                        + "".join(_exp_html) + '</div>', unsafe_allow_html=True)

        # ── Cannibalization Check — scoped to the locations where THIS game is actually live ──
        with st.expander("Is there any cannibalization? — did other games' revenue drop at the SAME locations after this game launched there?"):
            tgt_dt = pd.to_datetime(meta["launch_date"], errors="coerce")
            if pd.isna(tgt_dt):
                st.info("Launch date unavailable — cannot compute cannibalization window.")
            elif platform == "EdgeLabs":
                st.info("EdgeLabs is casino-based (no location dimension) — cannibalization can't be scoped by location here.")
            else:
                cw = 4
                before_start, before_end = tgt_dt - pd.Timedelta(weeks=cw), tgt_dt - pd.Timedelta(days=1)
                after_start, after_end = tgt_dt, min(tgt_dt + pd.Timedelta(weeks=cw), pd.Timestamp.today())
                with st.spinner("Loading location data…"):
                    geo_life = load_geo_detail(platform, str(_LIFETIME_START), str(dt.date.today()))
                    geo_pre = load_geo_detail(platform, str(before_start.date()), str(before_end.date()))
                    geo_post = load_geo_detail(platform, str(after_start.date()), str(after_end.date()))

                target_locs = set(geo_life[geo_life["game_id"].astype("Int64") == int(sel_id)]["location_name"]) if not geo_life.empty else set()
                if not target_locs:
                    st.info(f"No location history found for {top_title} — cannot scope cannibalization to its footprint.")
                elif geo_pre.empty or geo_post.empty:
                    st.info("Not enough location data both before and after launch to compute cannibalization.")
                else:
                    pre_s = geo_pre[geo_pre["location_name"].isin(target_locs) & (geo_pre["game_id"].astype("Int64") != int(sel_id))
                                    & ~geo_pre["game_id"].apply(L._is_hr)]
                    post_s = geo_post[geo_post["location_name"].isin(target_locs) & (geo_post["game_id"].astype("Int64") != int(sel_id))
                                      & ~geo_post["game_id"].apply(L._is_hr)]
                    pre_g = pre_s.groupby(["game_id", "game_name"])["bet"].sum().rename("bet_before")
                    post_g = post_s.groupby(["game_id", "game_name"])["bet"].sum().rename("bet_after")
                    cdf = pd.concat([pre_g, post_g], axis=1).reset_index().fillna(0)
                    cdf = cdf[cdf["bet_before"] > 0]
                    if cdf.empty:
                        st.info(f"No other games had bet activity at {top_title}'s {len(target_locs)} locations before its launch.")
                    else:
                        cdf["Change ($)"] = cdf["bet_after"] - cdf["bet_before"]
                        cdf["Change %"] = cdf["Change ($)"] / cdf["bet_before"] * 100
                        cdf = cdf.rename(columns={"game_name": "Game", "bet_before": "Bet Before ($)", "bet_after": "Bet After ($)"})
                        cdf = cdf.sort_values("Change ($)")
                        nd_tot, ng_tot = (cdf["Change ($)"] < 0).sum(), (cdf["Change ($)"] >= 0).sum()
                        st.caption(f"Scoped to the **{len(target_locs)}** locations where **{top_title}** is live. Since it launched "
                                  f"({meta['launch_date']}), **{nd_tot}** other games saw a bet-handle drop at those same locations "
                                  f"and **{ng_tot}** saw a rise, over the {cw}-week window.")

                        chart_c = pd.concat([cdf.head(min(10, nd_tot)), cdf.tail(min(5, ng_tot))]).drop_duplicates("Game").sort_values("Change ($)")
                        col_c = [MOVE_DOWN if v < 0 else MOVE_UP for v in chart_c["Change ($)"]]
                        fig_c = go.Figure(go.Bar(y=chart_c["Game"], x=chart_c["Change ($)"], orientation="h", marker_color=col_c,
                                                 text=[f'${abs(v):,.0f}' for v in chart_c["Change ($)"]], textposition="inside",
                                                 hovertemplate="<b>%{y}</b><br>Change: $%{x:,.0f}<extra></extra>"))
                        fig_c.update_layout(yaxis=dict(autorange="reversed"))
                        fig_c.update_xaxes(title_text="Change in bet handle ($), same locations", automargin=True)
                        st.markdown('<div class="card"><div class="card-title">Top drops / gains in other games\' bet handle — this game\'s locations only</div>', unsafe_allow_html=True)
                        st.plotly_chart(plotly_base(fig_c, h=max(340, len(chart_c) * 45 + 80)), use_container_width=True, theme=None)
                        st.markdown('</div>', unsafe_allow_html=True)

                        st.dataframe(
                            cdf[["Game", "Bet Before ($)", "Bet After ($)", "Change ($)", "Change %"]]
                            .style.format({"Bet Before ($)": "${:,.0f}", "Bet After ($)": "${:,.0f}",
                                          "Change ($)": "${:+,.0f}", "Change %": "{:+.1f}%"}),
                            use_container_width=True, hide_index=True,
                        )
                        st.markdown(
                            f'<div class="ann ann-a"><span class="ann-tag">NOTE</span><div class="ann-body">'
                            f'Scoped to locations where {top_title} is live, but a change in another game\'s bet handle there '
                            f'still does not prove cannibalization — seasonal patterns, concurrent launches, or terminal swaps '
                            f'at those same stores may also contribute. Use this as a starting point, not proof.</div></div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# TAB — FULL BREAKDOWN
# ══════════════════════════════════════════════════════════════════
with tab_full:
    render_generate_report_button("full")
    if is_all:
        ribbon("All Games — Performance Overview", f"{len(nonhr)} {platform} games, how each metric compares across the portfolio", T2, "DETAIL")

        rows5 = []
        for _, gc5 in nonhr.iterrows():
            gid5 = gc5["game_id"]
            g5 = df[df["game_id"] == gid5].sort_values("launch_week")
            if len(g5) < 2:
                continue
            lr5 = g5.iloc[-1]
            rows5.append({"game_id": gid5, "game_name": gc5["game_name"],
                          "Hold %": float(lr5["hold_pct"]) if pd.notna(lr5.get("hold_pct")) else np.nan,
                          "Net Revenue": float(g5["net_rev"].sum()),
                          "Bet Decay": float(lr5["bet_decay"]) if pd.notna(lr5.get("bet_decay")) else np.nan,
                          "ARPU": float(lr5["arpu"]) if pd.notna(lr5.get("arpu")) else np.nan,
                          "Weeks Live": int(lr5["launch_week"]) + 1})
        sc5 = pd.DataFrame(rows5)

        if len(sc5) >= 2:
            c1, c2 = st.columns(2)
            with c1:
                sub1 = sc5.dropna(subset=["Hold %", "Net Revenue"])
                fs1 = go.Figure(go.Scatter(x=sub1["Hold %"], y=sub1["Net Revenue"], mode="markers", text=sub1["game_name"],
                    marker=dict(color=sub1["Weeks Live"], colorscale=[[0, T3], [1, TEXT]], size=9, opacity=0.85,
                               colorbar=dict(title="Weeks Live", thickness=13, len=0.65)),
                    hovertemplate="<b>%{text}</b><br>Hold: %{x:.1f}%<br>Net Rev: $%{y:,.0f}<extra></extra>"))
                fs1.update_xaxes(title_text="Hold %"); fs1.update_yaxes(title_text="Net revenue ($)")
                st.markdown('<div class="card"><div class="card-title">Hold % vs. Net Revenue — color = weeks live</div>', unsafe_allow_html=True)
                st.plotly_chart(plotly_base(fs1, h=425), use_container_width=True, theme=None)
                st.markdown('</div>', unsafe_allow_html=True)
            with c2:
                sub2 = sc5.dropna(subset=["Bet Decay", "Net Revenue", "Hold %"])
                if len(sub2) >= 2:
                    fs2 = go.Figure(go.Scatter(x=sub2["Bet Decay"], y=sub2["Net Revenue"], mode="markers", text=sub2["game_name"],
                        marker=dict(color=sub2["Hold %"], colorscale=[[0, MOVE_DOWN], [0.5, "#E8DCC8"], [1, MOVE_UP]], size=9, opacity=0.85,
                                   colorbar=dict(title="Hold %", thickness=13, len=0.65)),
                        hovertemplate="<b>%{text}</b><br>Decay: %{x:.1f}%<br>Net Rev: $%{y:,.0f}<extra></extra>"))
                    fs2.update_xaxes(title_text="Bet decay % (latest)"); fs2.update_yaxes(title_text="Net revenue ($)")
                    st.markdown('<div class="card"><div class="card-title">Player Retention vs. Net Revenue — color = hold %</div>', unsafe_allow_html=True)
                    st.plotly_chart(plotly_base(fs2, h=425), use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)

            ribbon("How Metrics Are Distributed Across Launch Weeks", "Each ridge is one tenure checkpoint — shows how the fleet's whole spread shifts as games age, not just one snapshot", T2, "SPREAD")
            _nonhr_ids5 = set(nonhr["game_id"])
            _fleet_wk5 = df[df["game_id"].isin(_nonhr_ids5)]
            _checkpoints5 = [0, 4, 8, 13, 26, 52]
            ridge_kpis = [("hold_pct", "Hold %", "%"), ("bet_decay", "Bet Decay %", "%"), ("net_rev", "Net Revenue", "$")]
            if "arpu" in _fleet_wk5.columns and _fleet_wk5["arpu"].notna().sum() >= 15:
                ridge_kpis.append(("arpu", "ARPU", "$"))
            ridge_cols = st.columns(len(ridge_kpis))
            for rcol, (kcol, klabel, kunit) in zip(ridge_cols, ridge_kpis):
                with rcol:
                    groups5 = []
                    for wk5 in _checkpoints5:
                        vals5 = pd.to_numeric(_fleet_wk5.loc[_fleet_wk5["launch_week"] == wk5, kcol], errors="coerce").dropna()
                        vals5 = vals5[np.isfinite(vals5)]
                        if len(vals5) >= 5:
                            groups5.append((f"Wk {wk5}", vals5.values))
                    st.markdown(f'<div class="card"><div class="card-title">{klabel} — Distribution by Tenure</div>', unsafe_allow_html=True)
                    if len(groups5) < 2:
                        st.caption(f"Not enough weeks with 5+ games reporting {klabel} yet to show a distribution shift.")
                    else:
                        fig_rl5 = _ridgeline_fig(groups5, unit=kunit, height=330)
                        st.plotly_chart(fig_rl5, use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("Not enough data for fleet scatter analysis — need at least 2 games with 2+ weeks.")
    else:
        ribbon("Full Breakdown — Every Metric, Week by Week", top_title, T2, "DETAIL")
        st.caption("Pick any metric to see week-by-week performance vs. all other games. The shaded band is the normal range.")

        # Rank & Share — lifetime Game Net rank among all non-HR games on this platform.
        _life_net_all = df[~df["game_id"].apply(L._is_hr)].groupby("game_id")["net_rev"].sum()
        _total_net_all = float(_life_net_all.sum())
        _rank_sorted = _life_net_all.sort_values(ascending=False)
        _this_net_fb = float(_life_net_all.get(sel_id, 0.0))
        _this_rank_fb = (int(_rank_sorted.index.get_loc(sel_id)) + 1) if sel_id in _rank_sorted.index else None
        _n_ranked_fb = len(_rank_sorted)
        _share_pct_fb = min((_this_net_fb / _total_net_all * 100), 99.9) if _total_net_all > 0 else 0.0

        # ARPU and SPP are per-player-PER-WEEK rates (net_rev/players and spins/players on a
        # weekly row) — spell the /wk out, since "Spins/Player" alone reads as per-day next to
        # the genuinely per-day "Spins/Store/Day" on Is It On Track.
        KPI_CARDS = [("hold_pct", "Hold %", _pct, True), ("bet_decay", "Bet Decay (↓ better)", _pct, False),
                     ("arpu", "ARPU net (per player/wk)", _usd, True),
                     ("spp", "Spins (per player/wk)", lambda v: f"{v:,.0f}", True)]
        card_items = [{
            "label": "Rank & Share (lifetime)",
            "value": f"#{_this_rank_fb} of {_n_ranked_fb}" if _this_rank_fb else "–",
            "sub": f"{_share_pct_fb:.1f}% of {platform} lifetime Game Net",
        }]
        for kpi_c, lbl_c, fmtc, hi_c in KPI_CARDS:
            bnd_c = bkpi(df, kpi_c)
            gv_c = gdf[gdf[kpi_c].notna()] if kpi_c in gdf.columns else pd.DataFrame()
            if gv_c.empty:
                card_items.append({"label": lbl_c, "value": "–", "sub": "no data"})
                continue
            val_c, wk_c = float(gv_c.iloc[-1][kpi_c]), int(gv_c.iloc[-1]["launch_week"])
            brow_c = bnd_c[bnd_c["launch_week"] == wk_c] if not bnd_c.empty else pd.DataFrame()
            if not brow_c.empty:
                cls_c = _flag_cls(val_c, float(brow_c["p25"].iloc[0]), float(brow_c["p75"].iloc[0]),
                                  float(brow_c["p10"].iloc[0]), float(brow_c["p90"].iloc[0]), higher_is_better=hi_c)
                lbl_map = {"g": "OK", "a": "WATCH", "r": "FLAG", "n": "–"}
                sub_c = f"Platform median: {fmtc(float(brow_c['p50'].iloc[0]))} {badge(lbl_map[cls_c], cls_c)}"
            else:
                sub_c = f"wk {wk_c}"
            card_items.append({"label": f"{lbl_c} (wk {wk_c})", "value": fmtc(val_c), "sub": sub_c})
        krow(card_items)

        # Only the 4 KPIs not already covered elsewhere in the dashboard — Bet Decay %, Hold %,
        # and Avg Bet/Spin ($) are dropped here since Is It On Track / Similar Launches &
        # What to Expect / Compare already chart all three.
        KPI_PILLS = [("arpu", "ARPU net (per player/wk)", "ARPU ($ per player/wk)", _usd),
                     ("spp", "Spins (per player/wk)", "Spins per player/wk", lambda v: f"{v:,.0f}"),
                     ("stores", "Active casinos", "Casinos", lambda v: f"{v:,.0f}"),
                     ("player_decay", "Player decay %", "Pl. decay %", _pct)]
        pill_labels = [p[1] for p in KPI_PILLS]
        sel_pill = st.radio("KPI", pill_labels, horizontal=True, key="pill_v2", label_visibility="collapsed")
        sel_kpi_col, _, sel_ylab, sel_fmt = KPI_PILLS[pill_labels.index(sel_pill)]

        bnd_sel = bkpi(df, sel_kpi_col)
        gk = gdf[["launch_week", sel_kpi_col]].dropna() if sel_kpi_col in gdf.columns else pd.DataFrame()
        if gk.empty and bnd_sel.empty:
            st.info("No data for this metric yet.")

        fb = go.Figure()
        if not bnd_sel.empty:
            fill_band(fb, bnd_sel["launch_week"], bnd_sel["p10"], bnd_sel["p90"], "rgba(91,89,78,0.06)", "P10–P90")
            fill_band(fb, bnd_sel["launch_week"], bnd_sel["p25"], bnd_sel["p75"], "rgba(91,89,78,0.14)", "P25–P75")
            fb.add_trace(go.Scatter(x=bnd_sel["launch_week"], y=bnd_sel["p50"], line=dict(color=T3, width=1.5, dash="dot"), name="Platform median"))
        if not gk.empty:
            fb.add_trace(go.Scatter(x=gk["launch_week"], y=gk[sel_kpi_col], mode="lines+markers",
                                    line=dict(color=TEXT, width=2.5), marker=dict(size=5), name=top_title[:24]))
        fb.update_xaxes(title_text="Launch week")
        fb.update_yaxes(title_text=sel_ylab)
        st.markdown(f'<div class="card"><div class="card-title">{sel_pill} — this game vs. platform range</div>', unsafe_allow_html=True)
        st.plotly_chart(plotly_base(fb, h=395), use_container_width=True, theme=None)
        st.markdown('</div>', unsafe_allow_html=True)

        # A flat/near-zero "Active stores" line can look broken when it's actually a real,
        # highly-concentrated deployment (e.g. one flagship site running high volume) — badge
        # it so that reads as a finding instead of a chart bug.
        if sel_kpi_col == "stores":
            foot_fb = load_game_footprint(platform, sel_id)
            if foot_fb.empty:
                st.caption("No location-level footprint data available for this game." if platform != "EdgeLabs"
                           else "EdgeLabs is casino-based (no location dimension) — footprint doesn't apply.")
            else:
                n_locs_fb = int(foot_fb["loc_id"].nunique())
                total_net_fb = float(foot_fb["net_rev"].sum())
                top_net_fb = float(foot_fb.sort_values("net_rev", ascending=False)["net_rev"].iloc[0]) if total_net_fb != 0 else 0.0
                top_share_fb = (top_net_fb / total_net_fb * 100) if total_net_fb > 0 else None
                conc_kind_fb = "r" if (top_share_fb is not None and top_share_fb >= 40) else ("a" if (top_share_fb is not None and top_share_fb >= 25) else "g")
                conc_label_fb = "Concentrated" if conc_kind_fb == "r" else ("Watch" if conc_kind_fb == "a" else "Spread out")
                st.markdown(f'{badge(conc_label_fb, conc_kind_fb)} <span style="font-size:12px;color:{T2}">'
                           f'{n_locs_fb} location{"s" if n_locs_fb != 1 else ""} lifetime'
                           + (f" · top location is {_pct(top_share_fb)} of lifetime Game Net" if top_share_fb is not None else "")
                           + '</span>', unsafe_allow_html=True)

        # First 30 / 60 / 90 Days + Lifetime — average of the selected KPI through each
        # checkpoint (a rate/level metric like ARPU or Active Stores is more honestly
        # summarized as a window average than a cumulative sum).
        if not gk.empty:
            _bd_periods = [("First 30 Days", 30), ("First 60 Days", 60), ("First 90 Days", 90), ("Lifetime", None)]
            _bd_items = []
            for _lbl, _days in _bd_periods:
                _sub_bd = gk if _days is None else gk[gk["launch_week"] <= (_days // 7)]
                if _sub_bd.empty:
                    _bd_items.append({"label": _lbl, "value": "–", "sub": "Not reached yet" if _days is not None else "no data"})
                else:
                    _bd_items.append({"label": _lbl, "value": sel_fmt(float(_sub_bd[sel_kpi_col].mean())),
                                      "sub": f"avg over {len(_sub_bd)} wks" if _days is not None else f"avg over all {len(_sub_bd)} wks"})
            st.markdown(f'<div class="eyebrow" style="margin:10px 0 4px">{sel_pill} — First 30 / 60 / 90 Days & Lifetime</div>', unsafe_allow_html=True)
            krow(_bd_items)

        # ── Performance vs Similar Games — percentile-rank radar across 6 metrics ──
        ribbon("Performance vs Similar Games", f"{top_title}'s percentile rank across 6 metrics at week {max_wk} — further from center = stronger", T2, "RADAR")
        _peer_catalog_rad = load_sql_catalog()
        if not _peer_catalog_rad.empty:
            _peer_pool_rad = L.find_peer_pool(_peer_catalog_rad, df, sel_id, scale_tolerance=scale_x)
            # find_peer_pool's "ids" excludes the target itself — pct_rank needs the target's
            # own row present in the frame it's ranked within, or it always returns NaN.
            _peer_df_rad = df[df["game_id"].isin(_peer_pool_rad["ids"] | {sel_id})]
        else:
            _peer_pool_rad = {"family": "whole fleet (no catalog)", "pool_size": int(df["game_id"].nunique())}
            _peer_df_rad = df

        _rk_metrics = [("hold_pct", "Hold %"), ("bet_decay", "Decay"), ("arpu", "ARPU"),
                       ("avg_bet", "Avg bet"), ("spp", "Spins/P"), ("player_decay", "Pl. decay")]
        # A tight peer pool (often just a handful of games) can easily have zero OTHER members
        # at this exact week even when the pool itself is meaningful — pct_rank needs same-week
        # data across every peer, a much higher bar than a single metric/week lookup elsewhere.
        # Fall back to the whole platform per-metric rather than showing "no data" when that happens.
        _rk_pcts, _rk_widened = [], []
        for k, _ in _rk_metrics:
            p = L.pct_rank(_peer_df_rad, sel_id, k, max_wk)
            wid = False
            if pd.isna(p) and _peer_df_rad is not df:
                p = L.pct_rank(df, sel_id, k, max_wk)
                wid = True
            _rk_pcts.append(p)
            _rk_widened.append(wid)
        _rk_pcts_clean = [p if pd.notna(p) else 0.0 for p in _rk_pcts]
        _rk_labels = [l for _, l in _rk_metrics]

        col_radar, col_radar_legend = st.columns(2)
        with col_radar:
            fr = go.Figure(go.Scatterpolar(
                r=_rk_pcts_clean + [_rk_pcts_clean[0]], theta=_rk_labels + [_rk_labels[0]],
                fill="toself", fillcolor="rgba(61,122,106,0.14)",
                line=dict(color=MOVE_UP, width=2), name=top_title[:24]))
            fr.update_layout(
                polar=dict(bgcolor=SURFACE,
                           radialaxis=dict(range=[0, 100], tickfont=dict(size=9), gridcolor=BORDER),
                           angularaxis=dict(tickfont=dict(size=10), gridcolor=BORDER)),
                paper_bgcolor="rgba(0,0,0,0)", height=340,
                margin=dict(l=40, r=40, t=20, b=20), showlegend=False,
            )
            st.plotly_chart(fr, use_container_width=True, theme=None)
        def _ordinal(n):
            n = int(round(n))
            if 11 <= (n % 100) <= 13:
                return f"{n}th"
            return f"{n}{ {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th') }"

        with col_radar_legend:
            _rk_rows_html = []
            for (_, _lbl), _p, _wid in zip(_rk_metrics, _rk_pcts, _rk_widened):
                _p_str = f"{_ordinal(_p)} percentile" if pd.notna(_p) else "no data"
                if _wid:
                    _p_str += " (fleet-wide)"
                _rk_rows_html.append(
                    f'<div style="display:flex;justify-content:space-between;padding:9px 2px;border-bottom:1px solid {BORDER}">'
                    f'<span style="font-weight:600">{_lbl}</span><span style="color:{T2}">{_p_str}</span></div>')
            st.markdown(
                f'<div style="background:{SURFACE};border:1px solid {BORDER};border-radius:10px;padding:14px 16px">'
                f'<div class="card-title">Rank vs. {_peer_pool_rad["pool_size"]}-game peer pool'
                + (f" ({_peer_pool_rad['family']})" if _peer_pool_rad.get("family") else "") + '</div>'
                + "".join(_rk_rows_html) + '</div>', unsafe_allow_html=True)
            _cap_bits = ["Percentile = share of the peer pool this game beats or ties on that metric at its current week.",
                        "ARPU and Spins/P are per player per week, not per day."]
            if any(_rk_widened):
                _cap_bits.append("\"(fleet-wide)\" means the peer pool had no other game at this exact week for that metric, "
                                 "so it widened to the whole platform instead of showing nothing.")
            _cap_bits.append("Player decay isn't direction-flipped here — a high percentile on it means high decay, not low, "
                            "same as every other metric on this chart.")
            st.caption(" ".join(_cap_bits))

        col_players, col_mix = st.columns(2)
        with col_players:
            ribbon("Player Activity Week by Week", "Player count vs. engagement depth", T2, "PLAYERS")
            has_players = "players" in gdf.columns and gdf["players"].notna().any()
            pl_col = "players" if has_players else ("stores" if "stores" in gdf.columns else None)
            spp_col5 = "spp" if "spp" in gdf.columns else None
            if pl_col or spp_col5:
                pa_fig = go.Figure()
                if pl_col:
                    pa_data = gdf[["launch_week", pl_col]].dropna()
                    pa_fig.add_trace(go.Scatter(x=pa_data["launch_week"], y=pa_data[pl_col], mode="lines+markers",
                                                name="Players" if has_players else "Locations (proxy)", line=dict(color=TEXT, width=2.5), marker=dict(size=5)))
                if spp_col5:
                    spp5_data = gdf[["launch_week", spp_col5]].dropna()
                    bnd_spp5 = bkpi(df, spp_col5)
                    if not bnd_spp5.empty:
                        fill_band(pa_fig, bnd_spp5["launch_week"], bnd_spp5["p25"], bnd_spp5["p75"], "rgba(91,89,78,0.10)", "Fleet SPP P25–P75")
                    pa_fig.add_trace(go.Scatter(x=spp5_data["launch_week"], y=spp5_data[spp_col5], mode="lines+markers",
                                                name="Spins / Player / Week", line=dict(color=T3, width=2), marker=dict(size=5), yaxis="y2"))
                    pa_fig.update_layout(yaxis2=dict(title="Spins / player / week", overlaying="y", side="right", showgrid=False))
                pa_fig.update_xaxes(title_text="Launch week")
                pa_fig.update_yaxes(title_text="Players" if has_players else "Locations (proxy)")
                st.plotly_chart(plotly_base(pa_fig, h=395), use_container_width=True, theme=None)
            else:
                st.caption("Player count and spins-per-player columns not available for this platform.")

        with col_mix:
            # ── Revenue Mix — other titles sharing this game's underlying math model ──
            # SkinOf does NOT mean "same title, different theme" — confirmed against live SQL
            # data: it groups titles that share one math model/paytable but can have completely
            # different names and art (e.g. "Bugs to Riches Diamond - Gen2" and "Sword of Z"
            # sharing one parent math model). Labeled accordingly so it doesn't read as a bug
            # when names look unrelated.
            ribbon("Revenue Mix — Shared Math Model", f"Other titles built on the same underlying math/paytable as {top_title}", T2, "MIX")
            peer_catalog_fb = load_sql_catalog()
            _family_ids_fb = L._skin_group_ids(peer_catalog_fb, sel_id) if not peer_catalog_fb.empty else set()
            _df_ids_fb = set(df["game_id"].unique().tolist())
            _family_ids_fb = {i for i in _family_ids_fb if i in _df_ids_fb}
            if not _family_ids_fb:
                st.caption("This title has no tracked math-model relationship (GameCatalogView1.SkinOf) — nothing to split. "
                           "Only ~23% of titles have this populated, so absence doesn't necessarily mean this is a standalone math model.")
            else:
                _all_ids_fb = _family_ids_fb | {sel_id}
                _fam_sub = df[df["game_id"].isin(_all_ids_fb) & ~df["game_id"].apply(L._is_hr)]
                _fam_net = _fam_sub.groupby("game_id")["net_rev"].sum()
                _fam_names = _fam_sub.drop_duplicates("game_id").set_index("game_id")["game_name"]
                _fam_df = pd.DataFrame({"net": _fam_net, "name": _fam_names}).dropna()
                _fam_df = _fam_df[_fam_df["net"] > 0]
                if _fam_df.empty:
                    st.caption("No positive lifetime Game Net to chart across this math-model family.")
                else:
                    fig_mix = _treemap_fig(_fam_df["name"], _fam_df["net"], height=395)
                    st.markdown('<div class="card"><div class="card-title">Lifetime Game Net — titles sharing this math model</div>', unsafe_allow_html=True)
                    st.plotly_chart(fig_mix, use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)
                    st.caption(f"{len(_all_ids_fb)} titles on {platform} share this math model (via GameCatalogView1.SkinOf), "
                              "despite having different names/themes — this is expected, not an error. Scoped to this "
                              "platform only; the same math model on another platform isn't included here.")

# ══════════════════════════════════════════════════════════════════
# TAB — SOCIAL CASINO
# ══════════════════════════════════════════════════════════════════
with tab_social:
    # Platform comes from the sidebar selector rather than a second control here.
    # Every floor now has a loyalty-account signal: PFH/EdgeLabs via AccountNumber,
    # V1/V2 via the PlayerAccountNumber loyalty-card tap on each terminal-day row.
    # When a specific game is selected (not "All Games"), scope every query to that
    # one game_id and default the date range to its own lifetime instead of a
    # generic window — "which games" isn't a useful breakdown anymore once there's
    # only one.
    soc_platform = platform
    soc_scoped = not is_all
    soc_game_id = int(sel_id) if soc_scoped else None

    if soc_scoped:
        ribbon("Loyalty & Members", f"{meta['game_name']} ({FLOOR_LABELS.get(soc_platform, soc_platform)}) · this game's "
                                 "player-level activity — Stickiness, DAU/MAU, ARPDAU.", T2, "LOYALTY")
    else:
        ribbon("Loyalty & Members", f"{FLOOR_LABELS.get(soc_platform, soc_platform)} · player-level activity — Stickiness, "
                                "DAU/MAU, ARPDAU, from each floor's loyalty-account data.", T2, "LOYALTY")

    _soc_key = f"_{soc_game_id}" if soc_scoped else "_all"
    _soc_default_start = meta["launch_date"] if soc_scoped else dt.date(2025, 1, 1)
    sf2, sf3 = st.columns(2)
    with sf2:
        soc_start = st.date_input("Start", value=_soc_default_start, key=f"soc_start{_soc_key}",
                                   help="Defaults to this game's launch date." if soc_scoped else None)
    with sf3:
        soc_end = st.date_input("End", value=dt.date.today(), key=f"soc_end{_soc_key}")

    soc_casino, soc_agg, soc_ccy = "All", "All", None
    if soc_platform == "EdgeLabs":
        try:
            _soc_casinos, _soc_aggs = _load_social_filters("EdgeLabs")
        except Exception:
            _soc_casinos, _soc_aggs = [], []
        sf4, sf5, sf6 = st.columns(3)
        with sf4:
            soc_casino = st.selectbox("Member Lounge", ["All"] + _soc_casinos, key="soc_casino")
        with sf5:
            soc_agg = st.selectbox("Aggregator", ["All"] + _soc_aggs, key="soc_agg")
        with sf6:
            # No "All" here on purpose — EdgeLabs mixes GC/SC/WOW in one column and
            # summing money across them is meaningless (GC alone reads as 23,825% RTP).
            _ccy_opts = ["SC", "GC", "WOW"]
            _ccy_idx = 0
            _ccy_note = None
            if soc_scoped:
                # Some games only ever ran in one currency (e.g. never had SC volume) --
                # defaulting to SC there just shows "no activity" for no real reason.
                # Prefer SC first, else whichever of GC/WOW the game actually has.
                try:
                    _avail_ccy = _load_social_game_currencies(soc_platform, soc_game_id)
                except Exception:
                    _avail_ccy = []
                if _avail_ccy and "SC" not in _avail_ccy:
                    for _c in _ccy_opts[1:]:
                        if _c in _avail_ccy:
                            _ccy_idx = _ccy_opts.index(_c)
                            _ccy_note = f"Defaulted to {_c} — this game has no SC activity."
                            break
            soc_ccy = st.selectbox("Currency", _ccy_opts, index=_ccy_idx, key=f"soc_ccy{_soc_key}",
                                    help="EdgeLabs records Gold Coins (GC), Sweeps Coins (SC) and WOW in one "
                                         "table. Money metrics are only coherent within a single currency, so "
                                         "one must be picked. GC is play-money — its RTP/ARPDAU aren't real economics.")
            if _ccy_note:
                st.caption(_ccy_note)

    if soc_start > soc_end:
        st.warning("Start date is after end date.")
    else:
        _s, _e = str(soc_start), str(soc_end)
        try:
            soc_daily = _load_social_daily(soc_platform, soc_casino, soc_agg, _s, _e, soc_ccy, soc_game_id)
            soc_mau = _load_social_mau(soc_platform, soc_casino, soc_agg, _s, _e, soc_ccy, soc_game_id)
            # "Players by game" is meaningless once already scoped to one game.
            soc_by_game = (_load_social_by_game(soc_platform, soc_casino, soc_agg, _s, _e, soc_ccy)
                           if not soc_scoped else pd.DataFrame())
            soc_by_casino = (_load_social_by_casino(soc_platform, soc_agg, _s, _e, soc_ccy, soc_game_id)
                             if soc_platform == "EdgeLabs" else pd.DataFrame())
            soc_totals = _load_social_totals(soc_platform, soc_casino, soc_agg, _s, _e, soc_ccy, soc_game_id)
        except Exception as e:
            st.error(f"Could not load social casino data: {e}")
            soc_daily = pd.DataFrame()
            soc_totals = {"players": 0, "bet": 0.0, "win": 0.0}

        if soc_daily.empty:
            _soc_empty_who = f"{meta['game_name']} ({soc_platform})" if soc_scoped else soc_platform
            st.info(f"No player activity for {_soc_empty_who} in this window.")
        else:
            mau_map = ({(int(r["yr"]), int(r["mo"])): int(r["mau"]) for _, r in soc_mau.iterrows()}
                       if not soc_mau.empty else {})
            soc_daily["mau"] = soc_daily["d"].apply(lambda x: mau_map.get((x.year, x.month)))
            soc_daily["stickiness"] = np.where(soc_daily["mau"].notna() & (soc_daily["mau"] > 0),
                                                soc_daily["players"] / soc_daily["mau"] * 100, np.nan)
            # DAU is a subset of that same month's MAU by definition, so this should never
            # exceed 100% -- but MAU here is a per-calendar-month distinct count while `d`
            # can span a game's very first partial month, so clip defensively anyway.
            soc_daily["stickiness"] = soc_daily["stickiness"].clip(upper=99.9)
            soc_daily["net"] = soc_daily["bet"] - soc_daily["win"]
            soc_daily["arpdau"] = np.where(soc_daily["players"] > 0, soc_daily["net"] / soc_daily["players"], np.nan)

            _soc_net_tot = soc_totals["bet"] - soc_totals["win"]
            _ccy_sub = soc_ccy if soc_ccy else None
            krow([
                {"label": "Distinct Players (window)", "value": f"{soc_totals['players']:,}"},
                {"label": "Net Revenue (window)", "value": _usd(_soc_net_tot), "sub": _ccy_sub},
                {"label": "Hold %", "value": _pct(((soc_totals["bet"] - soc_totals["win"]) / soc_totals["bet"] * 100) if soc_totals["bet"] > 0 else 0),
                 "sub": _ccy_sub},
                {"label": "Avg Daily Players", "value": f"{soc_daily['players'].mean():,.0f}"},
                {"label": "Peak Daily Players", "value": f"{soc_daily['players'].max():,.0f}"},
            ])

            # ── Player Segments — what kind of players make up this window ──
            try:
                player_bets = _load_player_bets(soc_platform, soc_casino, soc_agg, _s, _e, soc_ccy, soc_game_id)
            except Exception:
                player_bets = pd.DataFrame(columns=["player", "bet"])
            if not player_bets.empty:
                seg, thresholds = _segment_players(player_bets)
                ribbon("Player Segments", "Loyalty accounts in this window, tiered by total wagering", T2, "PLAYERS")
                seg_cols = st.columns(4)
                seg_colors = [T3, T2, GOLD, MOVE_DOWN]
                for col, row, clr in zip(seg_cols, seg.itertuples(), seg_colors):
                    col.markdown(
                        f'<div style="background:{SURFACE};border:1px solid {BORDER};border-left:3px solid {clr};'
                        f'border-radius:10px;padding:12px 14px">'
                        f'<div class="eyebrow">{row.segment}</div>'
                        f'<div style="font-family:{SERIF};font-size:26px;font-weight:600;color:{TEXT};margin-top:4px">{row.players:,}</div>'
                        f'<div style="font-size:13px;color:{T2};margin-top:2px">players</div>'
                        f'<div style="font-size:14px;color:{T2};margin-top:8px">{_usd(row.total_bet)} wagered '
                        f'<span style="color:{clr};font-weight:700">· {row.share_pct:.1f}% of wagering</span></div>'
                        f'</div>', unsafe_allow_html=True)
                if len(thresholds) == 3:
                    # st.caption renders markdown, and markdown treats a $...$ pair as LaTeX --
                    # with 4 dollar amounts in one caption those pair up and swallow the text
                    # between them, so every "$" here is escaped to render literally.
                    _esc = lambda v: _usd(v).replace("$", r"\$")
                    st.caption(f"Tiers by percentile of wagering in this window: New/Casual (bottom 50%, under "
                               f"{_esc(thresholds[0])}) · Regular (next 30%, up to {_esc(thresholds[1])}) · "
                               f"Premium (next 15%, up to {_esc(thresholds[2])}) · VIP/High Roller (top 5%, "
                               f"{_esc(thresholds[2])}+).")

            if soc_ccy == "GC":
                st.markdown(f'<div class="ann ann-a"><span class="ann-tag">WATCH</span><div class="ann-body">'
                            f'<strong>Gold Coins is play-money.</strong> Player counts are real, but Net Revenue, '
                            f'Hold % and ARPDAU are not real economics in GC — its win/bet ratio runs far above 100%. '
                            f'Switch Currency to SC for meaningful money metrics.</div></div>',
                            unsafe_allow_html=True)

            fig_stick = go.Figure(go.Scatter(x=soc_daily["d"], y=soc_daily["stickiness"], mode="lines",
                line=dict(color=T2, width=1.5), fill="tozeroy", fillcolor=_hex_alpha(T2, 0.15),
                hovertemplate="%{x|%d %b %Y}<br>Stickiness: %{y:.2f}%<extra></extra>"))
            fig_stick.update_yaxes(title_text="Stickiness (%)", ticksuffix="%")
            fig_stick.update_xaxes(title_text="Date")
            st.markdown('<div class="card"><div class="card-title">Stickiness (%) by Date — DAU ÷ MAU</div>', unsafe_allow_html=True)
            st.plotly_chart(plotly_base(fig_stick, h=320), use_container_width=True, theme=None)
            st.markdown('</div>', unsafe_allow_html=True)

            pc1, pc2 = st.columns(2)
            with pc1:
                fig_dau = go.Figure(go.Scatter(x=soc_daily["d"], y=soc_daily["players"], mode="lines",
                    line=dict(color=MOVE_DOWN, width=1.5), fill="tozeroy", fillcolor=_hex_alpha(MOVE_DOWN, 0.12),
                    hovertemplate="%{x|%d %b %Y}<br>Players: %{y:,.0f}<extra></extra>"))
                fig_dau.update_yaxes(title_text="# Players")
                fig_dau.update_xaxes(title_text="Date")
                st.markdown('<div class="card"><div class="card-title"># Players by Date</div>', unsafe_allow_html=True)
                st.plotly_chart(plotly_base(fig_dau, h=340), use_container_width=True, theme=None)
                st.markdown('</div>', unsafe_allow_html=True)
            with pc2:
                fig_arp = go.Figure(go.Scatter(x=soc_daily["d"], y=soc_daily["arpdau"], mode="lines",
                    line=dict(color=T2, width=1.3),
                    hovertemplate="%{x|%d %b %Y}<br>ARPDAU: $%{y:.2f}<extra></extra>"))
                fig_arp.update_yaxes(title_text="ARPDAU ($)")
                fig_arp.update_xaxes(title_text="Date")
                st.markdown('<div class="card"><div class="card-title">ARPDAU by Date — Net Revenue ÷ Daily Players</div>', unsafe_allow_html=True)
                st.plotly_chart(plotly_base(fig_arp, h=340), use_container_width=True, theme=None)
                st.markdown('</div>', unsafe_allow_html=True)

            if soc_scoped:
                # Already scoped to one game — "# Players by Game" would just be a single
                # bar, so it's dropped entirely. "# Players by Casino" is still useful for
                # EdgeLabs (which casino brands carry this specific game); PFH has none.
                if soc_platform == "EdgeLabs" and not soc_by_casino.empty:
                    top_c = soc_by_casino.head(12)
                    fig_c = go.Figure(go.Bar(x=top_c["players"], y=top_c["CasinoName"], orientation="h",
                        marker_color=MOVE_UP, text=[f"{v:,.0f}" for v in top_c["players"]], textposition="outside",
                        hovertemplate="<b>%{y}</b><br>Players: %{x:,.0f}<extra></extra>"))
                    fig_c.update_layout(yaxis=dict(autorange="reversed"))
                    st.markdown(f'<div class="card"><div class="card-title"># Players by Casino — {meta["game_name"]} '
                                f'(top {len(top_c)} of {len(soc_by_casino)})</div>', unsafe_allow_html=True)
                    st.plotly_chart(plotly_base(fig_c, h=340, ml=120), use_container_width=True, theme=None)
                    st.markdown('</div>', unsafe_allow_html=True)
            else:
                pc3, pc4 = st.columns(2)
                with pc3:
                    top_g = soc_by_game.head(10) if not soc_by_game.empty else soc_by_game
                    if not top_g.empty:
                        fig_g = go.Figure(go.Bar(x=top_g["players"], y=top_g["Name"].str[:24], orientation="h",
                            marker_color=T2, text=[f"{v:,.0f}" for v in top_g["players"]], textposition="outside",
                            hovertemplate="<b>%{y}</b><br>Players: %{x:,.0f}<extra></extra>"))
                        fig_g.update_layout(yaxis=dict(autorange="reversed"))
                        st.markdown(f'<div class="card"><div class="card-title"># Players by Game (top {len(top_g)} of {len(soc_by_game)})</div>', unsafe_allow_html=True)
                        st.plotly_chart(plotly_base(fig_g, h=340, ml=160), use_container_width=True, theme=None)
                        st.markdown('</div>', unsafe_allow_html=True)
                    else:
                        st.info("No per-game player data in this window.")
                with pc4:
                    if soc_platform == "EdgeLabs" and not soc_by_casino.empty:
                        top_c = soc_by_casino.head(12)
                        fig_c = go.Figure(go.Bar(x=top_c["players"], y=top_c["CasinoName"], orientation="h",
                            marker_color=MOVE_UP, text=[f"{v:,.0f}" for v in top_c["players"]], textposition="outside",
                            hovertemplate="<b>%{y}</b><br>Players: %{x:,.0f}<extra></extra>"))
                        fig_c.update_layout(yaxis=dict(autorange="reversed"))
                        st.markdown(f'<div class="card"><div class="card-title"># Players by Casino (top {len(top_c)} of {len(soc_by_casino)})</div>', unsafe_allow_html=True)
                        st.plotly_chart(plotly_base(fig_c, h=340, ml=120), use_container_width=True, theme=None)
                        st.markdown('</div>', unsafe_allow_html=True)
                    else:
                        st.markdown('<div class="card"><div class="card-title"># Players by Casino</div>', unsafe_allow_html=True)
                        st.caption("Only the Private Members' Club (EdgeLabs) tracks separate casino brands — "
                                  "nothing to break out here for this floor. Switch the sidebar Gaming Floor to "
                                  "Private Members' Club to see the 20+ casino-brand breakdown.")
                        st.markdown('</div>', unsafe_allow_html=True)

            st.caption(
                "Stickiness % = daily distinct players ÷ that calendar month's distinct players (DAU/MAU). "
                "ARPDAU = Net Revenue ÷ daily distinct players. VIP Kiosk Network/Private Members' Club identify "
                "players via AccountNumber (their online-account identity); the gaming floors (V1/V2) identify "
                "them via the loyalty card tapped at each terminal (PlayerAccountNumber).")

# ══════════════════════════════════════════════════════════════════
# TAB — COMPARE
# ══════════════════════════════════════════════════════════════════
with tab_compare:
    ribbon("Compare Games", "Select up to 6 games — parallel reporting, no cross-scoring", T2, "COMPARE")

    cmp_game_names = nonhr.sort_values("launch_date", ascending=False)["game_name"].tolist()
    cmp_name_to_id = dict(zip(nonhr["game_name"], nonhr["game_id"]))

    cc1, cc2 = st.columns([3, 2])
    with cc1:
        cmp_sel = st.multiselect("Pick 2–6 games to compare", options=cmp_game_names, max_selections=6,
                                 placeholder="Select 2–6 games…", key="cmp_sel_v2")
    with cc2:
        cmp_all_dates = pd.to_datetime(df["launch_date"]).dt.date
        cmp_date_min, cmp_date_max = cmp_all_dates.min(), pd.Timestamp.today().date()
        cmp_date_range = st.date_input("Date range", value=(cmp_date_min, cmp_date_max),
                                       min_value=cmp_date_min, max_value=cmp_date_max, key="cmp_dates_v2")
        cmp_from, cmp_to = cmp_date_range if isinstance(cmp_date_range, (list, tuple)) and len(cmp_date_range) == 2 else (cmp_date_min, cmp_date_max)

    if len(cmp_sel) < 2:
        st.info("Select at least 2 games above to generate a parallel report.")
    else:
        cmp_colors = [TEXT, MOVE_UP, "#7A6A9C", "#5B7A9C", "#9A7A3D", MOVE_DOWN]
        cmp_games = []
        for cn in cmp_sel:
            cid = cmp_name_to_id[cn]
            cgdf = df[df["game_id"] == cid].sort_values("launch_week").reset_index(drop=True)
            cgdf["_cal_date"] = pd.to_datetime(cgdf["launch_date"]) + pd.to_timedelta(cgdf["launch_week"] * 7, unit="d")
            cgdf = cgdf[(cgdf["_cal_date"].dt.date >= cmp_from) & (cgdf["_cal_date"].dt.date <= cmp_to)].reset_index(drop=True)
            cmeta = nonhr[nonhr["game_id"] == cid].iloc[0]
            cmp_games.append({"id": cid, "name": cn, "date": cmeta["launch_date"], "gdf": cgdf,
                              "weeks": int(cgdf["launch_week"].max()) if not cgdf.empty else 0})
        n_cmp = len(cmp_games)

        ribbon("Key Metrics", tag="METRICS")
        kpi_cols = st.columns(n_cmp)
        for ci, cg in enumerate(cmp_games):
            gd = cg["gdf"]
            tot_bet, tot_net = float(gd["bet_handle"].sum()), float(gd["net_rev"].sum())
            tot_hold = (tot_net / tot_bet * 100) if tot_bet > 0 else np.nan
            w0_b = float(gd.loc[gd["launch_week"] == 0, "bet_handle"].sum() or 0)
            rec_b = float(gd.tail(4)["bet_handle"].mean() or 0)
            locs = int(gd["stores"].max() or 0) if "stores" in gd.columns else 0
            with kpi_cols[ci]:
                st.markdown(
                    f'<div class="card" style="border-top:3px solid {cmp_colors[ci]}">'
                    f'<div style="font-size:14px;font-weight:700;color:{cmp_colors[ci]};margin-bottom:4px">{cg["name"]}</div>'
                    f'<div style="font-size:11px;color:{T3};margin-bottom:8px">Launched {cg["date"]} · {cg["weeks"]}w</div>'
                    f'<table style="width:100%;border-collapse:collapse;font-size:13px">'
                    f'<tr><td style="color:{T2};padding:3px 0">Total Bet Handle</td><td style="text-align:right;font-weight:700">{_usd(tot_bet)}</td></tr>'
                    f'<tr><td style="color:{T2};padding:3px 0">Total Net Revenue</td><td style="text-align:right;font-weight:700">{_usd(tot_net)}</td></tr>'
                    f'<tr><td style="color:{T2};padding:3px 0">Avg Hold %</td><td style="text-align:right;font-weight:700">{_pct(tot_hold)}</td></tr>'
                    f'<tr><td style="color:{T2};padding:3px 0">Launch Wk Bet</td><td style="text-align:right;font-weight:700">{_usd(w0_b)}</td></tr>'
                    f'<tr><td style="color:{T2};padding:3px 0">Recent 4-wk Avg</td><td style="text-align:right;font-weight:700">{_usd(rec_b)}</td></tr>'
                    f'<tr><td style="color:{T2};padding:3px 0">Peak Locations</td><td style="text-align:right;font-weight:700">{locs:,}</td></tr>'
                    f'</table></div>', unsafe_allow_html=True)

        fig_cmp_bh = go.Figure()
        for ci, cg in enumerate(cmp_games):
            if cg["gdf"].empty:
                continue
            fig_cmp_bh.add_trace(go.Scatter(x=cg["gdf"]["launch_week"], y=cg["gdf"]["bet_handle"], mode="lines",
                                            name=cg["name"], line=dict(color=cmp_colors[ci], width=2.2),
                                            hovertemplate=f'<b>{cg["name"]}</b><br>Wk %{{x}}<br>Bet: $%{{y:,.0f}}<extra></extra>'))
        fig_cmp_bh.update_xaxes(title_text="Weeks since launch")
        fig_cmp_bh.update_yaxes(title_text="Bet handle ($)")
        st.markdown('<div class="card"><div class="card-title">Weekly Bet Handle</div>', unsafe_allow_html=True)
        st.plotly_chart(plotly_base(fig_cmp_bh, h=450), use_container_width=True, theme=None)
        st.markdown('</div>', unsafe_allow_html=True)

        fig_cmp_nr = go.Figure()
        for ci, cg in enumerate(cmp_games):
            if cg["gdf"].empty:
                continue
            fig_cmp_nr.add_trace(go.Scatter(x=cg["gdf"]["launch_week"], y=cg["gdf"]["net_rev"], mode="lines",
                                            name=cg["name"], line=dict(color=cmp_colors[ci], width=2.2),
                                            hovertemplate=f'<b>{cg["name"]}</b><br>Wk %{{x}}<br>Net: $%{{y:,.0f}}<extra></extra>'))
        fig_cmp_nr.update_xaxes(title_text="Weeks since launch")
        fig_cmp_nr.update_yaxes(title_text="Net revenue ($)")
        st.markdown('<div class="card"><div class="card-title">Weekly Net Revenue</div>', unsafe_allow_html=True)
        st.plotly_chart(plotly_base(fig_cmp_nr, h=395), use_container_width=True, theme=None)
        st.markdown('</div>', unsafe_allow_html=True)

        ribbon("26-Week Forecast", tag="FORECAST")
        fcst_cols = st.columns(n_cmp)
        for ci, cg in enumerate(cmp_games):
            gd = cg["gdf"]
            if gd.empty:
                continue
            mwk = int(gd["launch_week"].max())
            cmp_fcst = L.fit_linear_trend(gd["bet_handle"].values, mwk)
            fw, fv, fhi, flo = cmp_fcst["weeks"], cmp_fcst["values"], cmp_fcst["upper"], cmp_fcst["lower"]
            fig_fc = go.Figure()
            fig_fc.add_trace(go.Scatter(x=gd["launch_week"], y=gd["bet_handle"], mode="lines", name="Actual",
                                        line=dict(color=cmp_colors[ci], width=2.2), hovertemplate="Wk %{x}<br>Bet: $%{y:,.0f}<extra></extra>"))
            if fw:
                fig_fc.add_trace(go.Scatter(x=fw + fw[::-1], y=fhi + flo[::-1], fill="toself",
                                            fillcolor="rgba(47,125,110,0.12)", line=dict(width=0), showlegend=False, hoverinfo="skip"))
                fig_fc.add_trace(go.Scatter(x=fw, y=fv, mode="lines", name="Forecast",
                                            line=dict(color=MOVE_UP, width=2, dash="dash"), hovertemplate="Wk %{x}<br>Forecast: $%{y:,.0f}<extra></extra>"))
            fig_fc.update_xaxes(title_text="Weeks since launch")
            fig_fc.update_yaxes(title_text="Bet handle ($)")
            with fcst_cols[ci]:
                st.markdown(f'<div class="card"><div class="card-title" style="color:{cmp_colors[ci]}">{cg["name"]}</div>', unsafe_allow_html=True)
                st.plotly_chart(plotly_base(fig_fc, h=345), use_container_width=True, theme=None)
                st.markdown('</div>', unsafe_allow_html=True)

        ribbon("On-Track Status", tag="STATUS")
        ot_cols = st.columns(n_cmp)
        CMP_KPIS = [("bet_handle", "Bet Handle"), ("net_rev", "Net Revenue"), ("hold_pct", "Hold %")]
        for ci, cg in enumerate(cmp_games):
            gd = cg["gdf"]
            with ot_cols[ci]:
                rows_html = f'<div class="card"><div style="font-size:14px;font-weight:700;color:{cmp_colors[ci]};margin-bottom:10px">{cg["name"]}</div>'
                for kp, kp_lbl in CMP_KPIS:
                    if kp not in gd.columns or gd.empty:
                        continue
                    bnd = bkpi(df, kp)
                    cur = float(gd.tail(4)[kp].mean() or 0)
                    mwk = int(gd["launch_week"].max()) if not gd.empty else 0
                    bnd_row = bnd[bnd["launch_week"] == mwk] if not bnd.empty else pd.DataFrame()
                    if not bnd_row.empty:
                        p25, p50, p75 = float(bnd_row["p25"].iloc[0]), float(bnd_row["p50"].iloc[0]), float(bnd_row["p75"].iloc[0])
                        pvs = ((cur - p50) / p50 * 100) if p50 > 0 else np.nan
                        kind = "g" if cur >= p75 else ("a" if cur >= p25 else "r")
                        lbl2 = "Above peers" if cur >= p75 else ("In line" if cur >= p25 else "Below peers")
                        vs2 = f"{pvs:+.0f}% vs p50" if not np.isnan(pvs) else ""
                    else:
                        kind, lbl2, vs2 = "n", "No data", ""
                    rows_html += (f'<div style="display:flex;justify-content:space-between;padding:4px 0;border-bottom:1px solid {BORDER}">'
                                 f'<span style="color:{T2};font-size:12px">{kp_lbl}</span>'
                                 f'<span style="font-size:12px">{badge(lbl2, kind)} <span style="color:{T3}">{vs2}</span></span></div>')
                rows_html += "</div>"
                st.markdown(rows_html, unsafe_allow_html=True)

        ribbon("Player Retention Trend", tag="RETENTION")
        bnd_decay2 = bkpi(df, "bet_decay")
        fh2 = go.Figure()
        if not bnd_decay2.empty:
            fh2.add_trace(go.Scatter(x=bnd_decay2["launch_week"], y=bnd_decay2["p50"], mode="lines",
                                     line=dict(color=T3, width=1.5, dash="dot"), name="Typical game (P50)", hoverinfo="skip"))
        for ci, cg in enumerate(cmp_games):
            gd = cg["gdf"]
            if gd.empty or "bet_decay" not in gd.columns:
                continue
            fh2.add_trace(go.Scatter(x=gd["launch_week"], y=gd["bet_decay"], mode="lines",
                                     line=dict(color=cmp_colors[ci], width=2.2), name=cg["name"]))
        fh2.update_xaxes(title_text="Launch week")
        fh2.update_yaxes(title_text="Bet decay % (week 0 = 100)")
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.plotly_chart(plotly_base(fh2, h=395), use_container_width=True, theme=None)
        st.markdown('</div>', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════
# TAB — GAME CLUSTERS
# ══════════════════════════════════════════════════════════════════
with tab_clusters:
    # Always the full property catalog, every system combined — a cluster map is
    # far more useful (and interesting) grouping the whole 167-game portfolio by
    # performance pattern than whatever single system happens to be in scope.
    ribbon("Game Clusters", f"All {len(full_nonhr)} games across the property, grouped by performance pattern", T2, "CLUSTERS")
    st.caption("Games are grouped using K-means clustering (k=4) on 7 performance features "
              "(hold %, hold trend, retention, bet decay, ARPU, launch scale, revenue acceleration).")

    feat_df = L.compute_cluster_features(full_df, full_nonhr)
    feat_cols = ["avg_hold", "hold_slope", "retention_ratio", "avg_decay", "avg_arpu", "launch_scale", "rev_accel"]
    feat_df = feat_df.dropna(subset=["avg_hold", "avg_decay", "launch_scale"], how="any")

    if len(feat_df) < 6:
        st.info("Not enough games with sufficient data for clustering (need at least 6). Try a platform with more launches.")
    else:
        X_raw = feat_df[feat_cols].copy()
        for c in feat_cols:
            X_raw[c] = X_raw[c].fillna(X_raw[c].median())
        X_raw["avg_hold"] = X_raw["avg_hold"].clip(-50, 50)
        mu_, std_ = X_raw.values.mean(axis=0), X_raw.values.std(axis=0) + 1e-9
        X_norm = (X_raw.values - mu_) / std_

        k_actual = min(4, len(feat_df))
        labels = L.run_kmeans(X_norm, k=k_actual)
        feat_df["cluster"] = labels
        cluster_avgs = feat_df.groupby("cluster")[["avg_hold", "retention_ratio", "launch_scale"]].mean()
        sorted_clusters = cluster_avgs.sort_values(["avg_hold", "retention_ratio"], ascending=False).index.tolist()
        label_map = {orig: new for new, orig in enumerate(sorted_clusters)}
        feat_df["cluster_idx"] = feat_df["cluster"].map(label_map)

        sel_game_name = None if is_all else meta["game_name"]
        sel_cluster_idx = None
        if sel_game_name and sel_game_name in feat_df["game_name"].values:
            sel_cluster_idx = int(feat_df[feat_df["game_name"] == sel_game_name]["cluster_idx"].iloc[0])

        cols_cards = st.columns(k_actual)
        for ci in range(k_actual):
            cmeta = _CLUSTER_META[ci]
            cg = feat_df[feat_df["cluster_idx"] == ci]
            n, ah, ar = len(cg), cg["avg_hold"].mean(), cg["retention_ratio"].mean()
            is_sel_cl = (sel_cluster_idx == ci)
            border = f"border:2px solid {T2}" if is_sel_cl else f"border:1px solid {BORDER}"
            sel_tag = f'<div style="font-size:10px;font-weight:700;color:{T2};margin-bottom:4px">▶ {sel_game_name[:22]}</div>' if is_sel_cl else ""
            with cols_cards[ci]:
                st.markdown(f"""
<div style="background:{SURFACE};{border};border-top:4px solid {cmeta['color']};border-radius:10px;padding:18px 14px;text-align:center;min-height:170px">
  {sel_tag}
  <div style="font-size:15px;font-weight:700;color:{cmeta['color']};margin-bottom:6px">{cmeta['name']}</div>
  <div style="font-size:12px;color:{T2};margin-bottom:8px">{n} game{'s' if n != 1 else ''}</div>
  <div style="font-size:12px;color:{T2};margin-bottom:3px">Avg Hold: <strong style="color:{TEXT}">{ah:.1f}%</strong></div>
  <div style="font-size:12px;color:{T2}">Retention: <strong style="color:{TEXT}">{ar:.2f}x</strong> <span style="font-size:10px">(Wk4/Wk1)</span></div>
</div>""", unsafe_allow_html=True)

        scatter_fig = go.Figure()
        for ci in range(k_actual):
            cmeta = _CLUSTER_META[ci]
            cg = feat_df[feat_df["cluster_idx"] == ci]
            cg_sel = cg[cg["game_name"] == sel_game_name] if sel_game_name else cg.iloc[0:0]
            cg_rest = cg[cg["game_name"] != sel_game_name] if sel_game_name else cg
            scatter_fig.add_trace(go.Scatter(
                x=cg_rest["avg_hold"], y=cg_rest["avg_arpu"], mode="markers", name=cmeta["name"],
                marker=dict(size=9, color=cmeta["color"], opacity=0.8),
                customdata=cg_rest[["game_name", "avg_hold", "retention_ratio", "avg_arpu", "weeks_live"]].values,
                hovertemplate="<b>%{customdata[0]}</b><br>Avg Hold: %{customdata[1]:.1f}%<br>Avg ARPU: $%{customdata[3]:.2f}<br>"
                             "Wk4/Wk1: %{customdata[2]:.2f}x<br>Weeks Live: %{customdata[4]}<extra></extra>"))
            if not cg_sel.empty:
                scatter_fig.add_trace(go.Scatter(
                    x=cg_sel["avg_hold"], y=cg_sel["avg_arpu"], mode="markers+text", name=f"★ {sel_game_name}",
                    marker=dict(size=15, color=TEXT, symbol="star"), text=cg_sel["game_name"].str[:22], textposition="top center",
                    customdata=cg_sel[["game_name", "avg_hold", "retention_ratio", "avg_arpu", "weeks_live"]].values,
                    hovertemplate="<b>%{customdata[0]}</b> ★<br>Avg Hold: %{customdata[1]:.1f}%<br>Avg ARPU: $%{customdata[3]:.2f}<extra></extra>"))
        med_hold, med_arpu = float(feat_df["avg_hold"].median()), float(feat_df["avg_arpu"].median())
        scatter_fig.add_vline(x=med_hold, line_dash="dot", line_color=BORDER, line_width=1)
        scatter_fig.add_hline(y=med_arpu, line_dash="dot", line_color=BORDER, line_width=1)
        scatter_fig.update_xaxes(title_text="Avg hold % (higher = earns more per dollar wagered)")
        scatter_fig.update_yaxes(title_text="Avg ARPU $ (higher = players spend more)")
        st.markdown('<div class="card"><div class="card-title">Game Performance Cluster Map — Hold % vs. Player Spend (ARPU)</div>', unsafe_allow_html=True)
        st.plotly_chart(plotly_base(scatter_fig, h=630), use_container_width=True, theme=None)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown(f'<div style="font-size:14px;font-weight:700;color:{TEXT};margin:20px 0 10px">Games by Cluster</div>', unsafe_allow_html=True)
        for ci in range(k_actual):
            cmeta = _CLUSTER_META[ci]
            cg = feat_df[feat_df["cluster_idx"] == ci].sort_values("avg_hold", ascending=False)
            has_sel = sel_game_name in cg["game_name"].values if sel_game_name else False
            exp_label = f"{cmeta['name']} — {len(cg)} games — {cmeta['desc']}" + (f" · ★ {sel_game_name}" if has_sel else "")
            with st.expander(exp_label, expanded=(ci == 0 or has_sel)):
                disp = cg[["game_name", "mechanic", "avg_hold", "hold_slope", "retention_ratio", "avg_decay", "avg_arpu", "weeks_live"]].copy()
                disp.columns = ["Game", "Mechanic", "Avg Hold %", "Hold Trend", "Retention (Wk4/Wk1)", "Avg Bet Decay %", "Avg ARPU ($)", "Weeks Live"]
                st.dataframe(
                    disp.style.format({"Avg Hold %": "{:.1f}%", "Hold Trend": "{:.2f}", "Retention (Wk4/Wk1)": "{:.2f}x",
                                      "Avg Bet Decay %": "{:.1f}%", "Avg ARPU ($)": "${:.2f}"}, na_rep="–"),
                    use_container_width=True, hide_index=True,
                )

